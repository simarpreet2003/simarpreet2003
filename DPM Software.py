# EMAIL    - skdentalcare9988@gmail.com
# PASSWORD - ravjotskdc
# SECURITY OFF - https://myaccount.google.com/u/1/lesssecureapps?gar=1&pli=1&rapt=AEjHL4NdVGsjOGDeLDWHSR327F7F2M7KL1V6xRrx_IoLmm2aRzvpLx9Wa_Q6toKYZFRh3JC302JuKjIHqpx4SzlyGLtFXC4srw


############################### ADD  PYTHON  AS  ENVIRONMENT  VARIABLE  ################################

#                                    -> CONTROL PANEL - SYSTEM & SECURITY - SYSTEM -
#                                       ADVANCED SYSTEM SETTINGS - ENVIRONMENT VARIABLES -
#                                       SYSTEM VARIABLES - DOUBLE CLICK ON PATH - NEW -
#          ADD THESE 2 [                C:\Users\Acer\AppData\Local\Programs\Python\Python38-32\Scripts 
#                      [                C:\Users\Acer\AppData\Local\Programs\Python\Python38-32
# add to USER VARIABLES FOR ACER also........


#########################################  EXECUTE  ON  CMD   ##########################################

# python -m pip install pip;
# python -m pip install pillow;
# python -m pip install --upgrade pip;                     [20.0.2 for mine]
# python -m pip install pip install tkfilebrowser;
# python -m pip install mysql-connector-python;
# python -m pip install pywin32;
# python -m pip install pypiwin32;
# python -m pip install tkfilebrowser;
# python -m pip install temp;
# python -m pip install python-docx;
# python -m pip install email-to;
# python -m pip install paphra-tktable;
# python -m pip install pyscreenshot;
# pip install tkcalendar;



from tkinter import*
import tkinter as tk
from tkcalendar import*
from PIL import ImageTk, Image
from tkfilebrowser import askopenfilename
import tkinter.messagebox as MessageBox
import mysql.connector as mysql
import webbrowser
from tkinter import ttk
import win32api
import win32print
import tempfile
from tkinter import filedialog as fd
import time
from tkinter.scrolledtext import ScrolledText
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import smtplib
import math, random
from email.mime.text import MIMEText
import tkinter.scrolledtext as tkscrolled
from random import choice









#################################  TOPLEVEL   WINDOW   NEW   PASSWORD  SET #############################

def  top_for_FP():                                  
    def updatePASS():
        
        if a.get()!=b.get():
            MessageBox.showerror("PASSWORD RESET","Passwords not Matched",parent=top)

        else:  
            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
            cursor=con.cursor()
            cursor.execute("update c_hr set password='"+b.get()+"' ");
            cursor.execute("commit");
            con.close()
            MessageBox.showinfo("PASSWORD UPDATE ","Password Updated Successfully",parent=top)
            top.destroy()
            
    if OTPTYPE.get()==OTP:
        top=Toplevel()
        top.title('Generate Password ')
        top.geometry('500x150')
        top.config(bg='LIGHTCYAN')
        OTPTYPE.delete(0,END)
        C=Label(top,text='Enter New Password',font=("ROBOTO BOLD",15),bg='WHITE',fg='BLACK')
        C.place(x=10,y=10)
        a=Entry(top, show="*",font=("ARIAL",13))
        a.place(x=230,y=15)
        D=Label(top,text=' Confirm Password ',font=("ROBOTO BOLD",15),bg='WHITE',fg='BLACK')
        D.place(x=10,y=50)
        b=Entry(top,show="*",font=("ARIAL",13))
        b.place(x=230,y=55)
        buttonsubmit = Button(top, text=" Submit ",font=("ROBOTO BOLD",15),bg='WHITE',fg='BLACK',command=updatePASS)
        buttonsubmit.place(x=190,y=100)
        
    else:
        MessageBox.showerror("PASSWORD UPDATE ","WRONG OTP",parent=root)

        

    
    
    
digits = "0123456789"                      # OTP  GENERATE
OTP = ""   
for i in range(4) : 
    OTP = OTP + digits[math.floor(random.random() * 10)]     


########################################################################################################
###########################################   EMAIL   SEND  ############################################

def  mail():
    
    try:
        mailServer = smtplib.SMTP('smtp.gmail.com' , 587)
        mailServer.starttls()
        SEND='skdentalcare9988@gmail.com'
        PASS='ravjotskdc'
        RECIEVE='ravjotsingh.co@gmail.com'
        msg = MIMEText(OTP+'  is your OTP for password change request. Please enter this to verify your identity and proceed to change the password.')

        msg['Subject'] = 'Reset Your Software Password'
        msg['From'] = 'skdentalcare9988@gmail.com'
        msg['To'] = 'ravjotsingh.co@gmail.com'
        mailServer.login(SEND , PASS)
        mailServer.sendmail(SEND, RECIEVE ,msg.as_string())
        mailServer.quit()

        MessageBox.showinfo("OTP STATUS","OTP Sent Successfully at your mail - 'ravjotsingh.co@gmail.com' ",parent=root)

    except:
        MessageBox.showerror("OTP STATUS","   Access denied!  ",parent=root)
        


####################### ####################### ####################### ####################### ########
#######################     ON   CLICKING   LOGIN   BUTTON   ON  THE  ROOT  WINDOW ##################   



def login():
    
    DROPDOWNget=DROPDOWN.get()                       # IF USERNAME AND PASSWORD SAME AS MYSQL THEN LOGIN                   
    con = mysql.connect(host="localhost",user="root",
        password="admin",database="patientsoftware1")
    cursor=con.cursor()
    cursor.execute("select* from c_hr ");
    for row in cursor:
        username=str(row[0])
        password=str(row[1])
    if DROPDOWNget==username and passw.get()==password:
            
        top=Toplevel()                                 # TOP LEVEL WINDOW AFTER LOGIN WITH NOTEBOOK
        top.title('SK DENTAL CLINIC ~ ADMIN PORTAL ')
        top.geometry('3000x1200')
        top.configure(bg='WHITE')
        top.state('zoomed')
        passw.delete(0,'end')
        root.withdraw()
        
        nb = ttk.Notebook(top)                         #######   NOTEBOOK   #######

        HOME          = ttk.Frame(nb)        
        ADDPATIENT    = ttk.Frame(nb)
        SEARCHPATIENT = ttk.Frame(nb)
        APPOINTMENTS   = ttk.Frame(nb)
        ss = ttk.Style()
        ss.theme_use('clam')


        nb.add(HOME, text='       DASHBOARD       ')
        nb.add(ADDPATIENT, text='       ADD PATIENT        ')
        nb.add(SEARCHPATIENT, text='       OPERATIONS        ')
        nb.add(APPOINTMENTS, text='       APPOINTMENTS        ')
        nb.pack(expand=1, fill="both")
        
        
        
######################################  LOG  OUT  DESTROY  TOPLEVEL   ################################

        def logout():
            MsgBox = MessageBox.askquestion ('LOGOUT','Are you sure you want to logout?',icon = 'warning')
            if MsgBox == 'yes':
               top.destroy()
               root.deiconify()
            
                
        
########################################################################################################
#################################  CANVAS   FOR    HOME    TAB  DASHBOARD   ############################



        canvas_HOME = Canvas(HOME, width = 800, height = 400)            # CANVAS FORMED
        canvas_HOME.configure(bg='white')
        canvas_HOME.pack(side = LEFT, fill = BOTH, expand = True)
        

        
        entryContainezr = Frame(canvas_HOME)                               # FRAME FOR CANVAS
        entryContainezr.pack()



      
        
         
        gif1sX = Image.open("C:/Users/Acer/Desktop/pms files/ImageBlur (5).png")    # HOME BG IMAGE
        gif1sX = gif1sX.resize((1530, 765))
        imageDs = ImageTk.PhotoImage(gif1sX)
        label2g = tk.Label(canvas_HOME,image=imageDs)
        label2g.image=imageDs
        canvas_HOME.create_window((0, 0), anchor=NW, window=label2g)
        imageDs.image = imageDs


        wpotoiQZmage = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/Untitled.png")
        aj=Button(canvas_HOME, text = '', image = wpotoiQZmage,bg='white',state='disabled')
        aj.image=wpotoiQZmage        
        canvas_HOME.create_window((8, 15), anchor=NW, window=aj)
        wpotoiQZmage.image = wpotoiQZmage
        
        potoiQmagez = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/logout_icon.png")# LOGOUT BUTTON
        potoiQmagez = potoiQmagez.subsample(16, 16) 
        ajA=Button(canvas_HOME, text = '', image = potoiQmagez,bg='white',command=logout)
        ajA.image=potoiQmagez 
        canvas_HOME.create_window((1450, 10), anchor=NW, window=ajA)
        potoiQmagez.image = potoiQmagez


        variablexc = StringVar(canvas_HOME)              # DROPDOWN  FOR  SELECT  MONTH
        variablexc.set("JANUARY")
        wsd = OptionMenu(canvas_HOME, variablexc, "JANUARY", "FEBRUARY",'MARCH','APRIL','MAY','JUNE','JULY','AUGUST','SEPTEMBER','OCTOBER','NOVEMBER','DECEMBER')
        wsd.config(bg = "darkturquoise",fg='WHITE',font=('Bahnschrift Light SemiCondensed',23))
        canvas_HOME.create_window((550, 170), anchor=NW, window=wsd)


        variablexca = StringVar(canvas_HOME)              # DROPDOWN  FOR  SELECT  YEAR
        variablexca.set("2020")
        wsda = OptionMenu(canvas_HOME, variablexca, "2019", "2020",'2021','2022','2023','2024')
        wsda.config(bg = "darkturquoise",fg='WHITE',font=('Bahnschrift Light SemiCondensed',23))
        canvas_HOME.create_window((797, 170), anchor=NW, window=wsda)


    
            
        def month():                                        #MONTHWISE  PATIENTS  NUMBER

            mt=Text(HOME,font=('CAMBRIA',28),fg='white',bg='darkturquoise',width=6,height=1.3)
            mt.place(x=790,y=277)
            mt.config(insertbackground='darkturquoise')
            
#JAN ALL YEARS
            if variablexc.get()=='JANUARY' and variablexca.get()=='2019':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2019/01/01' and date <='2019/01/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ') 

                con.close()


            elif variablexc.get()=='JANUARY' and variablexca.get()=='2020':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2020/01/01' and date <='2020/01/31'")
                result=cursor.fetchone()
                mt.insert(END,result)

            if int(len(mt.get('1.0',END)))==2:
                mt.insert('1.0','      ')
            elif int(len(mt.get('1.0',END)))==3:
                mt.insert('1.0','     ')
            elif int(len(mt.get('1.0',END)))==4:
                mt.insert('1.0','    ')
            elif int(len(mt.get('1.0',END)))==5:
                mt.insert('1.0','   ')
            elif int(len(mt.get('1.0',END)))==6:
                mt.insert('1.0','  ')
                
                con.close()

            elif variablexc.get()=='JANUARY' and variablexca.get()=='2021':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2021/01/01' and date <='2021/01/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='JANUARY' and variablexca.get()=='2022':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2022/01/01' and date <='2022/01/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='JANUARY' and variablexca.get()=='2023':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2023/01/01' and date <='2023/01/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='JANUARY' and variablexca.get()=='2024':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2024/01/01' and date <='2024/01/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

#FEB ALL YEARS

            elif variablexc.get()=='FEBRUARY' and variablexca.get()=='2019':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2019/02/01' and date <='2019/02/28'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()


            elif variablexc.get()=='FEBRUARY' and variablexca.get()=='2020':
                            
                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2020/02/01' and date <='2020/02/29'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='FEBRUARY' and variablexca.get()=='2021':
               
                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2021/02/01' and date <='2021/02/29'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='FEBRUARY' and variablexca.get()=='2022':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2022/02/01' and date <='2022/02/29'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='FEBRUARY' and variablexca.get()=='2023':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2023/02/01' and date <='2023/02/29'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='FEBRUARY' and variablexca.get()=='2024':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2024/02/01' and date <='2024/02/28'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()


# MARCH ALL YEARS


            elif variablexc.get()=='MARCH' and variablexca.get()=='2019':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2019/03/01' and date <='2019/03/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()


            elif variablexc.get()=='MARCH' and variablexca.get()=='2020':
                            
                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2020/03/01' and date <='2020/03/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='MARCH' and variablexca.get()=='2021':
               
                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2021/03/01' and date <='2021/03/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='MARCH' and variablexca.get()=='2022':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2022/03/01' and date <='2022/03/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='MARCH' and variablexca.get()=='2023':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2023/03/01' and date <='2023/03/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='MARCH' and variablexca.get()=='2024':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2024/03/01' and date <='2024/03/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()


# APRIL ALL YEARS


            elif variablexc.get()=='APRIL' and variablexca.get()=='2019':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2019/04/01' and date <='2019/04/30'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()


            elif variablexc.get()=='APRIL' and variablexca.get()=='2020':
                            
                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2020/04/01' and date <='2020/04/30'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='APRIL' and variablexca.get()=='2021':
               
                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2021/04/01' and date <='2021/04/30'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='APRIL' and variablexca.get()=='2022':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2022/04/01' and date <='2022/04/30'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='APRIL' and variablexca.get()=='2023':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2023/04/01' and date <='2023/04/30'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='APRIL' and variablexca.get()=='2024':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2024/04/01' and date <='2024/04/30'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()



# MAY ALL YEARS


            elif variablexc.get()=='MAY' and variablexca.get()=='2019':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2019/05/01' and date <='2019/05/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()


            elif variablexc.get()=='MAY' and variablexca.get()=='2020':
                            
                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2020/05/01' and date <='2020/05/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='MAY' and variablexca.get()=='2021':
               
                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2021/05/01' and date <='2021/05/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='MAY' and variablexca.get()=='2022':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2022/05/01' and date <='2022/05/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='MAY' and variablexca.get()=='2023':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2023/05/01' and date <='2023/05/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='MAY' and variablexca.get()=='2024':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2024/05/01' and date <='2024/05/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()



# JUNE ALL YEARS


            elif variablexc.get()=='JUNE' and variablexca.get()=='2019':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2019/06/01' and date <='2019/06/30'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()


            elif variablexc.get()=='JUNE' and variablexca.get()=='2020':
                            
                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2020/06/01' and date <='2020/06/30'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='JUNE' and variablexca.get()=='2021':
               
                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2021/06/01' and date <='2021/06/30'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='JUNE' and variablexca.get()=='2022':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2022/06/01' and date <='2022/06/30'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='JUNE' and variablexca.get()=='2023':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2023/06/01' and date <='2023/06/30'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='JUNE' and variablexca.get()=='2024':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2024/06/01' and date <='2024/06/30'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()



# JULY ALL YEARS


            elif variablexc.get()=='JULY' and variablexca.get()=='2019':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2019/07/01' and date <='2019/07/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()


            elif variablexc.get()=='JULY' and variablexca.get()=='2020':
                            
                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2020/07/01' and date <='2020/07/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='JULY' and variablexca.get()=='2021':
               
                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2021/07/01' and date <='2021/07/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='JULY' and variablexca.get()=='2022':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2022/07/01' and date <='2022/07/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='JULY' and variablexca.get()=='2023':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2023/07/01' and date <='2023/07/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='JULY' and variablexca.get()=='2024':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2024/07/01' and date <='2024/07/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()



# AUG ALL YEARS


            elif variablexc.get()=='AUGUST' and variablexca.get()=='2019':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2019/08/01' and date <='2019/08/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()


            elif variablexc.get()=='AUGUST' and variablexca.get()=='2020':
                            
                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2020/08/01' and date <='2020/08/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='AUGUST' and variablexca.get()=='2021':
               
                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2021/08/01' and date <='2021/08/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='AUGUST' and variablexca.get()=='2022':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2022/08/01' and date <='2022/08/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='AUGUST' and variablexca.get()=='2023':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2023/08/01' and date <='2023/08/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='AUGUST' and variablexca.get()=='2024':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2024/08/01' and date <='2024/08/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()



# SEP ALL YEARS


            elif variablexc.get()=='SEPTEMBER' and variablexca.get()=='2019':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2019/09/01' and date <='2019/09/30'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()


            elif variablexc.get()=='SEPTEMBER' and variablexca.get()=='2020':
                            
                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2020/09/01' and date <='2020/09/30'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='SEPTEMBER' and variablexca.get()=='2021':
               
                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2021/09/01' and date <='2021/09/30'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='SEPTEMBER' and variablexca.get()=='2022':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2022/09/01' and date <='2022/09/30'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='SEPTEMBER' and variablexca.get()=='2023':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2023/09/01' and date <='2023/09/30'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='SEPTEMBER' and variablexca.get()=='2024':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2024/09/01' and date <='2024/09/30'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()



# OCT ALL YEARS


            elif variablexc.get()=='OCTOBER' and variablexca.get()=='2019':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2019/10/01' and date <='2019/10/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()


            elif variablexc.get()=='OCTOBER' and variablexca.get()=='2020':
                            
                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2020/10/01' and date <='2020/10/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='OCTOBER' and variablexca.get()=='2021':
               
                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2021/10/01' and date <='2021/10/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='OCTOBER' and variablexca.get()=='2022':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2022/10/01' and date <='2022/10/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='OCTOBER' and variablexca.get()=='2023':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2023/10/01' and date <='2023/10/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='OCTOBER' and variablexca.get()=='2024':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2024/10/01' and date <='2024/10/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()



# NOV ALL YEARS


            elif variablexc.get()=='NOVEMBER' and variablexca.get()=='2019':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2019/03/01' and date <='2019/03/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()


            elif variablexc.get()=='NOVEMBER' and variablexca.get()=='2020':
                            
                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2020/11/01' and date <='2020/11/30'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='NOVEMBER' and variablexca.get()=='2021':
               
                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2021/11/01' and date <='2021/11/30'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='NOVEMBER' and variablexca.get()=='2022':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2022/11/01' and date <='2022/11/30'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='NOVEMBER' and variablexca.get()=='2023':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2023/11/01' and date <='2023/11/30'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='NOVEMBER' and variablexca.get()=='2024':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2024/11/01' and date <='2024/11/30'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()



# DEC ALL YEARS


            elif variablexc.get()=='DECEMBER' and variablexca.get()=='2019':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2019/12/01' and date <='2019/12/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()


            elif variablexc.get()=='DECEMBER' and variablexca.get()=='2020':
                            
                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2020/12/01' and date <='2020/12/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='DECEMBER' and variablexca.get()=='2021':
               
                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2021/12/01' and date <='2021/12/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='DECEMBER' and variablexca.get()=='2022':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2022/12/01' and date <='2022/12/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='DECEMBER' and variablexca.get()=='2023':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2023/12/01' and date <='2023/12/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            elif variablexc.get()=='DECEMBER' and variablexca.get()=='2024':

                mt.delete('1.0',END)
            
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select count(*) from c_h1 where date >= '2024/12/01' and date <='2024/12/31'")
                result=cursor.fetchone()
                mt.insert(END,result)
                if int(len(mt.get('1.0',END)))==2:
                    mt.insert('1.0','      ')
                elif int(len(mt.get('1.0',END)))==3:
                    mt.insert('1.0','     ')
                elif int(len(mt.get('1.0',END)))==4:
                    mt.insert('1.0','    ')
                elif int(len(mt.get('1.0',END)))==5:
                    mt.insert('1.0','   ')
                elif int(len(mt.get('1.0',END)))==6:
                    mt.insert('1.0','  ')

                con.close()

            widgetssz = Button(canvas_HOME, text='  TOTAL PATIENTS  ',font=('Bahnschrift Light SemiCondensed',21), state='disabled', disabledforeground='white',bg='darkturquoise')
            widgetssz.pack()
            canvas_HOME.create_window(630, 300, window=widgetssz)

            def PH():
                mt.destroy()
                widgetssz.destroy()
                widgetclear.destroy()
                widgethidw.destroy()
                
            def clear():
                mt.delete('1.0',END)


            widgetclear = Button(canvas_HOME, text='  CLEAR  ',font=('Bahnschrift Light SemiCondensed',11), fg='white',bg='dodgerblue',command=clear)
            widgetclear.pack()
            canvas_HOME.create_window(970, 300, window=widgetclear)

            widgethidw = Button(canvas_HOME, text='   HIDE   ',font=('Bahnschrift Light SemiCondensed',12), fg='white',bg='dodgerblue',command=PH)
            widgethidw.pack()
            canvas_HOME.create_window(970, 195, window=widgethidw)
                
        widget = Button(canvas_HOME, text='  SHOW  ',font=('Bahnschrift Light SemiCondensed ',12), fg='white',bg='dodgerblue',command=month)
        widget.pack()
        canvas_HOME.create_window(970, 195, window=widget)


            
        
######  defining  dashboard  commands
            
        
        def showdashboard():
         
#ACTIVE           
            ac=Text(HOME,font=('CAMBRIA',25),fg='white',bg='darkturquoise',width=6,height=1.3)
            ac.config(insertbackground='darkturquoise')
            ac.place(x=320,y=175)

            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
            cursor=con.cursor()
            cursor.execute("select count(patient_status) from c_h1 where patient_status='active'")
            result=cursor.fetchone()
            ac.insert(END,result)
            if int(len(ac.get('1.0',END)))==2:
                ac.insert('1.0','      ')
            elif int(len(ac.get('1.0',END)))==3:
                ac.insert('1.0','     ')
            elif int(len(ac.get('1.0',END)))==4:
                ac.insert('1.0','    ')
            elif int(len(ac.get('1.0',END)))==5:
                ac.insert('1.0','   ')
            elif int(len(ac.get('1.0',END)))==6:
                ac.insert('1.0','  ') 
            con.close()

           
#RECOVERED            
            bc=Text(HOME,font=('CAMBRIA',25),fg='white',bg='darkturquoise',width=6,height=1.3)
            bc.place(x=320,y=275)
            bc.config(insertbackground='darkturquoise')

            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
            cursor=con.cursor()
            cursor.execute("select count(patient_status) from c_h1 where patient_status='recovered'")
            result=cursor.fetchone()
            bc.insert(END,result)
            if int(len(bc.get('1.0',END)))==2:
                bc.insert('1.0','      ')
            elif int(len(bc.get('1.0',END)))==3:
                bc.insert('1.0','     ')
            elif int(len(bc.get('1.0',END)))==4:
                bc.insert('1.0','    ')
            elif int(len(bc.get('1.0',END)))==5:
                bc.insert('1.0','   ')
            elif int(len(bc.get('1.0',END)))==6:
                bc.insert('1.0','  ')
            con.close()

#REFERRED            
            ec=Text(HOME,font=('CAMBRIA',25),fg='white',bg='darkturquoise',width=6,height=1.3)
            ec.place(x=320,y=375)
            ec.config(insertbackground='darkturquoise')

            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
            cursor=con.cursor()
            cursor.execute("select count(patient_status) from c_h1 where patient_status='referred'")
            result=cursor.fetchone()
            ec.insert(END,result)
            if int(len(ec.get('1.0',END)))==2:
                ec.insert('1.0','      ')
            elif int(len(ec.get('1.0',END)))==3:
                ec.insert('1.0','     ')
            elif int(len(ec.get('1.0',END)))==4:
                ec.insert('1.0','    ')
            elif int(len(ec.get('1.0',END)))==5:
                ec.insert('1.0','   ')
            elif int(len(ec.get('1.0',END)))==6:
                ec.insert('1.0','  ')
            con.close()
            
#TOTAL            
            cc=Text(HOME,font=('CAMBRIA bold',28),fg='white',bg='darkturquoise',width=6,height=1.3)
            cc.place(x=310,y=475)
            cc.config(insertbackground='darkturquoise')
            
            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
            cursor=con.cursor()
            cursor.execute("select count(patient_status) from c_h1")
            result=cursor.fetchone()
            cc.insert(END,result)
            if int(len(cc.get('1.0',END)))==2:
                cc.insert('1.0','      ')
            elif int(len(cc.get('1.0',END)))==3:
                cc.insert('1.0','     ')
            elif int(len(cc.get('1.0',END)))==4:
                cc.insert('1.0','    ')
            elif int(len(cc.get('1.0',END)))==5:
                cc.insert('1.0','   ')
            elif int(len(cc.get('1.0',END)))==5:
                cc.insert('1.0','  ')   
            con.close()

            

            def hide():
                ac.destroy()
                bc.destroy()
                cc.destroy()
                ec.destroy()
                widgethide.destroy()
               
                
            widgethide = Button(canvas_HOME, text='        HIDE        ',font=('Bahnschrift Light SemiCondensed',14), fg='white',bg='dodgerblue',command=hide)
            widgethide.pack()
            canvas_HOME.create_window(150, 560, window=widgethide)
            
            
        widget = Button(canvas_HOME, text='       ACTIVE      ',font=('Bahnschrift Light SemiCondensed',19),state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_HOME.create_window(150, 195, window=widget)

        widget = Button(canvas_HOME, text='   RECOVERED   ',font=('Bahnschrift Light SemiCondensed',19), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_HOME.create_window(150, 295, window=widget)        

        widget = Button(canvas_HOME, text='   REFERRED   ',font=('Bahnschrift Light SemiCondensed',19), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_HOME.create_window(150, 395, window=widget)
        
        widget = Button(canvas_HOME, text='       TOTAL       ',font=('Bahnschrift Light SemiCondensed BOLD',22), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_HOME.create_window(150, 495, window=widget)
        
        wisdget = Button(canvas_HOME, text='       SHOW       ',font=('Bahnschrift Light SemiCondensed BOLD',12), fg='white',bg='dodgerblue',command=showdashboard)
        wisdget.pack()
        canvas_HOME.create_window(150, 560, window=wisdget)


        ttk.Separator(canvas_HOME).place(x=490, y=5, relheight=10)
        ttk.Separator(canvas_HOME).place(x=1040, y=5, relheight=10)

#########
########################### DASHBOARD  PROGRESS
#########


       
                
        m4t=Text(HOME,font=('CAMBRIA',28),fg='white',bg='darkturquoise',width=7,height=1.3)
        m4t.place(x=1267,y=172)
        m4t.config(insertbackground='darkturquoise')

        
       
        def pro_6months():


            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
            cursor=con.cursor()
            cursor.execute("select count(*) from c_h1 where date>= now()-interval 7 month and date<= now()-interval 1 month")
            result1=cursor.fetchone()
            last6months=result1[0]
            con.close()
          
            avg6months=round(last6months/6,2)
           
            

            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
            cursor=con.cursor()
            cursor.execute("select count(*) from c_h1 where date>= now()-interval 1 month;")
            ky=cursor.fetchone()
            ongoingmonth=ky[0]
            
            con.close()

            m4t.delete('1.0',END)

            klo=str(round(((ongoingmonth-avg6months)/avg6months*100),2)) + ' %'

            m4t.insert(END,klo)

            

            if 25<=(ongoingmonth*100/avg6months)/2 < 50:
                
                s = ttk.Style()
                s.theme_use('clam')
                s.configure("salmon.Horizontal.TProgressbar",  foreground='salmon', background='salmon')
                style = ttk.Style(SEARCHPATIENT)
                style.configure("Treeview.Heading", font=('CAMBRRIA', 15),foreground='darkblue')
                style.configure("Treeview", font=('CAMBRRIA', 11))
                progressb = ttk.Progressbar(canvas_HOME,style="salmon.Horizontal.TProgressbar", orient = HORIZONTAL, length = 200, mode = 'determinate')
                progressb.pack()
                canvas_HOME.create_window(1250, 360, window=progressb)
                progressb['value'] = (ongoingmonth*100/avg6months)/2

                def lld():
                    progressb.destroy()
                    backAl.destroy()
                    m4t.delete('1.0',END)
    


                backAl=Button(canvas_HOME,text="  HIDE  ", font=('Bahnschrift Light SemiCondensed bold',12),bg='dodgerblue',fg='white',command=lld)
                canvas_HOME.create_window((1420, 180), anchor=NW, window=backAl)                         

            if (ongoingmonth*100/avg6months)/2 < 25:
                
                sm = ttk.Style()
                sm.theme_use('clam')
                sm.configure("red.Horizontal.TProgressbar",  foreground='red', background='red')
                style = ttk.Style(SEARCHPATIENT)
                style.configure("Treeview.Heading", font=('CAMBRRIA', 15),foreground='darkblue')
                style.configure("Treeview", font=('CAMBRRIA', 11))

                progressbz = ttk.Progressbar(canvas_HOME,style="red.Horizontal.TProgressbar", orient = HORIZONTAL, length = 200, mode = 'determinate')
                progressbz.pack()
                canvas_HOME.create_window(1250, 360, window=progressbz)
                progressbz['value'] = (ongoingmonth*100/avg6months)/2

                def ll():
                    progressbz.destroy()
                    backAlx.destroy()
                    m4t.delete('1.0',END)

    
                backAlx=Button(canvas_HOME,text="  HIDE  ", font=('Bahnschrift Light SemiCondensed bold',12),bg='dodgerblue',fg='white',command=ll)
                canvas_HOME.create_window((1420, 180), anchor=NW, window=backAlx)                         

        
            if (ongoingmonth*100/avg6months)/2 == 50:

                
                saa = ttk.Style()
                saa.theme_use('clam')
                saa.configure("yellow.Horizontal.TProgressbar", foreground='yellow', background='yellow')
                style = ttk.Style(SEARCHPATIENT)
                style.configure("Treeview.Heading", font=('CAMBRRIA', 15),foreground='darkblue')
                style.configure("Treeview", font=('CAMBRRIA', 11))
                                     

                progressv = ttk.Progressbar(canvas_HOME,style="yellow.Horizontal.TProgressbar", orient = HORIZONTAL, length = 200, mode = 'determinate')
                progressv.pack()
                canvas_HOME.create_window(1250, 360, window=progressv)
                progressv['value'] = (ongoingmonth*100/avg6months)/2

                def lllx():
                    progressv.destroy()
                    backAlSs.destroy()
                    m4t.delete('1.0',END)
                    
               

                backAlSs=Button(canvas_HOME,text="  HIDE  ", font=('Bahnschrift Light SemiCondensed bold',12),bg='dodgerblue',fg='white',command=lllx)
                canvas_HOME.create_window((1420, 180), anchor=NW, window=backAlSs)
                
            if 75<=(ongoingmonth*100/avg6months)/2 :
                
                sa = ttk.Style()
                sa.theme_use('clam')
                sa.configure("green.Horizontal.TProgressbar", foreground='green', background='green')
                style = ttk.Style(SEARCHPATIENT)
                style.configure("Treeview.Heading", font=('CAMBRRIA', 15),foreground='darkblue')
                style.configure("Treeview", font=('CAMBRRIA', 11))
                
                progress = ttk.Progressbar(canvas_HOME,style="green.Horizontal.TProgressbar", orient = HORIZONTAL, length = 200, mode = 'determinate')
                progress.pack()
                canvas_HOME.create_window(1250, 360, window=progress)  
                progress['value'] = (ongoingmonth*100/avg6months)/2

                def lll():
                    progress.destroy()
                    backAlS.destroy()
                    m4t.delete('1.0',END)

           
                backAlS=Button(canvas_HOME,text="  HIDE  ", font=('Bahnschrift Light SemiCondensed bold',12),bg='dodgerblue',fg='white',command=lll)
                canvas_HOME.create_window((1420, 180), anchor=NW, window=backAlS)
                    
            if 50<(ongoingmonth*100/avg6months)/2 < 75:
               
                
                sac = ttk.Style()
                sac.theme_use('clam')
                sac.configure("msgreen.Horizontal.TProgressbar", foreground='mediumspringgreen', background='mediumspringgreen')
                style = ttk.Style(SEARCHPATIENT)
                style.configure("Treeview.Heading", font=('CAMBRRIA', 15),foreground='darkblue')
                style.configure("Treeview", font=('CAMBRRIA', 11))
                
                progressm = ttk.Progressbar(canvas_HOME,style="msgreen.Horizontal.TProgressbar", orient = HORIZONTAL, length = 200, mode = 'determinate')
                progressm.pack()
                canvas_HOME.create_window(1250,360, window=progressm) 
                progressm['value'] = (ongoingmonth*100/avg6months)/2

                def llljg():
                    progressm.destroy()
                    backAlSc.destroy()
                    m4t.delete('1.0',END)
                                    

                backAlSc=Button(canvas_HOME,text="  HIDE  ", font=('Bahnschrift Light SemiCondensed bold',12),bg='dodgerblue',fg='white',command=llljg)
                canvas_HOME.create_window((1420, 180), anchor=NW, window=backAlSc)



        backAe=Button(canvas_HOME,text=" PROGRESS ", font=('Bahnschrift Light SemiCondensed bold',18),bg='DARKTURQUOISE',state='disabled', disabledforeground='white')
        canvas_HOME.create_window((1080, 170), anchor=NW, window=backAe)        

        backA=Button(canvas_HOME,text=" SHOW ", font=('Bahnschrift Light SemiCondensed bold',12),bg='dodgerblue',fg='white',command=pro_6months)
        canvas_HOME.create_window((1420, 180), anchor=NW, window=backA)




            
        potoiQmagzez = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/progressbar.png")
        #potoiQmagzez = potoiQmagzez.subsample(16, 16) 
        ajAa=Button(canvas_HOME, image = potoiQmagzez)
        ajAa.image=potoiQmagzez
        canvas_HOME.create_window((1147, 300), anchor=NW, window=ajAa)
        potoiQmagzez.image = potoiQmagzez                





        
        
########################################################################################################
########################################################################################################   

#####################################     ADD  PATIENT    TAB   ##################################### 


        
        nb1 = ttk.Notebook(ADDPATIENT)                         #######   NOTEBOOK   #######

        PERSONAL_DETAILS = ttk.Frame(nb1)        
        MEDICAL_DETAILS = ttk.Frame(nb1)
        OTHER_DETAILS   = ttk.Frame(nb1)
        IMAGES   = ttk.Frame(nb1)
  

        nb1.add(PERSONAL_DETAILS, text='                                          PERSONAL   DETAILS                                                                     ')
        nb1.add(MEDICAL_DETAILS, text='                                    MEDICAL    DETAILS                                                          ')
        nb1.add(OTHER_DETAILS , text='                          OTHER   DETAILS                                         ')               
        nb1.add(IMAGES , text='                                                    IMAGES                                                 ')
        nb1.pack(expand=1, fill="both")

        


 #####################        PERSONAL   DETAILS       ####################

        def phot1():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=fl)
            aA.image=p
            canvas.create_window((138,130),window=aA)
            p.image = p

        def phot2():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f2)
            aA.image=p
            canvas.create_window((438,130),window=aA)
            p.image = p

        def phot3():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f3)
            aA.image=p
            canvas.create_window((738,130),window=aA)
            p.image = p


        def phot4():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f4)
            aA.image=p
            canvas.create_window((1038,130),window=aA)
            p.image = p

        def phot5():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f5)
            aA.image=p
            canvas.create_window((1338,130),window=aA)
            p.image = p

        def phot6():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f6)
            aA.image=p
            canvas.create_window((138,460),window=aA)
            p.image = p


        def phot7():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f7)
            aA.image=p
            canvas.create_window((438,460),window=aA)
            p.image = p

            
        def phot8():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f8)
            aA.image=p
            canvas.create_window((738,460),window=aA)
            p.image = p

        def phot9():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f9)
            aA.image=p
            canvas.create_window((1038,460),window=aA)
            p.image = p

        def phot10():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f10)
            aA.image=p
            canvas.create_window((1338,460),window=aA)
            p.image = p                
            
        def phot11():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=fl1)
            aA.image=p
            canvas.create_window((138,790),window=aA)
            p.image = p

        def phot12():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f12)
            aA.image=p
            canvas.create_window((438,790),window=aA)
            p.image = p

        def phot13():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f13)
            aA.image=p
            canvas.create_window((738,790),window=aA)
            p.image = p


        def phot14():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f14)
            aA.image=p
            canvas.create_window((1038,790),window=aA)
            p.image = p

        def phot15():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f15)
            aA.image=p
            canvas.create_window((1338,790),window=aA)
            p.image = p

        def phot16():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f16)
            aA.image=p
            canvas.create_window((138,1120),window=aA)
            p.image = p


        def phot17():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f17)
            aA.image=p
            canvas.create_window((438,1120),window=aA)
            p.image = p

            
        def phot18():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f18)
            aA.image=p
            canvas.create_window((738,1120),window=aA)
            p.image = p

        def phot19():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f19)
            aA.image=p
            canvas.create_window((1038,1120),window=aA)
            p.image = p

        def phot20():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f20)
            aA.image=p
            canvas.create_window((1338,1120),window=aA)
            p.image = p

        def phot21():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f21)
            aA.image=p
            canvas.create_window((138,1450),window=aA)
            p.image = p

        def phot22():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f22)
            aA.image=p
            canvas.create_window((438,1450),window=aA)
            p.image = p

        def phot23():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f23)
            aA.image=p
            canvas.create_window((738,1450),window=aA)
            p.image = p


        def phot24():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f24)
            aA.image=p
            canvas.create_window((1038,1450),window=aA)
            p.image = p

        def phot25():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f25)
            aA.image=p
            canvas.create_window((1338,1450),window=aA)
            p.image = p

        def phot26():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f26)
            aA.image=p
            canvas.create_window((138,1780),window=aA)
            p.image = p


        def phot27():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f27)
            aA.image=p
            canvas.create_window((438,1780),window=aA)
            p.image = p

            
        def phot28():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f28)
            aA.image=p
            canvas.create_window((738,1780),window=aA)
            p.image = p

        def phot29():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f29)
            aA.image=p
            canvas.create_window((1038,1780),window=aA)
            p.image = p

        def phot30():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f30)
            aA.image=p
            canvas.create_window((1338,1780),window=aA)
            p.image = p                
            
        def phot31():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f31)
            aA.image=p
            canvas.create_window((138,2110),window=aA)
            p.image = p

        def phot32():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f32)
            aA.image=p
            canvas.create_window((438,2110),window=aA)
            p.image = p

        def phot33():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f33)
            aA.image=p
            canvas.create_window((738,2110),window=aA)
            p.image = p


        def phot34():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f34)
            aA.image=p
            canvas.create_window((1038,2110),window=aA)
            p.image = p

        def phot35():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f35)
            aA.image=p
            canvas.create_window((1338,2110),window=aA)
            p.image = p

        def phot36():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f36)
            aA.image=p
            canvas.create_window((138,2440),window=aA)
            p.image = p


        def phot37():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f37)
            aA.image=p
            canvas.create_window((438,2440),window=aA)
            p.image = p

            
        def phot38():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f38)
            aA.image=p
            canvas.create_window((738,2440),window=aA)
            p.image = p

        def phot39():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f39)
            aA.image=p
            canvas.create_window((1038,2440),window=aA)
            p.image = p

        def phot40():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f40)
            aA.image=p
            canvas.create_window((1338,2440),window=aA)
            p.image = p

        def phot41():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f41)
            aA.image=p
            canvas.create_window((138,2770),window=aA)
            p.image = p

        def phot42():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f42)
            aA.image=p
            canvas.create_window((438,2770),window=aA)
            p.image = p

        def phot43():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f43)
            aA.image=p
            canvas.create_window((738,2770),window=aA)
            p.image = p


        def phot44():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f44)
            aA.image=p
            canvas.create_window((1038,2770),window=aA)
            p.image = p

        def phot45():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f45)
            aA.image=p
            canvas.create_window((1338,2770),window=aA)
            p.image = p

        def phot46():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f46)
            aA.image=p
            canvas.create_window((138,3100),window=aA)
            p.image = p


        def phot47():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f47)
            aA.image=p
            canvas.create_window((438,3100),window=aA)
            p.image = p

            
        def phot48():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f48)
            aA.image=p
            canvas.create_window((738,3100),window=aA)
            p.image = p

        def phot49():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f49)
            aA.image=p
            canvas.create_window((1038,3100),window=aA)
            p.image = p

        def phot50():

            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
            p = p.subsample(1, 1) 
            aA=Button(canvas, text = '', image = p,bg='white',command=f50)
            aA.image=p
            canvas.create_window((1338,3100),window=aA)
            p.image = p


        def add1():
                    
            SER=e_SERIAL_NUMBER.get();   
            NAM=e_NAME.get();
            PAT=e_PATIENT_ID.get();
            DAT=e_DATE.get();
            AGE=e_AGE.get();
            SEX=e_SEX.get();
            OCC=e_OCCUPATION.get('1.0', END);
            ADD=e_ADDRESS.get('1.0', END);
            CONTACT=e_CONTACT.get('1.0', END);
            CAS=e_CASETYPE.get('1.0', END);
            TRE=e_TREATMENTTYPE.get('1.0', END);
            PT=e_PATIENTTYPE.get('1.0', END);
            CC=e_CHIEFCOMPLAINT.get('1.0', END);
            HOPI=e_HISTORYOFPRESENTILLNESS.get('1.0', END);
            MH=e_MEDICALHISTORY.get('1.0', END);
            ALL=e_ALLERGIES.get('1.0', END);   
            DH=e_DENTALHISTORY.get('1.0', END);
            GE=e_GENERALEXAMINATION.get('1.0', END);
            BP=e_BLOODPRESSURE.get('1.0', END);
            PR=e_PULSERATE.get('1.0', END);
            RR=e_RESPIRATORYRATE.get('1.0', END);
            HAB=e_HABITS.get('1.0', END);
            EOE=e_EXTRAORALEXAMINATION.get('1.0', END);
            TJ=e_TEMPOROMANDIBULARJOINT.get('1.0', END);
            EOL=e_EXTRAORALLYMPHNODES.get('1.0', END);
            SOF=e_SYMMETRYOFFACE.get('1.0', END);
            IOE=e_INTRAORALEXAMINATION.get('1.0', END);
            GIN=e_GINGIVA.get('1.0', END);
            COL=e_COLOR.get('1.0', END);
            CON=e_CONTOUR.get('1.0', END);
            CONSIS=e_CONSISTENCY.get('1.0', END);
            ST=e_SURFACETEXTURE.get('1.0', END);
            BOP=e_BLEEDINGONPROBING.get('1.0', END);
            PE=e_PERIODONTALEXAMINATION.get('1.0', END);
            PP=e_PERIODONTALPOCKETS.get('1.0', END);    
            STA=e_STAINS.get('1.0', END);
            CAL=e_CALCULUS.get('1.0', END);
            FI=e_FURCATIONINVOLVEMENT.get('1.0', END);
            REC=e_RECESSION.get('1.0', END);
            OF=e_OTHERFINDINGS.get('1.0', END);
            HP=e_HARDPALATE.get('1.0', END);
            SP=e_SOFTPALATE.get('1.0', END);       
            BM=e_BUCCALMUCOSA.get('1.0', END);
            TON=e_TONGUE.get('1.0', END);
            FOTM=e_FLOOROFTHEMOUTH.get('1.0', END);
            TE=e_TEETHEXAMINATION.get('1.0', END);
            NOT=e_NATUREOFTEETH.get('1.0', END);
            NOTP=e_NUMBEROFTEETHPRESENT.get('1.0', END);
            DT=e_DECAYEDTEETH.get('1.0', END);
            FT=e_FILLEDTEETH.get('1.0', END);
            TOP=e_TENDERNESSONPERCUSSION.get('1.0', END);
            RS=e_ROOTSTUMPS.get('1.0', END);
            MIS=e_MISSING.get('1.0', END);
            PD=e_PROVISIONALDIAGNOSIS.get('1.0', END);     
            DA=e_DIAGNOSTICAIDS.get('1.0', END);
            FD=e_FINALDIAGNOSIS.get('1.0', END);
            TP=e_TREATMENTPLAN.get('1.0', END);
            AI=e_ADDITIONALINFO.get('1.0', END);
            IMAGE=q1.get('1.0', '1.end');
            IMAGE1=q2.get('1.0',  '1.end');
            IMAGE2=q3.get('1.0',  '1.end');
            IMAGE3=q4.get('1.0',  '1.end');
            IMAGE4=q5.get('1.0',  '1.end');
            IMAGE5=q6.get('1.0',  '1.end');
            IMAGE6=q7.get('1.0',  '1.end');
            IMAGE7=q8.get('1.0',  '1.end');
            IMAGE8=q9.get('1.0',  '1.end');
            IMAGE9=q10.get('1.0',  '1.end');
            IMAGE10=q11.get('1.0',  '1.end');
            IMAGE11=q12.get('1.0',  '1.end');
            IMAGE12=q13.get('1.0',  '1.end');
            IMAGE13=q14.get('1.0',  '1.end');
            IMAGE14=q15.get('1.0',  '1.end');
            IMAGE15=q16.get('1.0',  '1.end');
            IMAGE16=q17.get('1.0',  '1.end');
            IMAGE17=q18.get('1.0',  '1.end');
            IMAGE18=q19.get('1.0',  '1.end');
            IMAGE19=q20.get('1.0',  '1.end');
            IMAGE20=q21.get('1.0',  '1.end');
            IMAGE21=q22.get('1.0',  '1.end');
            IMAGE22=q23.get('1.0',  '1.end');
            IMAGE23=q24.get('1.0',  '1.end');
            IMAGE24=q25.get('1.0',  '1.end');
            IMAGE25=q26.get('1.0',  '1.end');
            IMAGE26=q27.get('1.0',  '1.end');
            IMAGE27=q28.get('1.0',  '1.end');
            IMAGE28=q29.get('1.0',  '1.end');
            IMAGE29=q30.get('1.0',  '1.end');
            IMAGE30=q31.get('1.0',  '1.end');
            IMAGE31=q32.get('1.0',  '1.end');
            IMAGE32=q33.get('1.0',  '1.end');
            IMAGE33=q34.get('1.0',  '1.end');
            IMAGE34=q35.get('1.0',  '1.end');
            IMAGE35=q36.get('1.0',  '1.end');
            IMAGE36=q37.get('1.0',  '1.end');
            IMAGE37=q38.get('1.0',  '1.end');
            IMAGE38=q39.get('1.0',  '1.end');
            IMAGE39=q40.get('1.0',  '1.end');
            IMAGE40=q41.get('1.0',  '1.end');
            IMAGE41=q42.get('1.0',  '1.end');
            IMAGE42=q43.get('1.0',  '1.end');
            IMAGE43=q44.get('1.0',  '1.end');
            IMAGE44=q45.get('1.0',  '1.end');
            IMAGE45=q46.get('1.0',  '1.end');
            IMAGE46=q47.get('1.0',  '1.end');
            IMAGE47=q48.get('1.0',  '1.end');
            IMAGE48=q49.get('1.0',  '1.end');
            IMAGE49=q50.get('1.0',  '1.end');
            PS=variable.get();
            D1=img_desc1.get('1.0', END)
            D2=img_desc2.get('1.0', END)
            D3=img_desc3.get('1.0', END)
            D4=img_desc4.get('1.0', END)
            D5=img_desc5.get('1.0', END)
            D6=img_desc6.get('1.0', END)
            D7=img_desc7.get('1.0', END)
            D8=img_desc8.get('1.0', END)
            D9=img_desc9.get('1.0', END)
            D10=img_desc10.get('1.0', END)
            D11=img_desc11.get('1.0', END)
            D12=img_desc12.get('1.0', END)
            D13=img_desc13.get('1.0', END)
            D14=img_desc14.get('1.0', END)
            D15=img_desc15.get('1.0', END)
            D16=img_desc16.get('1.0', END)
            D17=img_desc17.get('1.0', END)
            D18=img_desc18.get('1.0', END)
            D19=img_desc19.get('1.0', END)
            D20=img_desc20.get('1.0', END)
            D21=img_desc21.get('1.0', END)            
            D22=img_desc22.get('1.0', END)
            D23=img_desc23.get('1.0', END)
            D24=img_desc24.get('1.0', END)
            D25=img_desc25.get('1.0', END)
            D26=img_desc26.get('1.0', END)
            D27=img_desc27.get('1.0', END)
            D28=img_desc28.get('1.0', END)
            D29=img_desc29.get('1.0', END)
            D30=img_desc30.get('1.0', END)
            D31=img_desc31.get('1.0', END)
            D32=img_desc32.get('1.0', END)
            D33=img_desc33.get('1.0', END)
            D34=img_desc34.get('1.0', END)
            D35=img_desc35.get('1.0', END)
            D36=img_desc36.get('1.0', END)
            D37=img_desc37.get('1.0', END)
            D38=img_desc38.get('1.0', END)
            D39=img_desc39.get('1.0', END)
            D40=img_desc40.get('1.0', END)
            D41=img_desc41.get('1.0', END)
            D42=img_desc42.get('1.0', END)
            D43=img_desc43.get('1.0', END)
            D44=img_desc44.get('1.0', END)
            D45=img_desc45.get('1.0', END)
            D46=img_desc46.get('1.0', END)
            D47=img_desc47.get('1.0', END)
            D48=img_desc48.get('1.0', END)
            D49=img_desc49.get('1.0', END)
            D50=img_desc50.get('1.0', END)

            
            if SER=='' :
                MessageBox.showerror("Insert Status","ENTER SERIAL NUMBER");
            elif AGE=='' :
                MessageBox.showerror("Insert Status","ENTER AGE");
            elif DAT=='':
                MessageBox.showerror("Insert Status","ENTER DATE");
            elif DAT[4]!="/":
                MessageBox.showerror("Insert Status","INVALID DATE FORMAT")
            elif DAT[7]!="/":
                MessageBox.showerror("Insert Status","INVALID DATE FORMAT")
            elif DAT[5]>"1":
                MessageBox.showerror("Insert Status","INVALID DATE FORMAT")
            elif DAT[8]>"3":
                MessageBox.showerror("Insert Status","INVALID DATE FORMAT")


            
                   
            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
            cursor=con.cursor()
            cursor.execute("select serial_number from c_h1 where serial_number='"+SER+"'");
            for row in cursor:
                aaa= str(row[0])
                if SER==aaa:
                    MessageBox.showerror("Insert Status","Serial Number already exists")
                    con.close()
               
                   
                      
            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
            cursor=con.cursor()
            cursor.execute("insert into c_h1 values('"+SER+"','"+NAM+"','"+PAT+"','"+DAT+"','"+AGE+"','"+SEX+"','"+OCC+"','"+ADD+"','"+CONTACT+"','"+CAS+"','"+TRE+"','"+PT+"','"+CC+"','"+HOPI+"','"+MH+"','"+ALL+"','"+DH+"','"+GE+"','"+BP+"','"+PR+"','"+RR+"','"+HAB+"','"+EOE+"','"+TJ+"','"+EOL+"','"+SOF+"','"+IOE+"','"+GIN+"','"+COL+"','"+CON+"','"+CONSIS+"','"+ST+"','"+BOP+"','"+PE+"','"+PP+"','"+STA+"','"+CAL+"','"+FI+"','"+REC+"','"+OF+"','"+HP+"','"+SP+"','"+BM+"','"+TON+"','"+FOTM+"','"+TE+"','"+NOT+"','"+NOTP+"','"+DT+"','"+FT+"','"+TOP+"','"+RS+"','"+MIS+"','"+PD+"','"+DA+"','"+FD+"','"+TP+"','"+AI+"','"+IMAGE+"','"+IMAGE1+"','"+IMAGE2+"','"+IMAGE3+"','"+IMAGE4+"','"+IMAGE5+"','"+IMAGE6+"','"+IMAGE7+"','"+IMAGE8+"','"+IMAGE9+"','"+IMAGE10+"','"+IMAGE11+"','"+IMAGE12+"','"+IMAGE13+"','"+IMAGE14+"','"+IMAGE15+"','"+IMAGE16+"','"+IMAGE17+"','"+IMAGE18+"','"+IMAGE19+"','"+IMAGE20+"','"+IMAGE21+"','"+IMAGE22+"','"+IMAGE23+"','"+IMAGE24+"','"+IMAGE25+"','"+IMAGE26+"','"+IMAGE27+"','"+IMAGE28+"','"+IMAGE29+"','"+IMAGE30+"','"+IMAGE31+"','"+IMAGE32+"','"+IMAGE33+"','"+IMAGE34+"','"+IMAGE35+"','"+IMAGE36+"','"+IMAGE37+"','"+IMAGE38+"','"+IMAGE39+"','"+IMAGE40+"','"+IMAGE41+"','"+IMAGE42+"','"+IMAGE43+"','"+IMAGE44+"','"+IMAGE45+"','"+IMAGE46+"','"+IMAGE47+"','"+IMAGE48+"','"+IMAGE49+"','"+PS+"','"+D1+"','"+D2+"','"+D3+"','"+D4+"','"+D5+"','"+D6+"','"+D7+"','"+D8+"','"+D9+"','"+D10+"','"+D11+"','"+D12+"','"+D13+"','"+D14+"','"+D15+"','"+D16+"','"+D17+"','"+D18+"','"+D19+"','"+D20+"','"+D21+"','"+D22+"','"+D23+"','"+D24+"','"+D25+"','"+D26+"','"+D27+"','"+D28+"','"+D29+"','"+D30+"','"+D31+"','"+D32+"','"+D33+"','"+D34+"','"+D35+"','"+D36+"','"+D37+"','"+D38+"','"+D39+"','"+D40+"','"+D41+"','"+D42+"','"+D43+"','"+D44+"','"+D45+"','"+D46+"','"+D47+"','"+D48+"','"+D49+"','"+D50+"') ");                                                                
            cursor.execute("commit")
            
                                                                          
            e_SERIAL_NUMBER.delete(0,'end')   
            e_NAME.delete(0,'end')  
            e_PATIENT_ID.delete(0,'end')  
            e_DATE.delete(0,'end')  
            e_AGE.delete(0,'end')  
               
            e_OCCUPATION.delete('1.0', END);
            e_ADDRESS.delete('1.0', END);
            e_CONTACT.delete('1.0', END);
            e_CASETYPE.delete('1.0', END);
            e_TREATMENTTYPE.delete('1.0', END);
            e_PATIENTTYPE.delete('1.0', END);
            e_CHIEFCOMPLAINT.delete('1.0', END);
            e_HISTORYOFPRESENTILLNESS.delete('1.0', END);
            e_MEDICALHISTORY.delete('1.0', END);
            e_ALLERGIES.delete('1.0', END);   
            e_DENTALHISTORY.delete('1.0', END);
            e_GENERALEXAMINATION.delete('1.0', END);
            e_BLOODPRESSURE.delete('1.0', END);
            e_PULSERATE.delete('1.0', END);
            e_RESPIRATORYRATE.delete('1.0', END);
            e_HABITS.delete('1.0', END);
            e_EXTRAORALEXAMINATION.delete('1.0', END);
            e_TEMPOROMANDIBULARJOINT.delete('1.0', END);
            e_EXTRAORALLYMPHNODES.delete('1.0', END);
            e_SYMMETRYOFFACE.delete('1.0', END);
            e_INTRAORALEXAMINATION.delete('1.0', END);
            e_GINGIVA.delete('1.0', END);
            e_COLOR.delete('1.0', END);
            e_CONTOUR.delete('1.0', END);
            e_CONSISTENCY.delete('1.0', END);
            e_SURFACETEXTURE.delete('1.0', END);
            e_BLEEDINGONPROBING.delete('1.0', END);
            e_PERIODONTALEXAMINATION.delete('1.0', END);
            e_PERIODONTALPOCKETS.delete('1.0', END);    
            e_STAINS.delete('1.0', END);
            e_CALCULUS.delete('1.0', END);
            e_FURCATIONINVOLVEMENT.delete('1.0', END);
            e_RECESSION.delete('1.0', END);
            e_OTHERFINDINGS.delete('1.0', END);
            e_HARDPALATE.delete('1.0', END);
            e_SOFTPALATE.delete('1.0', END);       
            e_BUCCALMUCOSA.delete('1.0', END);
            e_TONGUE.delete('1.0', END);
            e_FLOOROFTHEMOUTH.delete('1.0', END);
            e_TEETHEXAMINATION.delete('1.0', END);
            e_NATUREOFTEETH.delete('1.0', END);
            e_NUMBEROFTEETHPRESENT.delete('1.0', END);
            e_DECAYEDTEETH.delete('1.0', END);
            e_FILLEDTEETH.delete('1.0', END);
            e_TENDERNESSONPERCUSSION.delete('1.0', END);
            e_ROOTSTUMPS.delete('1.0', END);
            e_MISSING.delete('1.0', END);
            e_PROVISIONALDIAGNOSIS.delete('1.0', END);     
            e_DIAGNOSTICAIDS.delete('1.0', END);
            e_FINALDIAGNOSIS.delete('1.0', END);
            e_TREATMENTPLAN.delete('1.0', END);
            e_ADDITIONALINFO.delete('1.0', END);
            q1.delete('1.0', '1.end');
            q2.delete('1.0',  '1.end');
            q3.delete('1.0',  '1.end');
            q4.delete('1.0',  '1.end');
            q5.delete('1.0',  '1.end');
            q6.delete('1.0',  '1.end');
            q7.delete('1.0',  '1.end');
            q8.delete('1.0',  '1.end');
            q9.delete('1.0',  '1.end');
            q10.delete('1.0',  '1.end');
            q11.delete('1.0',  '1.end');
            q12.delete('1.0',  '1.end');
            q13.delete('1.0',  '1.end');
            q14.delete('1.0',  '1.end');
            q15.delete('1.0',  '1.end');
            q16.delete('1.0',  '1.end');
            q17.delete('1.0',  '1.end');
            q18.delete('1.0',  '1.end');
            q19.delete('1.0',  '1.end');
            q20.delete('1.0',  '1.end');
            q21.delete('1.0',  '1.end');
            q22.delete('1.0',  '1.end');
            q23.delete('1.0',  '1.end');
            q24.delete('1.0',  '1.end');
            q25.delete('1.0',  '1.end');
            q26.delete('1.0',  '1.end');
            q27.delete('1.0',  '1.end');
            q28.delete('1.0',  '1.end');
            q29.delete('1.0',  '1.end');
            q30.delete('1.0',  '1.end');
            q31.delete('1.0',  '1.end');
            q32.delete('1.0',  '1.end');
            q33.delete('1.0',  '1.end');
            q34.delete('1.0',  '1.end');
            q35.delete('1.0',  '1.end');
            q36.delete('1.0',  '1.end');
            q37.delete('1.0',  '1.end');
            q38.delete('1.0',  '1.end');
            q39.delete('1.0',  '1.end');
            q40.delete('1.0',  '1.end');
            q41.delete('1.0',  '1.end');
            q42.delete('1.0',  '1.end');
            q43.delete('1.0',  '1.end');
            q44.delete('1.0',  '1.end');
            q45.delete('1.0',  '1.end');
            q46.delete('1.0',  '1.end');
            q47.delete('1.0',  '1.end');
            q48.delete('1.0',  '1.end');
            q49.delete('1.0',  '1.end');
            q50.delete('1.0',  '1.end');
            img_desc1.delete('1.0', END)
            img_desc2.delete('1.0', END)
            img_desc3.delete('1.0', END)
            img_desc4.delete('1.0', END)
            img_desc5.delete('1.0', END)
            img_desc6.delete('1.0', END)
            img_desc7.delete('1.0', END)
            img_desc8.delete('1.0', END)
            img_desc9.delete('1.0', END)
            img_desc10.delete('1.0', END)
            img_desc11.delete('1.0', END)
            img_desc12.delete('1.0', END)
            img_desc13.delete('1.0', END)
            img_desc14.delete('1.0', END)
            img_desc15.delete('1.0', END)
            img_desc16.delete('1.0', END)
            img_desc17.delete('1.0', END)
            img_desc18.delete('1.0', END)
            img_desc19.delete('1.0', END)
            img_desc20.delete('1.0', END)
            img_desc21.delete('1.0', END)            
            img_desc22.delete('1.0', END)
            img_desc23.delete('1.0', END)
            img_desc24.delete('1.0', END)
            img_desc25.delete('1.0', END)
            img_desc26.delete('1.0', END)
            img_desc27.delete('1.0', END)
            img_desc28.delete('1.0', END)
            img_desc29.delete('1.0', END)
            img_desc30.delete('1.0', END)
            img_desc31.delete('1.0', END)
            img_desc32.delete('1.0', END)
            img_desc33.delete('1.0', END)
            img_desc34.delete('1.0', END)
            img_desc35.delete('1.0', END)
            img_desc36.delete('1.0', END)
            img_desc37.delete('1.0', END)
            img_desc38.delete('1.0', END)
            img_desc39.delete('1.0', END)
            img_desc40.delete('1.0', END)
            img_desc41.delete('1.0', END)
            img_desc42.delete('1.0', END)
            img_desc43.delete('1.0', END)
            img_desc44.delete('1.0', END)
            img_desc45.delete('1.0', END)
            img_desc46.delete('1.0', END)
            img_desc47.delete('1.0', END)
            img_desc48.delete('1.0', END)
            img_desc49.delete('1.0', END)
            img_desc50.delete('1.0', END)            


            phot1()
            phot2()
            phot3()
            phot4()
            phot5()
            phot6()
            phot7()
            phot8()
            phot9()
            phot10()
            phot11()
            phot12()
            phot13()
            phot14()
            phot15()
            phot16()
            phot17()
            phot18()
            phot19()
            phot20()            
            phot21()
            phot22()
            phot23()
            phot24()
            phot25()
            phot26()
            phot27()
            phot28()
            phot29()
            phot30()
            phot31()
            phot32()
            phot33()
            phot34()
            phot35()
            phot36()
            phot37()
            phot38()
            phot39()
            phot40()
            phot41()
            phot42()
            phot43()
            phot44()
            phot45()
            phot46()
            phot47()
            phot48()
            phot49()
            phot50()


               
            MessageBox.showinfo("Insert Status","ADDED SUCCESSFULLY");
            con.close();



    



        
        canvas_PERSONAL_DETAILS = Canvas(PERSONAL_DETAILS, width = 800, height = 400)
        #canvas_PERSONAL_DETAILS.configure(bg='lightsteelblue')
        canvas_PERSONAL_DETAILS.pack(side = LEFT, fill = BOTH, expand = True)



        
        entryContainerZr = Frame(canvas_PERSONAL_DETAILS)                # FRAME FOR CANVAS
        entryContainerZr.pack(fill = BOTH)

        gif1sX = Image.open("C:/Users/Acer/Desktop/pms files/ImageBlur (5).png")    # PERSONAL DETAILS BG IMAGE
        gif1sX = gif1sX.resize((1550, 765))
        imageD = ImageTk.PhotoImage(gif1sX)
        labels2 = tk.Label(canvas_PERSONAL_DETAILS,image=imageD)
        labels2.image=imageD
        canvas_PERSONAL_DETAILS.create_window((0, 0), anchor=NW, window=labels2)
        imageD.image = imageD        

                                                                    # TOP LABEL WITH ADDPATIENT TEXT       
        photoiQmage = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/addpatientlabel1.png") 
        aj=Button(canvas_PERSONAL_DETAILS, text = '', state='disabled', image = photoiQmage)
        aj.image=photoiQmage
        canvas_PERSONAL_DETAILS.create_window((8, 15), anchor=NW, window=aj)
        photoiQmage.image = photoiQmage

        





# DATA  INSERT  BUTTONS  AND  ENTRY  BOXES

        def focus_next_widget(event):
            event.widget.tk_focusNext().focus()
            return("break")
        
        widget = Button(canvas_PERSONAL_DETAILS, text='SERIAL NUMBER *',font=('arial bold',19), fg='white',bg='darkturquoise', state='disabled', disabledforeground='white')
        widget.pack()
        canvas_PERSONAL_DETAILS.create_window(270, 195, window=widget)
        e_SERIAL_NUMBER=Entry(canvas_PERSONAL_DETAILS,font=('verdana',11),bg='snow')
        canvas_PERSONAL_DETAILS.create_window((520,195),window=e_SERIAL_NUMBER)
        

        widget = Button(canvas_PERSONAL_DETAILS, text='      NAME       ',font=('arial ',16), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_PERSONAL_DETAILS.create_window(270, 285, window=widget)    
        e_NAME=Entry(canvas_PERSONAL_DETAILS,font=('verdana ',13),bg='snow')
        canvas_PERSONAL_DETAILS.create_window((520,285),window=e_NAME)
        

        
        widget = Button(canvas_PERSONAL_DETAILS, text='  PATIENT ID   ',font=('arial ',16), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_PERSONAL_DETAILS.create_window(270, 370, window=widget)
        e_PATIENT_ID=Entry(canvas_PERSONAL_DETAILS,font=('verdana',11),bg='snow')
        canvas_PERSONAL_DETAILS.create_window((520,370),window=e_PATIENT_ID)
            
        e_DATE= DateEntry(canvas_PERSONAL_DETAILS,date_pattern='y/mm/dd',selectmode='day',locale='en_INDIA')
        e_DATE.pack()                   
        canvas_PERSONAL_DETAILS.create_window(470, 455, window=e_DATE)
        
        widget = Button(canvas_PERSONAL_DETAILS, text='       DATE *       ',font=('arial',16),state='disabled',disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_PERSONAL_DETAILS.create_window(270, 455, window=widget)
        
        

        

        
        widget = Button(canvas_PERSONAL_DETAILS, text='        AGE *        ',font=('arial',16), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_PERSONAL_DETAILS.create_window(270, 535, window=widget)
        e_AGE=Entry(canvas_PERSONAL_DETAILS,font=('verdana',11),bg='snow')
        canvas_PERSONAL_DETAILS.create_window((520,535),window=e_AGE)

        


        widget = Button(canvas_PERSONAL_DETAILS, text='        SEX         ',font=('arial',16), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_PERSONAL_DETAILS.create_window(270, 615, window=widget)

        
        e_SEX = StringVar(canvas_PERSONAL_DETAILS)
        e_SEX.set("Male")
        waas = OptionMenu(canvas_PERSONAL_DETAILS, e_SEX, "Male", "Female")
        waas.config(font='arial')
        canvas_PERSONAL_DETAILS.create_window((460,615),window=waas)
        

       
        widget = Button(canvas_PERSONAL_DETAILS, text='  OCCUPATION  ',font=('arial',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_PERSONAL_DETAILS.create_window(960, 250, window=widget)
        e_OCCUPATION=Text(canvas_PERSONAL_DETAILS, height=3, width=40,font=('VERDANA',11),bg='snow')
        canvas_PERSONAL_DETAILS.create_window((1280,250),window=e_OCCUPATION)
        e_OCCUPATION.bind("<Tab>", focus_next_widget)
        
        
        widget = Button(canvas_PERSONAL_DETAILS, text='     ADDRESS     ',font=('arial',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_PERSONAL_DETAILS.create_window(960,350, window=widget)
        e_ADDRESS=Text(canvas_PERSONAL_DETAILS, height=3, width=40,font=('verdana',11),bg='snow')
        canvas_PERSONAL_DETAILS.create_window((1280,350),window=e_ADDRESS)
        e_ADDRESS.bind("<Tab>", focus_next_widget)
        
        

        
        widget = Button(canvas_PERSONAL_DETAILS, text='     CONTACT     ',font=('arial',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_PERSONAL_DETAILS.create_window(960, 450, window=widget)
        e_CONTACT=Text(canvas_PERSONAL_DETAILS, height=3, width=40,font=('VERDANA',11),bg='snow')
        canvas_PERSONAL_DETAILS.create_window((1280,450),window=e_CONTACT)
        e_CONTACT.bind("<Tab>", focus_next_widget)


        
 # ADD  RECORD  BUTTON         
        
        button = Button(canvas_PERSONAL_DETAILS, text=" ADD       ", font=("ROBOTO BOLD",18),bg='dodgerblue',fg='snow',command=add1)
        canvas_PERSONAL_DETAILS.create_window((1380, 130), anchor=NW, window=button)



 #####################        MEDICAL   DETAILS       ####################





        canvas_MEDICAL_DETAILS = Canvas(MEDICAL_DETAILS, width = 800, height = 400,scrollregion=(0,0,4650,6500))
        canvas_MEDICAL_DETAILS.pack(side = LEFT, fill = BOTH, expand = True)
        #canvas_MEDICAL_DETAILS.config(bg='gainsboro')

        canvas_MEDICAL_DETAILS.bind("<Down>",lambda event: canvas_MEDICAL_DETAILS.yview_scroll(2, "units"))
        canvas_MEDICAL_DETAILS.bind("<Up>",lambda event: canvas_MEDICAL_DETAILS.yview_scroll(-2, "units"))

        vsb = Scrollbar(canvas_MEDICAL_DETAILS, command=canvas_MEDICAL_DETAILS.yview)                        # SCROLLBAR FOR CANVAS
        canvas_MEDICAL_DETAILS.config(yscrollcommand = vsb.set)
        canvas_MEDICAL_DETAILS.pack(side = LEFT, fill = BOTH, expand = True)
        vsb.pack(side = RIGHT, fill = Y)


        
        entryContainerAZr = Frame(canvas_MEDICAL_DETAILS)                # FRAME FOR CANVAS
        entryContainerAZr.pack(fill = BOTH)

        def _on_mousewheel(event):
            canvas_MEDICAL_DETAILS.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas_MEDICAL_DETAILS.bind_all("<MouseWheel>",_on_mousewheel)




# DATA  INSERT  BUTTONS  AND  ENTRY  BOXES

        widget = Button(canvas_MEDICAL_DETAILS, text='CHIEF COMPLAINT',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(174, 100, window=widget)
        e_CHIEFCOMPLAINT=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,100),window=e_CHIEFCOMPLAINT)
        e_CHIEFCOMPLAINT.bind("<Tab>", focus_next_widget)
        
        

    

        
        widget = Button(canvas_MEDICAL_DETAILS, text='HISTORY OF PRESENT ILLNESS',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(174, 240, window=widget)
        e_HISTORYOFPRESENTILLNESS=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,240),window=e_HISTORYOFPRESENTILLNESS)
        e_HISTORYOFPRESENTILLNESS.bind("<Tab>", focus_next_widget)
        

       
        widget = Button(canvas_MEDICAL_DETAILS, text='MEDICAL HISTORY',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170,380, window=widget)
        e_MEDICALHISTORY=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,380),window=e_MEDICALHISTORY)
        e_MEDICALHISTORY.bind("<Tab>", focus_next_widget)
        
       
        widget = Button(canvas_MEDICAL_DETAILS, text='ALLERGIES',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 520, window=widget)
        e_ALLERGIES=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,520),window=e_ALLERGIES)
        e_ALLERGIES.bind("<Tab>", focus_next_widget)
        

        
        widget = Button(canvas_MEDICAL_DETAILS, text='DENTAL HISTORY',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 660, window=widget)
        e_DENTALHISTORY=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,660),window=e_DENTALHISTORY)
        e_DENTALHISTORY.bind("<Tab>", focus_next_widget)

        
        widget = Button(canvas_MEDICAL_DETAILS, text='GENERAL EXAMINATION',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 800, window=widget)
        e_GENERALEXAMINATION=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,800),window=e_GENERALEXAMINATION)
        e_GENERALEXAMINATION.bind("<Tab>", focus_next_widget)

        
        widget = Button(canvas_MEDICAL_DETAILS, text='BLOOD PRESSURE',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 940, window=widget)
        e_BLOODPRESSURE=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,940),window=e_BLOODPRESSURE)
        e_BLOODPRESSURE.bind("<Tab>", focus_next_widget)

       
        widget = Button(canvas_MEDICAL_DETAILS, text='PULSE RATE',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 1080, window=widget)
        e_PULSERATE=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,1080),window=e_PULSERATE)
        e_PULSERATE.bind("<Tab>", focus_next_widget)

       
        widget = Button(canvas_MEDICAL_DETAILS, text='RESPIRATORY RATE',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 1220, window=widget)
        e_RESPIRATORYRATE=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,1220),window=e_RESPIRATORYRATE)
        e_RESPIRATORYRATE.bind("<Tab>", focus_next_widget)

       
        widget = Button(canvas_MEDICAL_DETAILS, text='HABITS',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 1360, window=widget)
        e_HABITS=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,1360),window=e_HABITS)
        e_HABITS.bind("<Tab>", focus_next_widget)

        
        widget = Button(canvas_MEDICAL_DETAILS, text='EXTRA ORAL EXAMINATION',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170,1500, window=widget)
        e_EXTRAORALEXAMINATION=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,1500),window=e_EXTRAORALEXAMINATION)
        e_EXTRAORALEXAMINATION.bind("<Tab>", focus_next_widget)

       
        widget = Button(canvas_MEDICAL_DETAILS, text='TEMPOROMANDIBULAR JOINT',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 1640, window=widget)
        e_TEMPOROMANDIBULARJOINT=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,1640),window=e_TEMPOROMANDIBULARJOINT)
        e_TEMPOROMANDIBULARJOINT.bind("<Tab>", focus_next_widget)

        
        widget = Button(canvas_MEDICAL_DETAILS, text='EXTRA ORAL LYMPHNODES',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 1780, window=widget)
        e_EXTRAORALLYMPHNODES=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,1780),window=e_EXTRAORALLYMPHNODES)
        e_EXTRAORALLYMPHNODES.bind("<Tab>", focus_next_widget)

        
        widget = Button(canvas_MEDICAL_DETAILS, text='SYMMETRY OF FACE',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 1920, window=widget)
        e_SYMMETRYOFFACE=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,1920),window=e_SYMMETRYOFFACE)
        e_SYMMETRYOFFACE.bind("<Tab>", focus_next_widget)

        
        widget = Button(canvas_MEDICAL_DETAILS, text='INTRA ORAL EXAMINATION',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 2060, window=widget)
        e_INTRAORALEXAMINATION=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,2060),window=e_INTRAORALEXAMINATION)
        e_INTRAORALEXAMINATION.bind("<Tab>", focus_next_widget)
        

       
        widget = Button(canvas_MEDICAL_DETAILS, text='GINGIVA',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 2200, window=widget)
        e_GINGIVA=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,2200),window=e_GINGIVA)
        e_GINGIVA.bind("<Tab>", focus_next_widget)
        

        widget = Button(canvas_MEDICAL_DETAILS, text='COLOR',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 2340, window=widget)
        e_COLOR=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,2340),window=e_COLOR)
        e_COLOR.bind("<Tab>", focus_next_widget)
        
      

        widget = Button(canvas_MEDICAL_DETAILS, text='CONTOUR',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 2480, window=widget)
        e_CONTOUR=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,2480),window=e_CONTOUR)
        e_CONTOUR.bind("<Tab>", focus_next_widget)
        e_CONTOUR.bind("<Down>",lambda event: canvas_MEDICAL_DETAILS.yview_scroll(2, "units"))
        e_CONTOUR.bind("<Up>",lambda event: canvas_MEDICAL_DETAILS.yview_scroll(-2, "units"))
      


        widget = Button(canvas_MEDICAL_DETAILS, text='CONSISTENCY',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 2620, window=widget)
        e_CONSISTENCY=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,2620),window=e_CONSISTENCY)
        e_CONSISTENCY.bind("<Tab>", focus_next_widget)
        
        

        widget = Button(canvas_MEDICAL_DETAILS, text='SURFACE TEXTURE',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 2760, window=widget)
        e_SURFACETEXTURE=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,2760),window=e_SURFACETEXTURE)
        e_SURFACETEXTURE.bind("<Tab>", focus_next_widget)
       


        widget = Button(canvas_MEDICAL_DETAILS, text='BLEEDING ON PROBING',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 2900, window=widget)
        e_BLEEDINGONPROBING=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,2900),window=e_BLEEDINGONPROBING)
        e_BLEEDINGONPROBING.bind("<Tab>", focus_next_widget)
        


        widget = Button(canvas_MEDICAL_DETAILS, text='PERIODONTAL EXAMINATION',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 3040, window=widget)
        e_PERIODONTALEXAMINATION=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,3040),window=e_PERIODONTALEXAMINATION)
        e_PERIODONTALEXAMINATION.bind("<Tab>", focus_next_widget)
      


        widget = Button(canvas_MEDICAL_DETAILS, text='PERIODONTAL POCKETS',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 3180, window=widget)
        e_PERIODONTALPOCKETS=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,3180),window=e_PERIODONTALPOCKETS)
        e_PERIODONTALPOCKETS.bind("<Tab>", focus_next_widget)



        widget = Button(canvas_MEDICAL_DETAILS, text='STAINS',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 3320, window=widget)
        e_STAINS=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,3320),window=e_STAINS)
        e_STAINS.bind("<Tab>", focus_next_widget)



        widget = Button(canvas_MEDICAL_DETAILS, text='CALCULUS',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 3460, window=widget)
        e_CALCULUS=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,3460),window=e_CALCULUS)
        e_CALCULUS.bind("<Tab>", focus_next_widget)



        widget = Button(canvas_MEDICAL_DETAILS, text='FURCATION INVOLVEMENT',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 3600, window=widget)
        e_FURCATIONINVOLVEMENT=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,3600),window=e_FURCATIONINVOLVEMENT)
        e_FURCATIONINVOLVEMENT.bind("<Tab>", focus_next_widget)
      


        widget = Button(canvas_MEDICAL_DETAILS, text='RECESSION',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 3740, window=widget)
        e_RECESSION=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,3740),window=e_RECESSION)
        e_RECESSION.bind("<Tab>", focus_next_widget)
       

        widget = Button(canvas_MEDICAL_DETAILS, text='OTHER FINDINGS',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 3880, window=widget)
        e_OTHERFINDINGS=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,3880),window=e_OTHERFINDINGS)
        e_OTHERFINDINGS.bind("<Tab>", focus_next_widget)
       


        widget = Button(canvas_MEDICAL_DETAILS, text='HARD PALATE',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 4020, window=widget)
        e_HARDPALATE=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,4020),window=e_HARDPALATE)
        e_HARDPALATE.bind("<Tab>", focus_next_widget)
      


        widget = Button(canvas_MEDICAL_DETAILS, text='SOFT PALATE',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 4160, window=widget)
        e_SOFTPALATE=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,4160),window=e_SOFTPALATE)
        e_SOFTPALATE.bind("<Tab>", focus_next_widget)



        widget = Button(canvas_MEDICAL_DETAILS, text='BUCCAL MUCOSA',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 4300, window=widget)
        e_BUCCALMUCOSA=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,4300),window=e_BUCCALMUCOSA)
        e_BUCCALMUCOSA.bind("<Tab>", focus_next_widget)
  



        widget = Button(canvas_MEDICAL_DETAILS, text='TONGUE',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 4440, window=widget)
        e_TONGUE=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,4440),window=e_TONGUE)
        e_TONGUE.bind("<Tab>", focus_next_widget)


        widget = Button(canvas_MEDICAL_DETAILS, text='FLOOR OF THE MOUTH',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 4580, window=widget)
        e_FLOOROFTHEMOUTH=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,4580),window=e_FLOOROFTHEMOUTH)
        e_FLOOROFTHEMOUTH.bind("<Tab>", focus_next_widget)



        widget = Button(canvas_MEDICAL_DETAILS, text='TEETH EXAMINATION',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 4720, window=widget)
        e_TEETHEXAMINATION=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,4720),window=e_TEETHEXAMINATION)
        e_TEETHEXAMINATION.bind("<Tab>", focus_next_widget)



        widget = Button(canvas_MEDICAL_DETAILS, text='NATURE OF TEETH',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 4860, window=widget)
        e_NATUREOFTEETH=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,4860),window=e_NATUREOFTEETH)
        e_NATUREOFTEETH.bind("<Tab>", focus_next_widget)
       


        widget = Button(canvas_MEDICAL_DETAILS, text='NUMBER OF TEETH PRESENT',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 5000, window=widget)
        e_NUMBEROFTEETHPRESENT=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,5000),window=e_NUMBEROFTEETHPRESENT)
        e_NUMBEROFTEETHPRESENT.bind("<Tab>", focus_next_widget)
        


        widget = Button(canvas_MEDICAL_DETAILS, text='DECAYED TEETH',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 5140, window=widget)
        e_DECAYEDTEETH=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,5140),window=e_DECAYEDTEETH)
        e_DECAYEDTEETH.bind("<Tab>", focus_next_widget)
      

        widget = Button(canvas_MEDICAL_DETAILS, text='FILLED TEETH',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 5280, window=widget)
        e_FILLEDTEETH=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,5280),window=e_FILLEDTEETH)
        e_FILLEDTEETH.bind("<Tab>", focus_next_widget)
       


        widget = Button(canvas_MEDICAL_DETAILS, text='TENDERNESS ON PERCUSSION',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 5420, window=widget)
        e_TENDERNESSONPERCUSSION=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,5420),window=e_TENDERNESSONPERCUSSION)
        e_TENDERNESSONPERCUSSION.bind("<Tab>", focus_next_widget)



        widget = Button(canvas_MEDICAL_DETAILS, text='ROOT STUMPS',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 5560, window=widget)
        e_ROOTSTUMPS=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,5560),window=e_ROOTSTUMPS)
        e_ROOTSTUMPS.bind("<Tab>", focus_next_widget)



        widget = Button(canvas_MEDICAL_DETAILS, text='MISSING',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 5700, window=widget)
        e_MISSING=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,5700),window=e_MISSING)
        e_MISSING.bind("<Tab>", focus_next_widget)



        widget = Button(canvas_MEDICAL_DETAILS, text='PROVISIONAL DIAGNOSIS',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 5840, window=widget)
        e_PROVISIONALDIAGNOSIS=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,5840),window=e_PROVISIONALDIAGNOSIS)
        e_PROVISIONALDIAGNOSIS.bind("<Tab>", focus_next_widget)



        widget = Button(canvas_MEDICAL_DETAILS, text='DIAGNOSTIC AIDS',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 5980, window=widget)
        e_DIAGNOSTICAIDS=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,5980),window=e_DIAGNOSTICAIDS)
        e_DIAGNOSTICAIDS.bind("<Tab>", focus_next_widget)



        widget = Button(canvas_MEDICAL_DETAILS, text='FINAL DIAGNOSIS',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 6120, window=widget)
        e_FINALDIAGNOSIS=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,6120),window=e_FINALDIAGNOSIS)
        e_FINALDIAGNOSIS.bind("<Tab>", focus_next_widget)



        widget = Button(canvas_MEDICAL_DETAILS, text='TREATMENT PLAN',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 6260, window=widget)
        e_TREATMENTPLAN=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,6260),window=e_TREATMENTPLAN)
        e_TREATMENTPLAN.bind("<Tab>", focus_next_widget)



        widget = Button(canvas_MEDICAL_DETAILS, text='ADDITIONAL INFO',font=('arial ',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_MEDICAL_DETAILS.create_window(170, 6400, window=widget)
        e_ADDITIONALINFO=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
        canvas_MEDICAL_DETAILS.create_window((920,6400),window=e_ADDITIONALINFO)
        e_ADDITIONALINFO.bind("<Tab>", focus_next_widget)




 #####################        OTHER   DETAILS       ####################




        canvas_OTHER_DETAILS = Canvas(OTHER_DETAILS, width = 800, height = 400)
        canvas_OTHER_DETAILS.pack(side = LEFT, fill = BOTH, expand = True)



        
        AentryContainerAZr = Frame(canvas_OTHER_DETAILS)                # FRAME FOR CANVAS
        AentryContainerAZr.pack(fill = BOTH)

                                       


# DATA  INSERT  BUTTONS  AND  ENTRY  BOXES



        widget = Button(canvas_OTHER_DETAILS, text='    CASE TYPE    ',font=('VERDANA',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_OTHER_DETAILS.create_window(320, 80, window=widget)
        e_CASETYPE=Text(canvas_OTHER_DETAILS, height=3, width=45,font=('VERDANA',11),bg='snow')
        canvas_OTHER_DETAILS.create_window((690,80),window=e_CASETYPE)
        e_CASETYPE.bind("<Tab>", focus_next_widget)
    
      



        widget = Button(canvas_OTHER_DETAILS, text='TREATMENT TYPE',font=('VERDANA',14),state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_OTHER_DETAILS.create_window(320, 180, window=widget)
        e_TREATMENTTYPE=Text(canvas_OTHER_DETAILS, height=3, width=45,font=('VERDANA',11),bg='snow')
        canvas_OTHER_DETAILS.create_window((690,180),window=e_TREATMENTTYPE)
        e_TREATMENTTYPE.bind("<Tab>", focus_next_widget)



        widget = Button(canvas_OTHER_DETAILS, text='  PATIENT TYPE  ',font=('VERDANA',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widget.pack()
        canvas_OTHER_DETAILS.create_window(320, 280, window=widget)
        e_PATIENTTYPE=Text(canvas_OTHER_DETAILS, height=3, width=45,font=('VERDANA',11),bg='snow')
        canvas_OTHER_DETAILS.create_window((690,280),window=e_PATIENTTYPE)
        e_PATIENTTYPE.bind("<Tab>", focus_next_widget)

        

        widgetb = Button(canvas_OTHER_DETAILS, text='PATIENT STATUS',font=('VERDANA',14), state='disabled', disabledforeground='white',bg='darkturquoise')
        widgetb.pack()
        canvas_OTHER_DETAILS.create_window(320, 380, window=widgetb)
        variable = StringVar(canvas_OTHER_DETAILS)
        variable.set("Active")
        was = OptionMenu(canvas_OTHER_DETAILS, variable, "Active", "Recovered",'Referred')
        was.config(font='arial')
        canvas_OTHER_DETAILS.create_window((550,380),window=was)

#############################  IMAGE  ADD   ############################
        q1=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q1)

        q2=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q2)

        q3=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q3)

        q4=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q4)

        q5=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q5)

        q6=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q6)

        q7=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q7)

        q8=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q8)

        q9=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q9)

        q10=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q10)

        q11=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q11)

        q12=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q12)

        q13=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q13)

        q14=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q14)

        q15=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q15)

        q16=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q16)

        q17=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q17)

        q18=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q18)

        q19=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q19)

        q20=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q20)

        q21=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q21)

        q22=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q22)

        q23=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q23)

        q24=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q24)

        q25=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q25)

        q26=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q26)

        q27=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q27)

        q28=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q28)

        q29=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q29)

        q30=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q30)

        q31=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q31)

        q32=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q32)

        q33=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q33)

        q34=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q34)

        q35=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q35)

        q36=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q36)

        q37=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q37)

        q38=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q38)

        q39=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q39)

        q40=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q40)

        q41=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q41)

        q42=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q42)

        q43=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q43)

        q44=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q44)

        q45=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q45)

        q46=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q46)

        q47=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q47)

        q48=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q48)

        q49=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q49)

        q50=Text(canvas_OTHER_DETAILS,height=3,width=100)
        canvas_OTHER_DETAILS.create_window((10000,20000),window=q50)

        
      

        def fl():
            filename = fd.askopenfilename(parent=canvas,title='Choose a file')
            q1.insert(END,filename)
            
            imagxxadfeae = Image.open(filename)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t1=Text(canvas,width=24,height=11)
            canvas.create_window((138,130),window=t1)
            t1.image=imagxxadfeae
            
            t1.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae




        def f2():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q2.insert(END,filenamee)
            
            imagxxadfeae1 = Image.open(filenamee)
            imagxxadfeae1 = imagxxadfeae1.resize((196,176))
            imagxxadfeae1 = ImageTk.PhotoImage(imagxxadfeae1)
            
            t2=Text(canvas,width=24,height=11)
            canvas.create_window((438,130),window=t2)                
            t2.image=imagxxadfeae1
            
            t2.image_create(END, image =imagxxadfeae1)
            imagxxadfeae1.image=imagxxadfeae1


        def f3():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q3.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t3=Text(canvas,width=24,height=11)
            canvas.create_window((738,130),window=t3)                
            t3.image=imagxxadfeae
            
            t3.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae


        def f4():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q4.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t4=Text(canvas,width=24,height=11)
            canvas.create_window((1038,130),window=t4)                
            t4.image=imagxxadfeae
            
            t4.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae


        def f5():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q5.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t5=Text(canvas,width=24,height=11)
            canvas.create_window((1338,130),window=t5)                
            t5.image=imagxxadfeae
            
            t5.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae

        def f6():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q6.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t6=Text(canvas,width=24,height=11)
            canvas.create_window((138,460),window=t6)                
            t6.image=imagxxadfeae
            
            t6.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae                

        def f7():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q7.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t7=Text(canvas,width=24,height=11)
            canvas.create_window((438,460),window=t7)                
            t7.image=imagxxadfeae
            
            t7.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae


        def f8():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q8.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t8=Text(canvas,width=24,height=11)
            canvas.create_window((738,460),window=t8)                
            t8.image=imagxxadfeae
            
            t8.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae


        def f9():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q9.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t9=Text(canvas,width=24,height=11)
            canvas.create_window((1038,460),window=t9)                
            t9.image=imagxxadfeae
            
            t9.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae
            
        def f10():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q10.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t10=Text(canvas,width=24,height=11)
            canvas.create_window((1338,460),window=t10)                
            t10.image=imagxxadfeae
            
            t10.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae

            
        def fl1():
            filename = fd.askopenfilename(parent=canvas,title='Choose a file')
            q11.insert(END,filename)
            
            imagxxadfeae = Image.open(filename)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t11=Text(canvas,width=24,height=11)
            canvas.create_window((138,790),window=t11)
            t11.image=imagxxadfeae
            
            t11.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae




        def f12():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q12.insert(END,filenamee)
            
            imagxxadfeae1 = Image.open(filenamee)
            imagxxadfeae1 = imagxxadfeae1.resize((196,176))
            imagxxadfeae1 = ImageTk.PhotoImage(imagxxadfeae1)
            
            t12=Text(canvas,width=24,height=11)
            canvas.create_window((438,790),window=t12)                
            t12.image=imagxxadfeae1
            
            t12.image_create(END, image =imagxxadfeae1)
            imagxxadfeae1.image=imagxxadfeae1


        def f13():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q13.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t13=Text(canvas,width=24,height=11)
            canvas.create_window((738,790),window=t13)                
            t13.image=imagxxadfeae
            
            t13.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae


        def f14():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q14.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t14=Text(canvas,width=24,height=11)
            canvas.create_window((1038,790),window=t14)                
            t14.image=imagxxadfeae
            
            t14.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae


        def f15():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q15.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t15=Text(canvas,width=24,height=11)
            canvas.create_window((1338,790),window=t15)                
            t15.image=imagxxadfeae
            
            t15.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae

        def f16():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q16.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t16=Text(canvas,width=24,height=11)
            canvas.create_window((138,1120),window=t16)                
            t16.image=imagxxadfeae
            
            t16.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae                

        def f17():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q17.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t17=Text(canvas,width=24,height=11)
            canvas.create_window((438,1120),window=t17)                
            t17.image=imagxxadfeae
            
            t17.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae


        def f18():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q18.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t18=Text(canvas,width=24,height=11)
            canvas.create_window((738,1120),window=t18)                
            t18.image=imagxxadfeae
            
            t18.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae


        def f19():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q19.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t19=Text(canvas,width=24,height=11)
            canvas.create_window((1038,1120),window=t19)                
            t19.image=imagxxadfeae
            
            t19.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae
            
        def f20():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q20.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t20=Text(canvas,width=24,height=11)
            canvas.create_window((1338,1120),window=t20)                
            t20.image=imagxxadfeae
            
            t20.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae

            
        def f21():
            filename = fd.askopenfilename(parent=canvas,title='Choose a file')
            q21.insert(END,filename)
            
            imagxxadfeae = Image.open(filename)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t21=Text(canvas,width=24,height=11)
            canvas.create_window((138,1450),window=t21)
            t21.image=imagxxadfeae
            
            t21.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae




        def f22():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q22.insert(END,filenamee)
            
            imagxxadfeae1 = Image.open(filenamee)
            imagxxadfeae1 = imagxxadfeae1.resize((196,176))
            imagxxadfeae1 = ImageTk.PhotoImage(imagxxadfeae1)
            
            t12=Text(canvas,width=24,height=11)
            canvas.create_window((438,1450),window=t12)                
            t12.image=imagxxadfeae1
            
            t12.image_create(END, image =imagxxadfeae1)
            imagxxadfeae1.image=imagxxadfeae1


        def f23():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q23.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t13=Text(canvas,width=24,height=11)
            canvas.create_window((738,1450),window=t13)                
            t13.image=imagxxadfeae
            
            t13.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae


        def f24():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q24.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t14=Text(canvas,width=24,height=11)
            canvas.create_window((1038,1450),window=t14)                
            t14.image=imagxxadfeae
            
            t14.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae


        def f25():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q25.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t15=Text(canvas,width=24,height=11)
            canvas.create_window((1338,1450),window=t15)                
            t15.image=imagxxadfeae
            
            t15.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae

        def f26():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q26.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t16=Text(canvas,width=24,height=11)
            canvas.create_window((138,1780),window=t16)                
            t16.image=imagxxadfeae
            
            t16.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae                

        def f27():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q27.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t17=Text(canvas,width=24,height=11)
            canvas.create_window((438,1780),window=t17)                
            t17.image=imagxxadfeae
            
            t17.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae


        def f28():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q28.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t18=Text(canvas,width=24,height=11)
            canvas.create_window((738,1780),window=t18)                
            t18.image=imagxxadfeae
            
            t18.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae


        def f29():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q29.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t19=Text(canvas,width=24,height=11)
            canvas.create_window((1038,1780),window=t19)                
            t19.image=imagxxadfeae
            
            t19.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae
            
        def f30():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q30.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t20=Text(canvas,width=24,height=11)
            canvas.create_window((1338,1780),window=t20)                
            t20.image=imagxxadfeae
            
            t20.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae

            
        def f31():
            filename = fd.askopenfilename(parent=canvas,title='Choose a file')
            q31.insert(END,filename)
            
            imagxxadfeae = Image.open(filename)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t1=Text(canvas,width=24,height=11)
            canvas.create_window((138,2110),window=t1)
            t1.image=imagxxadfeae
            
            t1.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae




        def f32():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q32.insert(END,filenamee)
            
            imagxxadfeae1 = Image.open(filenamee)
            imagxxadfeae1 = imagxxadfeae1.resize((196,176))
            imagxxadfeae1 = ImageTk.PhotoImage(imagxxadfeae1)
            
            t2=Text(canvas,width=24,height=11)
            canvas.create_window((438,2110),window=t2)                
            t2.image=imagxxadfeae1
            
            t2.image_create(END, image =imagxxadfeae1)
            imagxxadfeae1.image=imagxxadfeae1


        def f33():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q33.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t3=Text(canvas,width=24,height=11)
            canvas.create_window((738,2110),window=t3)                
            t3.image=imagxxadfeae
            
            t3.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae


        def f34():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q34.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t4=Text(canvas,width=24,height=11)
            canvas.create_window((1038,2110),window=t4)                
            t4.image=imagxxadfeae
            
            t4.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae


        def f35():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q35.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t5=Text(canvas,width=24,height=11)
            canvas.create_window((1338,2110),window=t5)                
            t5.image=imagxxadfeae
            
            t5.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae

        def f36():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q36.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t6=Text(canvas,width=24,height=11)
            canvas.create_window((138,2440),window=t6)                
            t6.image=imagxxadfeae
            
            t6.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae                

        def f37():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q37.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t7=Text(canvas,width=24,height=11)
            canvas.create_window((438,2440),window=t7)                
            t7.image=imagxxadfeae
            
            t7.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae


        def f38():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q38.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t8=Text(canvas,width=24,height=11)
            canvas.create_window((738,2440),window=t8)                
            t8.image=imagxxadfeae
            
            t8.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae


        def f39():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q39.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t9=Text(canvas,width=24,height=11)
            canvas.create_window((1038,2440),window=t9)                
            t9.image=imagxxadfeae
            
            t9.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae
            
        def f40():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q40.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t10=Text(canvas,width=24,height=11)
            canvas.create_window((1338,2440),window=t10)                
            t10.image=imagxxadfeae
            
            t10.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae

            
        def f41():
            filename = fd.askopenfilename(parent=canvas,title='Choose a file')
            q41.insert(END,filename)
            
            imagxxadfeae = Image.open(filename)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t11=Text(canvas,width=24,height=11)
            canvas.create_window((138,2770),window=t11)
            t11.image=imagxxadfeae
            
            t11.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae




        def f42():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q42.insert(END,filenamee)
            
            imagxxadfeae1 = Image.open(filenamee)
            imagxxadfeae1 = imagxxadfeae1.resize((196,176))
            imagxxadfeae1 = ImageTk.PhotoImage(imagxxadfeae1)
            
            t12=Text(canvas,width=24,height=11)
            canvas.create_window((438,2770),window=t12)                
            t12.image=imagxxadfeae1
            
            t12.image_create(END, image =imagxxadfeae1)
            imagxxadfeae1.image=imagxxadfeae1


        def f43():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q43.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t13=Text(canvas,width=24,height=11)
            canvas.create_window((738,2770),window=t13)                
            t13.image=imagxxadfeae
            
            t13.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae


        def f44():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q44.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t14=Text(canvas,width=24,height=11)
            canvas.create_window((1038,2770),window=t14)                
            t14.image=imagxxadfeae
            
            t14.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae


        def f45():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q45.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t15=Text(canvas,width=24,height=11)
            canvas.create_window((1338,2770),window=t15)                
            t15.image=imagxxadfeae
            
            t15.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae

        def f46():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q46.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t16=Text(canvas,width=24,height=11)
            canvas.create_window((138,3100),window=t16)                
            t16.image=imagxxadfeae
            
            t16.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae                

        def f47():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q47.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t17=Text(canvas,width=24,height=11)
            canvas.create_window((438,3100),window=t17)                
            t17.image=imagxxadfeae
            
            t17.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae


        def f48():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q48.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t18=Text(canvas,width=24,height=11)
            canvas.create_window((738,3100),window=t18)                
            t18.image=imagxxadfeae
            
            t18.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae


        def f49():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q49.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t19=Text(canvas,width=24,height=11)
            canvas.create_window((1038,3100),window=t19)                
            t19.image=imagxxadfeae
            
            t19.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae
            
        def f50():
            filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
            q50.insert(END,filenamee)
            
            imagxxadfeae = Image.open(filenamee)
            imagxxadfeae = imagxxadfeae.resize((196,176))
            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
            
            t40=Text(canvas,width=24,height=11)
            canvas.create_window((1338,3100),window=t40)                
            t40.image=imagxxadfeae
            
            t40.image_create(END, image =imagxxadfeae)
            imagxxadfeae.image=imagxxadfeae

             
        def d1():
            phot1()
            q1.delete(1.0,END)


        def d2():
            phot2()
            q2.delete(1.0,END)


        def d3():
            phot3()
            q3.delete(1.0,END)

        def d4():
            phot4()
            q4.delete(1.0,END)

        def d5():
            phot5()
            q5.delete(1.0,END)

        def d6():
            phot6()
            q6.delete(1.0,END)
                
        def d7():
            phot7()
            q7.delete(1.0,END)

        def d8():
            phot8()
            q8.delete(1.0,END)

        def d9():
            phot9()
            q9.delete(1.0,END)

        def d10():
            phot10()
            q10.delete(1.0,END)

        def d11():
            phot11()
            q11.delete(1.0,END)


        def d12():
            phot12()
            q12.delete(1.0,END)


        def d13():
            phot13()
            q13.delete(1.0,END)

        def d14():
            phot14()
            q14.delete(1.0,END)

        def d15():
            phot15()
            q15.delete(1.0,END)

        def d16():
            phot16()
            q16.delete(1.0,END)
                
        def d17():
            phot17()
            q17.delete(1.0,END)

        def d18():
            phot18()
            q18.delete(1.0,END)

        def d19():
            phot19()
            q19.delete(1.0,END)

        def d20():
            phot20()
            q20.delete(1.0,END)

             
        def d21():
            phot21()
            q21.delete(1.0,END)


        def d22():
            phot22()
            q22.delete(1.0,END)


        def d23():
            phot23()
            q23.delete(1.0,END)

        def d24():
            phot24()
            q24.delete(1.0,END)

        def d25():
            phot25()
            q25.delete(1.0,END)

        def d26():
            phot26()
            q26.delete(1.0,END)
                
        def d27():
            phot27()
            q27.delete(1.0,END)

        def d28():
            phot28()
            q28.delete(1.0,END)

        def d29():
            phot29()
            q29.delete(1.0,END)

        def d30():
            phot30()
            q30.delete(1.0,END)

        def d31():
            phot31()
            q31.delete(1.0,END)


        def d32():
            phot32()
            q32.delete(1.0,END)


        def d33():
            phot33()
            q33.delete(1.0,END)

        def d34():
            phot34()
            q34.delete(1.0,END)

        def d35():
            phot35()
            q35.delete(1.0,END)

        def d36():
            phot36()
            q36.delete(1.0,END)
                
        def d37():
            phot37()
            q37.delete(1.0,END)

        def d38():
            phot38()
            q38.delete(1.0,END)

        def d39():
            phot39()
            q39.delete(1.0,END)

        def d40():
            phot40()
            q40.delete(1.0,END)

        def d41():
            phot41()
            q41.delete(1.0,END)


        def d42():
            phot42()
            q42.delete(1.0,END)


        def d43():
            phot43()
            q43.delete(1.0,END)

        def d44():
            phot44()
            q44.delete(1.0,END)

        def d45():
            phot45()
            q45.delete(1.0,END)

        def d46():
            phot46()
            q46.delete(1.0,END)
                
        def d47():
            phot47()
            q47.delete(1.0,END)

        def d48():
            phot48()
            q48.delete(1.0,END)

        def d49():
            phot49()
            q49.delete(1.0,END)

        def d50():
            phot50()
            q50.delete(1.0,END)



            

        frame=Frame(IMAGES,width=500,height=800)
        frame.pack(expand=True, fill=BOTH) 

        canvas=tk.Canvas(frame,width=300,height=500,scrollregion=(0,0,1500,3400))
        canvas.configure(bg='whitesmoke')



    

# add image icon
            
        phot1()
        phot2()
        phot3()
        phot4()
        phot5()
        phot6()
        phot7()
        phot8()
        phot9()
        phot10()
        phot11()
        phot12()
        phot13()
        phot14()
        phot15()
        phot16()
        phot17()
        phot18()
        phot19()
        phot20()            
        phot21()
        phot22()
        phot23()
        phot24()
        phot25()
        phot26()
        phot27()
        phot28()
        phot29()
        phot30()
        phot31()
        phot32()
        phot33()
        phot34()
        phot35()
        phot36()
        phot37()
        phot38()
        phot39()
        phot40()
        phot41()
        phot42()
        phot43()
        phot44()
        phot45()
        phot46()
        phot47()
        phot48()
        phot49()
        phot50()

        
# DELETE BUTTONS

        de=Button(canvas,text='Delete',command=d1)
        canvas.create_window((138,21),window=de)

        de=Button(canvas,text='Delete',command=d2)
        canvas.create_window((438,21),window=de)

        de=Button(canvas,text='Delete',command=d3)
        canvas.create_window((738,21),window=de)

        de=Button(canvas,text='Delete',command=d4)
        canvas.create_window((1038,21),window=de)

        de=Button(canvas,text='Delete',command=d5)
        canvas.create_window((1338,21),window=de)
        
        de=Button(canvas,text='Delete',command=d6)
        canvas.create_window((138,350),window=de)

        de=Button(canvas,text='Delete',command=d7)
        canvas.create_window((438,350),window=de)

        de=Button(canvas,text='Delete',command=d8)
        canvas.create_window((738,350),window=de)
        
        de=Button(canvas,text='Delete',command=d9)
        canvas.create_window((1038,350),window=de)

        de=Button(canvas,text='Delete',command=d10)
        canvas.create_window((1338,350),window=de)
        
        de=Button(canvas,text='Delete',command=d11)
        canvas.create_window((138,679),window=de)

        de=Button(canvas,text='Delete',command=d12)
        canvas.create_window((438,679),window=de)

        de=Button(canvas,text='Delete',command=d13)
        canvas.create_window((738,679),window=de)

        de=Button(canvas,text='Delete',command=d14)
        canvas.create_window((1038,679),window=de)

        de=Button(canvas,text='Delete',command=d15)
        canvas.create_window((1338,679),window=de)
        
        de=Button(canvas,text='Delete',command=d16)
        canvas.create_window((138,1008),window=de)

        de=Button(canvas,text='Delete',command=d17)
        canvas.create_window((438,1008),window=de)

        de=Button(canvas,text='Delete',command=d18)
        canvas.create_window((738,1008),window=de)
        
        de=Button(canvas,text='Delete',command=d19)
        canvas.create_window((1038,1008),window=de)

        de=Button(canvas,text='Delete',command=d20)
        canvas.create_window((1338,1008),window=de)

        de=Button(canvas,text='Delete',command=d21)
        canvas.create_window((138,1337),window=de)

        de=Button(canvas,text='Delete',command=d22)
        canvas.create_window((438,1337),window=de)

        de=Button(canvas,text='Delete',command=d23)
        canvas.create_window((738,1337),window=de)

        de=Button(canvas,text='Delete',command=d24)
        canvas.create_window((1038,1337),window=de)

        de=Button(canvas,text='Delete',command=d25)
        canvas.create_window((1338,1337),window=de)
        
        de=Button(canvas,text='Delete',command=d26)
        canvas.create_window((138,1666),window=de)

        de=Button(canvas,text='Delete',command=d27)
        canvas.create_window((438,1666),window=de)

        de=Button(canvas,text='Delete',command=d28)
        canvas.create_window((738,1666),window=de)
        
        de=Button(canvas,text='Delete',command=d29)
        canvas.create_window((1038,1666),window=de)

        de=Button(canvas,text='Delete',command=d30)
        canvas.create_window((1338,1666),window=de)
        
        de=Button(canvas,text='Delete',command=d31)
        canvas.create_window((138,1995),window=de)

        de=Button(canvas,text='Delete',command=d32)
        canvas.create_window((438,1995),window=de)

        de=Button(canvas,text='Delete',command=d33)
        canvas.create_window((738,1995),window=de)

        de=Button(canvas,text='Delete',command=d34)
        canvas.create_window((1038,1995),window=de)

        de=Button(canvas,text='Delete',command=d35)
        canvas.create_window((1338,1995),window=de)
        
        de=Button(canvas,text='Delete',command=d36)
        canvas.create_window((138,2324),window=de)

        de=Button(canvas,text='Delete',command=d37)
        canvas.create_window((438,2324),window=de)

        de=Button(canvas,text='Delete',command=d38)
        canvas.create_window((738,2324),window=de)
        
        de=Button(canvas,text='Delete',command=d39)
        canvas.create_window((1038,2324),window=de)

        de=Button(canvas,text='Delete',command=d40)
        canvas.create_window((1338,2324),window=de)

        de=Button(canvas,text='Delete',command=d41)
        canvas.create_window((138,2653),window=de)

        de=Button(canvas,text='Delete',command=d42)
        canvas.create_window((438,2653),window=de)

        de=Button(canvas,text='Delete',command=d43)
        canvas.create_window((738,2653),window=de)

        de=Button(canvas,text='Delete',command=d44)
        canvas.create_window((1038,2653),window=de)

        de=Button(canvas,text='Delete',command=d45)
        canvas.create_window((1338,2653),window=de)
        
        de=Button(canvas,text='Delete',command=d46)
        canvas.create_window((138,2982),window=de)

        de=Button(canvas,text='Delete',command=d47)
        canvas.create_window((438,2982),window=de)

        de=Button(canvas,text='Delete',command=d48)
        canvas.create_window((738,2982),window=de)
        
        de=Button(canvas,text='Delete',command=d49)
        canvas.create_window((1038,2982),window=de)

        de=Button(canvas,text='Delete',command=d50)
        canvas.create_window((1338,2982),window=de)
        
# image path textbox            
                
# descripption boxes



        img_desc1=Text(canvas,width=26,height=3)
        canvas.create_window((138,273),window=img_desc1)

        img_desc2=Text(canvas,width=26,height=3)
        canvas.create_window((438,273),window=img_desc2)


        img_desc3=Text(canvas,width=26,height=3)
        canvas.create_window((738,273),window=img_desc3)            


        img_desc4=Text(canvas,width=26,height=3)
        canvas.create_window((1038,273),window=img_desc4)


        img_desc5=Text(canvas,width=26,height=3)
        canvas.create_window((1338,273),window=img_desc5)


        img_desc6=Text(canvas,width=26,height=3)
        canvas.create_window((138,603),window=img_desc6)

        img_desc7=Text(canvas,width=26,height=3)
        canvas.create_window((438,603),window=img_desc7)

        img_desc8=Text(canvas,width=26,height=3)
        canvas.create_window((738,603),window=img_desc8)

        img_desc9=Text(canvas,width=26,height=3)
        canvas.create_window((1038,603),window=img_desc9)            

        img_desc10=Text(canvas,width=26,height=3)
        canvas.create_window((1338,603),window=img_desc10)

        img_desc11=Text(canvas,width=26,height=3)
        canvas.create_window((138,933),window=img_desc11)

        img_desc12=Text(canvas,width=26,height=3)
        canvas.create_window((438,933),window=img_desc12)


        img_desc13=Text(canvas,width=26,height=3)
        canvas.create_window((738,933),window=img_desc13)            


        img_desc14=Text(canvas,width=26,height=3)
        canvas.create_window((1038,933),window=img_desc14)


        img_desc15=Text(canvas,width=26,height=3)
        canvas.create_window((1338,933),window=img_desc15)


        img_desc16=Text(canvas,width=26,height=3)
        canvas.create_window((138,1263),window=img_desc16)

        img_desc17=Text(canvas,width=26,height=3)
        canvas.create_window((438,1263),window=img_desc17)

        img_desc18=Text(canvas,width=26,height=3)
        canvas.create_window((738,1263),window=img_desc18)

        img_desc19=Text(canvas,width=26,height=3)
        canvas.create_window((1038,1263),window=img_desc19)            

        img_desc20=Text(canvas,width=26,height=3)
        canvas.create_window((1338,1263),window=img_desc20)

        img_desc21=Text(canvas,width=26,height=3)
        canvas.create_window((138,1593),window=img_desc21)

        img_desc22=Text(canvas,width=26,height=3)
        canvas.create_window((438,1593),window=img_desc22)


        img_desc23=Text(canvas,width=26,height=3)
        canvas.create_window((738,1593),window=img_desc23)            


        img_desc24=Text(canvas,width=26,height=3)
        canvas.create_window((1038,1593),window=img_desc24)


        img_desc25=Text(canvas,width=26,height=3)
        canvas.create_window((1338,1593),window=img_desc25)


        img_desc26=Text(canvas,width=26,height=3)
        canvas.create_window((138,1923),window=img_desc26)

        img_desc27=Text(canvas,width=26,height=3)
        canvas.create_window((438,1923),window=img_desc27)

        img_desc28=Text(canvas,width=26,height=3)
        canvas.create_window((738,1923),window=img_desc28)

        img_desc29=Text(canvas,width=26,height=3)
        canvas.create_window((1038,1923),window=img_desc29)            

        img_desc30=Text(canvas,width=26,height=3)
        canvas.create_window((1338,1923),window=img_desc30)

        img_desc31=Text(canvas,width=26,height=3)
        canvas.create_window((138,2253),window=img_desc31)

        img_desc32=Text(canvas,width=26,height=3)
        canvas.create_window((438,2253),window=img_desc32)


        img_desc33=Text(canvas,width=26,height=3)
        canvas.create_window((738,2253),window=img_desc33)            


        img_desc34=Text(canvas,width=26,height=3)
        canvas.create_window((1038,2253),window=img_desc34)


        img_desc35=Text(canvas,width=26,height=3)
        canvas.create_window((1338,2253),window=img_desc35)


        img_desc36=Text(canvas,width=26,height=3)
        canvas.create_window((138,2583),window=img_desc36)

        img_desc37=Text(canvas,width=26,height=3)
        canvas.create_window((438,2583),window=img_desc37)

        img_desc38=Text(canvas,width=26,height=3)
        canvas.create_window((738,2583),window=img_desc38)

        img_desc39=Text(canvas,width=26,height=3)
        canvas.create_window((1038,2583),window=img_desc39)            

        img_desc40=Text(canvas,width=26,height=3)
        canvas.create_window((1338,2583),window=img_desc40)

        img_desc41=Text(canvas,width=26,height=3)
        canvas.create_window((138,2913),window=img_desc41)

        img_desc42=Text(canvas,width=26,height=3)
        canvas.create_window((438,2913),window=img_desc42)


        img_desc43=Text(canvas,width=26,height=3)
        canvas.create_window((738,2913),window=img_desc43)            


        img_desc44=Text(canvas,width=26,height=3)
        canvas.create_window((1038,2913),window=img_desc44)


        img_desc45=Text(canvas,width=26,height=3)
        canvas.create_window((1338,2913),window=img_desc45)


        img_desc46=Text(canvas,width=26,height=3)
        canvas.create_window((138,3243),window=img_desc46)

        img_desc47=Text(canvas,width=26,height=3)
        canvas.create_window((438,3243),window=img_desc47)

        img_desc48=Text(canvas,width=26,height=3)
        canvas.create_window((738,3243),window=img_desc48)

        img_desc49=Text(canvas,width=26,height=3)
        canvas.create_window((1038,3243),window=img_desc49)            

        img_desc50=Text(canvas,width=26,height=3)
        canvas.create_window((1338,3243),window=img_desc50)

        v1=Button(top,text=' 1 ',bg='dodgerblue',fg='white')
        canvas.create_window((15,120),window=v1)

        v1=Button(top,text=' 2 ',bg='dodgerblue',fg='white')
        canvas.create_window((310,120),window=v1)

        v1=Button(top,text=' 3 ',bg='dodgerblue',fg='white')
        canvas.create_window((613,120),window=v1)

        v1=Button(top,text=' 4 ',bg='dodgerblue',fg='white')
        canvas.create_window((915,120),window=v1)

        v1=Button(top,text=' 5 ',bg='dodgerblue',fg='white')
        canvas.create_window((1215,120),window=v1)
        
        v1=Button(top,text=' 6 ',bg='dodgerblue',fg='white')
        canvas.create_window((15,460),window=v1)

        v1=Button(top,text=' 7 ',bg='dodgerblue',fg='white')
        canvas.create_window((310,460),window=v1)

        v1=Button(top,text=' 8 ',bg='dodgerblue',fg='white')
        canvas.create_window((613,460),window=v1)

        v1=Button(top,text=' 9 ',bg='dodgerblue',fg='white')
        canvas.create_window((915,460),window=v1)

        v1=Button(top,text=' 10 ',bg='dodgerblue',fg='white')
        canvas.create_window((1215,460),window=v1)
        
        v1=Button(top,text=' 11 ',bg='dodgerblue',fg='white')
        canvas.create_window((15,790),window=v1)

        v1=Button(top,text=' 12 ',bg='dodgerblue',fg='white')
        canvas.create_window((310,790),window=v1)

        v1=Button(top,text=' 13 ',bg='dodgerblue',fg='white')
        canvas.create_window((613,790),window=v1)

        v1=Button(top,text=' 14 ',bg='dodgerblue',fg='white')
        canvas.create_window((915,790),window=v1)

        v1=Button(top,text=' 15 ',bg='dodgerblue',fg='white')
        canvas.create_window((1215,790),window=v1)
        
        v1=Button(top,text=' 16 ',bg='dodgerblue',fg='white')
        canvas.create_window((15,1110),window=v1)

        v1=Button(top,text=' 17 ',bg='dodgerblue',fg='white')
        canvas.create_window((310,1110),window=v1)

        v1=Button(top,text=' 18 ',bg='dodgerblue',fg='white')
        canvas.create_window((613,1110),window=v1)

        v1=Button(top,text=' 19 ',bg='dodgerblue',fg='white')
        canvas.create_window((915,1110),window=v1)

        v1=Button(top,text=' 20 ',bg='dodgerblue',fg='white')
        canvas.create_window((1215,1110),window=v1)
        
        v1=Button(top,text=' 21 ',bg='dodgerblue',fg='white')
        canvas.create_window((15,1440),window=v1)

        v1=Button(top,text=' 22 ',bg='dodgerblue',fg='white')
        canvas.create_window((310,1440),window=v1)

        v1=Button(top,text=' 23 ',bg='dodgerblue',fg='white')
        canvas.create_window((613,1440),window=v1)

        v1=Button(top,text=' 24 ',bg='dodgerblue',fg='white')
        canvas.create_window((915,1440),window=v1)

        v1=Button(top,text=' 25 ',bg='dodgerblue',fg='white')
        canvas.create_window((1215,1440),window=v1)
        
        v1=Button(top,text=' 26 ',bg='dodgerblue',fg='white')
        canvas.create_window((15,1770),window=v1)

        v1=Button(top,text=' 27 ',bg='dodgerblue',fg='white')
        canvas.create_window((310,1770),window=v1)

        v1=Button(top,text=' 28 ',bg='dodgerblue',fg='white')
        canvas.create_window((613,1770),window=v1)

        v1=Button(top,text=' 29 ',bg='dodgerblue',fg='white')
        canvas.create_window((915,1770),window=v1)

        v1=Button(top,text=' 30 ',bg='dodgerblue',fg='white')
        canvas.create_window((1215,1770),window=v1)
        
        v1=Button(top,text=' 31 ',bg='dodgerblue',fg='white')
        canvas.create_window((15,2100),window=v1)

        v1=Button(top,text=' 32 ',bg='dodgerblue',fg='white')
        canvas.create_window((310,2100),window=v1)

        v1=Button(top,text=' 33 ',bg='dodgerblue',fg='white')
        canvas.create_window((613,2100),window=v1)

        v1=Button(top,text=' 34 ',bg='dodgerblue',fg='white')
        canvas.create_window((915,2100),window=v1)

        v1=Button(top,text=' 35 ',bg='dodgerblue',fg='white')
        canvas.create_window((1215,2100),window=v1)
        
        v1=Button(top,text=' 36 ',bg='dodgerblue',fg='white')
        canvas.create_window((15,2430),window=v1)

        v1=Button(top,text=' 37 ',bg='dodgerblue',fg='white')
        canvas.create_window((310,2430),window=v1)

        v1=Button(top,text=' 38 ',bg='dodgerblue',fg='white')
        canvas.create_window((613,2430),window=v1)

        v1=Button(top,text=' 39 ',bg='dodgerblue',fg='white')
        canvas.create_window((915,2430),window=v1)

        v1=Button(top,text=' 40 ',bg='dodgerblue',fg='white')
        canvas.create_window((1215,2430),window=v1)
        
        v1=Button(top,text=' 41 ',bg='dodgerblue',fg='white')
        canvas.create_window((15,2760),window=v1)

        v1=Button(top,text=' 42 ',bg='dodgerblue',fg='white')
        canvas.create_window((310,2760),window=v1)

        v1=Button(top,text=' 43 ',bg='dodgerblue',fg='white')
        canvas.create_window((613,2760),window=v1)

        v1=Button(top,text=' 44 ',bg='dodgerblue',fg='white')
        canvas.create_window((915,2760),window=v1)

        v1=Button(top,text=' 45 ',bg='dodgerblue',fg='white')
        canvas.create_window((1215,2760),window=v1)

        v1=Button(top,text=' 46 ',bg='dodgerblue',fg='white')
        canvas.create_window((15,3080),window=v1)

        v1=Button(top,text=' 47 ',bg='dodgerblue',fg='white')
        canvas.create_window((310,3080),window=v1)

        v1=Button(top,text=' 48 ',bg='dodgerblue',fg='white')
        canvas.create_window((613,3080),window=v1)

        v1=Button(top,text=' 49 ',bg='dodgerblue',fg='white')
        canvas.create_window((915,3080),window=v1)

        v1=Button(top,text=' 50 ',bg='dodgerblue',fg='white')
        canvas.create_window((1215,3080),window=v1)


            
        vbar=Scrollbar(frame,orient=VERTICAL,width=25)
        vbar.pack(side=RIGHT,fill=Y)
        vbar.config(command=canvas.yview)
        canvas.bind('<Up>', lambda event : canvas.yview_scroll( -1, 'units'))
        canvas.bind('<Down>', lambda event : canvas.yview_scroll(1, 'units'))
 
  


        canvas.config(width=1550,height=1250)
        canvas.config( yscrollcommand=vbar.set)
        canvas.pack(side=LEFT,expand=True,fill=BOTH)
            




        

#########################################################################################################
################################   CANVAS   FOR  SEARCH  PATIENT    TAB   #############################
#########################################################################################################

    
        

        def showrecords_serialsno():

            sss=Entry(SEARCHPATIENT)
            sss.place(x=10,y=10)


            agsif1dsaad = Image.open("C:/Users/Acer/Desktop/pms files/Untitledf.png")    #  BG IMAGE
            diqmageeda = ImageTk.PhotoImage(agsif1dsaad)
            labelz2d = tk.Label(SEARCHPATIENT,image=diqmageeda)
            labelz2d.image=diqmageeda
            labelz2d.place(x=200,y=550)
            diqmageeda.image = diqmageeda                
         



            
            treeview=ttk.Treeview(SEARCHPATIENT,columns=('Name','', 'Serial Number','Patient ID','Date','Age','Sex','Status'),height=32)

            style = ttk.Style(SEARCHPATIENT)
            style.configure("Treeview.Heading", font=('CAMBRRIA', 15),foreground='darkblue')
            style.configure("Treeview", font=('CAMBRRIA', 11))
            

            treeview.pack()
            treeview.heading('#2', text='Serial No.')
            treeview.heading('#1', text='')
            treeview.heading('#3', text='Patient ID')
            treeview.heading('#0', text='                                                                                           Name')
            treeview.heading('#4', text='Date')
            treeview.heading('#5', text='Age')
            treeview.heading('#6', text='Sex')
            treeview.heading('#7', text='Contact ')
            treeview.heading('#8', text='Status')
            treeview.column('#1', width=1,stretch=tk.YES)
            treeview.column('#2',width=120, stretch=tk.YES)
            treeview.column('#3', width=120,stretch=tk.YES)
            treeview.column('#0',width=740, stretch=tk.YES)
            treeview.column('#4',width=110, stretch=tk.YES)
            treeview.column('#5',width=90, stretch=tk.YES)
            treeview.column('#6',width=105, stretch=tk.YES)
            treeview.column('#7',width=120, stretch=tk.YES)
            treeview.column('#8',width=125, stretch=tk.YES)
            
            treeview.grid(row=5, columnspan=5, sticky='nsew')

            gsif1dsaad = Image.open("C:/Users/Acer/Desktop/pms files/ImageBlur (5) - Copy.png")    #  BG IMAGE
            gsif1dsaad = gsif1dsaad.resize((548, 762))
            diqmageed = ImageTk.PhotoImage(gsif1dsaad)
            labelz2d = tk.Label(SEARCHPATIENT,image=diqmageed)
            labelz2d.image=diqmageed 
            labelz2d.place(x=0,y=0)
            diqmageed.image = diqmageed                
         
 
            ttk.Separator(SEARCHPATIENT).place(x=0, y=320,relwidth=0.36)



            
            def my():
                treeview.insert('', 'end', text='',values=('','','','','','','',''))

                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select* from c_h1")
                rows=cursor.fetchall()
                
             #very long..
                for row in rows:
                    treeview.insert('', 'end', text='                                                                                                                                       '+str(row[1]),values=('|',str(row[0]),'|   '+str(row[2]),'|    '+str(row[3]), '|     '+str(row[4]),'|      '+str(row[5]),'|     '+str(row[8]),'|     '+str(row[108])))
            my()

            def refresh():
                cleartreeview()
                my()

            def myy():
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select* from c_h1 order by "+sorta.get()+" "+sortaa.get()+"")
                rows=cursor.fetchall()
                cleartreeview()
                treeview.insert('', 'end', text='',values=('','','','','','','',''))   
                
    #very long..
                for row in rows:
                    treeview.insert('', 'end', text='                                                                                                                                       '+str(row[1]),values=('|',str(row[0]),'|   '+str(row[2]),'|    '+str(row[3]), '|     '+str(row[4]),'|      '+str(row[5]),'|     '+str(row[8]),'|     '+str(row[108])))
                
                       
            def cleartreeview():
                x = treeview.get_children()

                if x != '()': # checks if there is something in the first row
                    for child in x:
                        treeview.delete(child)

                        
            def selectItem():

                def installed_printer():
                    printers = win32print.EnumPrinters(2)
                    for p in printers:
                        return(p)

                printerdef = ''

                def locprinter():
                    
                    var1 = StringVar()
                    LABEL = Label(topa, text="select Printer").pack()
                    PRCOMBO = ttk.Combobox(pt, width=35,textvariable=var1)
                    print_list = []
                    printers = list(win32print.EnumPrinters(2))
                    for i in printers:
                        print_list.append(i[2])
                    print(print_list)
                    # Put printers in combobox
                    PRCOMBO['values'] = print_list
                    PRCOMBO.pack()
                    def select():
                        global printerdef
                        printerdef = PRCOMBO.get()
                        pt.destroy()
                    BUTTON = ttk.Button(topa, text="Done",command=select).pack()


               



                def INFO():
                    printText = text_serialno.get("1.0", END)
                    
                    print(printerdef)
                    filename = tempfile.mktemp(".txt")
                    open(filename, "w").write(printText)
                    # Bellow is call to print text from T2 textbox
                    win32api.ShellExecute(
                        0,
                        "printto",
                        filename,
                        '"%s"' % win32print.GetDefaultPrinter(),
                        ".",
                        0
                    )


                curItem = treeview.focus()
                values = treeview.item(curItem)['values']
                cc=(values[1])
                sss.delete(0,END)
                sss.insert(END,cc)

                top=Toplevel()
                scrollbar = Scrollbar(top,width=30)
                scrollbar.pack(side=RIGHT, fill=Y)     
             
                text_serialno=Text(top, height=80, width=100,font=("calibri",17),bg='white',fg='black',wrap=WORD,spacing1=5,spacing2=5)
                text_serialno.pack()
                
                text_serialno.config(yscrollcommand=scrollbar.set)
                scrollbar.config(command=text_serialno.yview)

                    
                printt=Button(top,text=" PRINT ", font=('Bahnschrift Light SemiCondensed',15),bg='dodgerblue',fg='white',command=INFO)
                printt.place(x=1410,y=30)
                
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor(named_tuple=True)
                cursor.execute("select* from c_h1 where serial_number ='"+ sss.get()+"' ")
                text_serialno.delete('1.0', END)

                for row in cursor:
                    
                    
                    top='     '
                    line='                        SK  DENTAL  CLINIC '
                    line1=''
                    pe='                                                                                        PERSONAL  DETAILS'
                    pe1=''
                    a= 'SERIAL NUMBER    -       '+ str(row[0])
                    a1=''
                    b= 'NAME                         -       '+str(row[1])
                    b1=''
                    c='PATIENT ID             -        '+str(row[2])
                    c1=''
                    d='DATE                         -        '+str(row[3])
                    d1=''
                    e='AGE                            -        '+str(row[4])
                    e1=''
                    f='SEX                             -        '+str(row[5])
                    f1=''
                    g='OCCUPATION         -        '+str(row[6])
                    g1=''
                    h='ADDRESS                 -        '+str(row[7])
                    h1=''
                    i='CONTACT                -        '+str(row[8])
                    i1=''
                    md=''
                    md1='                                                                                 DENTAL  DETAILS'
                    md2=''
                    j='CHIEF COMPLAINT   -   '+str(row[12])
                    j1=''
                    k='HISTORY OF PRESENT ILLNESS  -  '+str(row[13])
                    k1=''
                    l='MEDICAL HISTORY  -  '+str(row[14])
                    l1=''
                    m='ALLERGIES  -  '+str(row[15])
                    m1=''
                    n='DENTAL HISTORY  -  '+str(row[16])
                    n1=''
                    o='GENERAL EXAMINATION  -  '+str(row[17])
                    o1=''
                    p='BLOOD PRESSURE  -  '+str(row[18])
                    p1=''
                    q='PULSE RATE  -  '+str(row[19])
                    q1=''
                    r='RESPIRATORY RATE  -  '+str(row[20])
                    r1=''
                    s='HABITS  -  '+str(row[21])
                    s1=''
                    t='EXTRA ORAL EXAMINATION  -  '+str(row[22])
                    t1=''
                    u='TEMPOROMANDIBULAR JOINT  -  '+str(row[23])
                    u1=''
                    v='EXTRA ORAL LYMPHNODES  -  '+str(row[24])
                    v1=''
                    w='SYMMETRY OF FACE  -  '+str(row[25])
                    w1=''
                    x='INTRA ORAL EXAMINATION  -  '+str(row[26])
                    x1=''
                    y='GINGIVA  -  '+str(row[27])
                    y1=''
                    z='COLOR  -  '+str(row[28])
                    z1=''
                    aa='CONTOUR  -  '+str(row[29])
                    aa1=''
                    bb='CONSISTENCY  -  '+str(row[30])
                    bb1=''
                    cc='SURFACE TEXTURE  -  '+str(row[31])
                    cc1=''
                    dd='BLEEDING ON PROBING  -  '+str(row[32])
                    dd1=''
                    ee='PERIODONTAL EXAMINATION  -  '+str(row[33])
                    ee1=''
                    ff='PERIODONTAL POCKETS  -  '+str(row[34])
                    ff1=''
                    gg='STAINS  -  '+str(row[35])
                    gg1=''
                    hh='CALCULUS  -  '+str(row[36])
                    hh1=''
                    ii='FURCATION INVOLVEMENT  -  '+str(row[37])
                    ii1=''
                    jj='RECESSION  -  '+str(row[38])
                    jj1=''
                    kk='OTHER FINDINGS  -  '+str(row[39])
                    kk1=''
                    ll='HARD PALATE  -  '+str(row[40])
                    ll1=''
                    mm='SOFT PALATE  -  '+str(row[41])
                    mm1=''
                    nn='BUCCAL MUCOSA  -  '+str(row[42])
                    nn1=''
                    oo='TONGUE  -  '+str(row[43])
                    oo1=''
                    pp='FLOOR OF THE MOUTH  -  '+str(row[44])
                    pp1=''
                    qq='TEETH EXAMINATION  -  '+str(row[45])
                    qq1=''
                    rr='NATURE OF TEETH  -  '+str(row[46])
                    rr1=''
                    ss='NUMBER OF TEETH PRESENT  -  '+str(row[47])
                    ss1=''
                    tt='DECAYED TEETH  -  '+str(row[48])
                    tt1=''
                    uu='FILLED TEETH  -  '+str(row[49])
                    uu1=''
                    vv='TENDERNESS ON PERCUSSION  -  '+str(row[50])
                    vv1=''
                    ww='ROOT STUMPS  -  '+str(row[51])
                    ww1=''
                    xx='MISSING  -  '+str(row[52])
                    xx1=''
                    yy='PROVISIONAL DIAGNOSIS  -  '+str(row[53])
                    yy1=''
                    zz='FINAL DIAGNOSIS  -  '+str(row[54])
                    zz1=''
                    la='TREATMENT PLAN  -  '+str(row[55])
                    la1=''
                    la2='ADDITIONAL INFO  -  '+str(row[56])
                    la3=''
                    tqa=str(row[58])
                    topp2=str(row[59])
                    topp3=str(row[60])
                    topp4=str(row[61])
                    topp5=str(row[62])
                    topp6=str(row[63])
                    topp7=str(row[64])
                    topp8=str(row[65])
                    topp9=str(row[66])
                    topp10=str(row[67])
                    topp11=str(row[68])
                    topp12=str(row[69])
                    topp13=str(row[70])
                    topp14=str(row[71])
                    topp15=str(row[72])
                    topp16=str(row[73])
                    topp17=str(row[74])
                    topp18=str(row[75])
                    topp19=str(row[76])
                    topp20=str(row[77])
                    topp21=str(row[78])
                    topp22=str(row[79])
                    topp23=str(row[80])
                    topp24=str(row[81])
                    topp25=str(row[82])
                    topp26=str(row[83])
                    topp27=str(row[84])
                    topp28=str(row[85])
                    topp29=str(row[86])
                    topp30=str(row[87])
                    topp31=str(row[88])
                    topp32=str(row[89])
                    topp33=str(row[90])
                    topp34=str(row[91])
                    topp35=str(row[92])
                    topp36=str(row[93])
                    topp37=str(row[94])
                    topp38=str(row[95])
                    topp39=str(row[96])
                    topp40=str(row[97])
                    topp41=str(row[98])
                    topp42=str(row[99])
                    topp43=str(row[100])
                    topp44=str(row[101])
                    topp45=str(row[102])
                    topp46=str(row[103])
                    topp47=str(row[104])
                    topp48=str(row[105])
                    topp49=str(row[106])
                    topp50=str(row[107])
                    dsc1=str(row[109])
                    dsc2=str(row[110])
                    dsc3=str(row[111])
                    dsc4=str(row[112])
                    dsc5=str(row[113])
                    dsc6=str(row[114])
                    dsc7=str(row[115])
                    dsc8=str(row[116])
                    dsc9=str(row[117])
                    dsc10=str(row[118])
                    dsc11=str(row[119])
                    dsc12=str(row[120])
                    dsc13=str(row[121])
                    dsc14=str(row[122])
                    dsc15=str(row[123])
                    dsc16=str(row[124])
                    dsc17=str(row[125])
                    dsc18=str(row[126])
                    dsc19=str(row[127])
                    dsc20=str(row[128])
                    dsc21=str(row[129])
                    dsc22=str(row[130])
                    dsc23=str(row[131])
                    dsc24=str(row[132])
                    dsc25=str(row[133])
                    dsc26=str(row[134])
                    dsc27=str(row[135])
                    dsc28=str(row[136])
                    dsc29=str(row[137])
                    dsc30=str(row[138])
                    dsc31=str(row[139])
                    dsc32=str(row[140])
                    dsc33=str(row[141])
                    dsc34=str(row[142])
                    dsc35=str(row[143])
                    dsc36=str(row[144])
                    dsc37=str(row[145])
                    dsc38=str(row[146])
                    dsc39=str(row[147])
                    dsc40=str(row[148])
                    dsc41=str(row[149])
                    dsc42=str(row[150])
                    dsc43=str(row[151])
                    dsc44=str(row[152])
                    dsc45=str(row[153])
                    dsc46=str(row[154])
                    dsc47=str(row[155])
                    dsc48=str(row[156])
                    dsc49=str(row[157])
                    dsc50=str(row[158])
                    
                    insertData=(line1,line,line1,line1,pe,pe1,line1, a,a1,b,b1,c,c1,d,d1,e,e1,f,f1,g,g1,h,h1,i,i1,md,md1,md2,md2,j,j1,k,k1,l,l1,m,m1,n,n1,o,o1,p,p1,q,q1,r,r1,s,s1,t,t1,u,u1,v,v1,w,w1,x,x1,y,y1,z,z1,aa,aa1,bb,bb1,cc,cc1,dd,dd1,ee,ee1,ff,ff1,gg,gg1,hh,hh1,ii,ii1,jj,jj1,kk,kk1,ll,ll1,mm,mm1,nn,nn1,oo,oo1,pp,pp1,qq,qq1,rr,rr1,ss,ss1,tt,tt1,uu,uu1,vv,vv1,ww,ww1,xx,xx1,yy,yy1,zz,zz1,la,la1,la2,la3)
                        
                    for a in insertData:
                                         
                        text_serialno.insert(END, a+"\n")
                    
                    if tqa!='' and tqa!='None' and tqa!='none':
                        text_serialno.insert(END, '                                                                  CASE  IMAGES  '+"\n")
                        text_serialno.insert(END,'')
                        text_serialno.insert(END, '-  '+dsc1+"\n")
                        imga = Image.open(tqa)
                        imga = imga.resize((498,372))
                        imga = ImageTk.PhotoImage(imga)
                        text_serialno.image_create(END, image =imga)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imga.image=imga




                            
                    if topp2!='' and topp2!='None' and topp2!='none':
                        imgaq = Image.open(topp2)
                        imgaq = imgaq.resize((498,372))
                        imgaq = ImageTk.PhotoImage(imgaq)
                        text_serialno.insert(END, '-  '+dsc2+"\n")
                        text_serialno.image_create(END, image =imgaq)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imgaq.qimage=imgaq

                    if topp3!='' and topp3!='None' and topp3!='none':
                        imgaw = Image.open(topp3)
                        imgaw = imgaw.resize((498,372))
                        imgaw = ImageTk.PhotoImage(imgaw)
                        text_serialno.insert(END, '-  '+dsc3+"\n")
                        text_serialno.image_create(END, image =imgaw)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imgaw.image=imgaw

                    if topp4!='' and topp4!='None' and topp4!='none':
                        imgae = Image.open(topp4)
                        imgae = imgae.resize((498,372))
                        imgae = ImageTk.PhotoImage(imgae)
                        text_serialno.insert(END, '-  '+dsc4+"\n")
                        text_serialno.image_create(END, image =imgae)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imgae.image=imgae

                    if topp5!='' and topp5!='None' and topp5!='none':
                        imgar = Image.open(topp5)
                        imgar = imgar.resize((498,372))
                        imgar = ImageTk.PhotoImage(imgar)
                        text_serialno.insert(END, '-  '+dsc5+"\n")
                        text_serialno.image_create(END, image =imgar)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imgar.image=imgar

                    if topp6!='' and topp6!='None' and topp6!='none':
                        imgat = Image.open(topp6)
                        imgat = imgat.resize((498,372))
                        imgat = ImageTk.PhotoImage(imgat)
                        text_serialno.insert(END, '-  '+dsc6+"\n")
                        text_serialno.image_create(END, image =imgat)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imgat.image=imgat

                    if topp7!='' and topp7!='None' and topp7!='none':
                        imgay = Image.open(topp7)
                        imgay = imgay.resize((498,372))
                        imgay = ImageTk.PhotoImage(imgay)
                        text_serialno.insert(END, '-  '+dsc7+"\n")
                        text_serialno.image_create(END, image =imgay)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imgay.image=imgay

                    if topp8!='' and topp8!='None' and topp8!='none':
                        imgau = Image.open(topp8)
                        imgau = imgau.resize((498,372))
                        imgau = ImageTk.PhotoImage(imgau)
                        text_serialno.insert(END, '-  '+dsc8+"\n")
                        text_serialno.image_create(END, image =imgau)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imgau.image=imgau

                    if topp9!='' and topp9!='None' and topp9!='none':
                        imgai = Image.open(topp9)
                        imgai = imgai.resize((498,372))
                        imgai = ImageTk.PhotoImage(imgai)
                        text_serialno.insert(END, '-  '+dsc9+"\n")
                        text_serialno.image_create(END, image =imgai)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imgai.image=imgai

                    if topp10!='' and topp10!='None' and topp10!='none':
                        imgao = Image.open(topp10)
                        imgao = imgao.resize((498,372))
                        imgao = ImageTk.PhotoImage(imgao)
                        text_serialno.insert(END, '-  '+dsc10+"\n")
                        text_serialno.image_create(END, image =imgao)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imgao.image=imgao

                    if topp11!='' and topp11!='None' and topp11!='none':
                        imgap = Image.open(topp11)
                        imgap = imgap.resize((498,372))
                        imgap = ImageTk.PhotoImage(imgap)
                        text_serialno.insert(END, '-  '+dsc11+"\n")
                        text_serialno.image_create(END, image =imgap)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imgap.image=imgap

                    if topp12!='' and topp12!='None' and topp12!='none':
                        imgaa = Image.open(topp12)
                        imgaa = imgaa.resize((498,372))
                        imgaa = ImageTk.PhotoImage(imgaa)
                        text_serialno.insert(END, '-  '+dsc12+"\n")
                        text_serialno.image_create(END, image =imgaa)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imgaa.image=imgaa

                    if topp13!='' and topp13!='None' and topp13!='none':
                        imgas = Image.open(topp13)
                        imgas = imgas.resize((498,372))
                        imgas = ImageTk.PhotoImage(imgas)
                        text_serialno.insert(END, '-  '+dsc13+"\n")
                        text_serialno.image_create(END, image =imgas)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imgas.image=imgas

                    if topp14!='' and topp14!='None' and topp14!='none':
                        imgad = Image.open(topp14)
                        imgad = imgad.resize((498,372))
                        imgad = ImageTk.PhotoImage(imgad)
                        text_serialno.insert(END, '-  '+dsc14+"\n")
                        text_serialno.image_create(END, image =imgad)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imgad.image=imgad

                    if topp15!='' and topp15!='None' and topp15!='none':
                        imgaf = Image.open(topp15)
                        imgaf = imgaf.resize((498,372))
                        imgaf = ImageTk.PhotoImage(imgaf)
                        text_serialno.insert(END, '-  '+dsc15+"\n")
                        text_serialno.image_create(END, image =imgaf)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imgaf.image=imgaf

                    if topp16!='' and topp16!='None' and topp16!='none':
                        imgag = Image.open(topp16)
                        imgag = imgag.resize((498,372))
                        imgag = ImageTk.PhotoImage(imgag)
                        text_serialno.insert(END, '-  '+dsc16+"\n")
                        text_serialno.image_create(END, image =imgag)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imgag.image=imgag

                    if topp17!='' and topp17!='None' and topp17!='none':
                        imgah = Image.open(topp17)
                        imgah = imgah.resize((498,372))
                        imgah = ImageTk.PhotoImage(imgah)
                        text_serialno.insert(END, '-  '+dsc17+"\n")
                        text_serialno.image_create(END, image =imgah)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imgah.image=imgah

                    if topp18!='' and topp18!='None' and topp18!='none':
                        imgahj = Image.open(topp18)
                        imgahj = imgahj.resize((498,372))
                        imgahj = ImageTk.PhotoImage(imgahj)
                        text_serialno.insert(END, '-  '+dsc18+"\n")
                        text_serialno.image_create(END, image =imgahj)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imgahj.image=imgahj

                    if topp19!='' and topp19!='None' and topp19!='none':
                        imgahk = Image.open(topp19)
                        imgahk = imgahk.resize((498,372))
                        imgahk = ImageTk.PhotoImage(imgahk)
                        text_serialno.insert(END, '-  '+dsc19+"\n")
                        text_serialno.image_create(END, image =imgahk)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imgahk.image=imgahk

                    if topp20!='' and topp20!='None' and topp20!='none':
                        imgahl = Image.open(topp20)
                        imgahl = imgahl.resize((498,372))
                        imgahl = ImageTk.PhotoImage(imgahl)
                        text_serialno.insert(END, '-  '+dsc20+"\n")
                        text_serialno.image_create(END, image =imgahl)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imgahl.image=imgahl

                    if topp21!='' and topp21!='None' and topp21!='none':
                        imgahz = Image.open(topp21)
                        imgahz = imgahz.resize((498,372))
                        imgahz = ImageTk.PhotoImage(imgahz)
                        text_serialno.insert(END, '-  '+dsc21+"\n")
                        text_serialno.image_create(END, image =imgahz)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imgahz.image=imgahz

                    if topp22!='' and topp22!='None' and topp22!='none':
                        imgahx = Image.open(topp22)
                        imgahx = imgahx.resize((498,372))
                        imgahx = ImageTk.PhotoImage(imgahx)
                        text_serialno.insert(END, '-  '+dsc22+"\n")
                        text_serialno.image_create(END, image =imgahx)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imgahx.image=imgahx

                    if topp23!='' and topp23!='None' and topp23!='none':
                        imagxs = Image.open(topp23)
                        imagxs = imagxs.resize((498,372))
                        imagxs = ImageTk.PhotoImage(imagxs)
                        text_serialno.insert(END, '-  '+dsc23+"\n")
                        text_serialno.image_create(END, image =imagxs)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxs.image=imagxs

                    if topp24!='' and topp24!='None' and topp24!='none':
                        imagxq = Image.open(topp24)
                        imagxq = imagxq.resize((498,372))
                        imagxq = ImageTk.PhotoImage(imagxq)
                        text_serialno.insert(END, '-  '+dsc24+"\n")
                        text_serialno.image_create(END, image =imagxq)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxq.image=imagxq

                    if topp25!='' and topp25!='None' and topp25!='none':
                        imagxw = Image.open(topp25)
                        imagxw = imagxw.resize((498,372))
                        imagxw = ImageTk.PhotoImage(imagxw)
                        text_serialno.insert(END, '-  '+dsc25+"\n")
                        text_serialno.image_create(END, image =imagxw)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxw.image=imagxw

                    if topp26!='' and topp26!='None' and topp26!='none':
                        imagxe = Image.open(topp26)
                        imagxe = imagxe.resize((498,372))
                        imagxe = ImageTk.PhotoImage(imagxe)
                        text_serialno.insert(END, '-  '+dsc26+"\n")
                        text_serialno.image_create(END, image =imagxe)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxe.image=imagxe

                    if topp27!='' and topp27!='None' and topp27!='none':
                        imagxr = Image.open(topp27)
                        imagxr = imagxr.resize((498,372))
                        imagxr = ImageTk.PhotoImage(imagxr)
                        text_serialno.insert(END, '-  '+dsc27+"\n")
                        text_serialno.image_create(END, image =imagxr)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxr.image=imagxr

                    if topp28!='' and topp28!='None' and topp28!='none':
                        imagxt = Image.open(topp28)
                        imagxt = imagxt.resize((498,372))
                        imagxt = ImageTk.PhotoImage(imagxt)
                        text_serialno.insert(END, '-  '+dsc28+"\n")
                        text_serialno.image_create(END, image =imagxt)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxt.image=imagxt

                    if topp29!='' and topp29!='None' and topp29!='none':
                        imagxy = Image.open(topp29)
                        imagxy = imagxy.resize((498,372))
                        imagxy = ImageTk.PhotoImage(imagxy)
                        text_serialno.insert(END, '-  '+dsc29+"\n")
                        text_serialno.image_create(END, image =imagxy)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxy.image=imagxy

                    if topp30!='' and topp30!='None' and topp30!='none':
                        imagxu = Image.open(topp30)
                        imagxu = imagxu.resize((498,372))
                        imagxu = ImageTk.PhotoImage(imagxu)
                        text_serialno.insert(END, '-  '+dsc30+"\n")
                        text_serialno.image_create(END, image =imagxu)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxu.image=imagxu

                    if topp31!='' and topp31!='None' and topp31!='none':
                        imagxi = Image.open(topp31)
                        imagxi = imagxi.resize((498,372))
                        imagxi = ImageTk.PhotoImage(imagxi)
                        text_serialno.insert(END, '-  '+dsc31+"\n")
                        text_serialno.image_create(END, image =imagxi)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxi.image=imagxi

                    if topp32!='' and topp32!='None' and topp32!='none':
                        imagxo = Image.open(topp32)
                        imagxo = imagxo.resize((498,372))
                        imagxo = ImageTk.PhotoImage(imagxo)
                        text_serialno.insert(END, '-  '+dsc32+"\n")
                        text_serialno.image_create(END, image =imagxo)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxo.image=imagxo

                    if topp33!='' and topp33!='None' and topp33!='none':
                        imagxp = Image.open(topp33)
                        imagxp = imagxp.resize((498,372))
                        imagxp = ImageTk.PhotoImage(imagxp)
                        text_serialno.insert(END, '-  '+dsc33+"\n")
                        text_serialno.image_create(END, image =imagxp)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxp.image=imagxp

                    if topp34!='' and topp34!='None' and topp34!='none':
                        imagxd = Image.open(topp34)
                        imagxd = imagxd.resize((498,372))
                        imagxd = ImageTk.PhotoImage(imagxd)
                        text_serialno.insert(END, '-  '+dsc34+"\n")
                        text_serialno.image_create(END, image =imagxd)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxd.image=imagxd

                    if topp35!='' and topp35!='None' and topp35!='none':
                        imagxf = Image.open(topp35)
                        imagxf = imagxf.resize((498,372))
                        imagxf = ImageTk.PhotoImage(imagxf)
                        text_serialno.insert(END, '-  '+dsc35+"\n")
                        text_serialno.image_create(END, image =imagxf)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxf.image=imagxf

                    if topp36!='' and topp36!='None' and topp36!='none':
                        imagxg = Image.open(topp36)
                        imagxg = imagxg.resize((498,372))
                        imagxg = ImageTk.PhotoImage(imagxg)
                        text_serialno.insert(END, '-  '+dsc36+"\n")
                        text_serialno.image_create(END, image =imagxg)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxg.image=imagxg

                    if topp37!='' and topp37!='None' and topp37!='none':
                        imagxh = Image.open(topp37)
                        imagxh = imagxh.resize((498,372))
                        imagxh = ImageTk.PhotoImage(imagxh)
                        text_serialno.insert(END, '-  '+dsc37+"\n")
                        text_serialno.image_create(END, image =imagxh)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxh.image=imagxh

                    if topp38!='' and topp38!='None' and topp38!='none':
                        imagxj = Image.open(topp38)
                        imagxj = imagxj.resize((498,372))
                        imagxj = ImageTk.PhotoImage(imagxj)
                        text_serialno.insert(END, '-  '+dsc38+"\n")
                        text_serialno.image_create(END, image =imagxj)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxj.image=imagxj

                    if topp39!='' and topp39!='None' and topp39!='none':
                        imagxk = Image.open(topp39)
                        imagxk = imagxk.resize((498,372))
                        imagxk = ImageTk.PhotoImage(imagxk)
                        text_serialno.insert(END, '-  '+dsc39+"\n")
                        text_serialno.image_create(END, image =imagxk)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxk.image=imagxk

                    if topp40!='' and topp40!='None' and topp40!='none':
                        imagxv = Image.open(topp40)
                        imagxv = imagxv.resize((498,372))
                        imagxv = ImageTk.PhotoImage(imagxv)
                        text_serialno.insert(END, '-  '+dsc40+"\n")
                        text_serialno.image_create(END, image =imagxv)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxv.image=imagxv

                    if topp41!='' and topp41!='None' and topp41!='none':
                        imagxxz = Image.open(topp41)
                        imagxxz = imagxxz.resize((498,372))
                        imagxxz = ImageTk.PhotoImage(imagxxz)
                        text_serialno.insert(END, '-  '+dsc41+"\n")
                        text_serialno.image_create(END, image =imagxxz)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxxz.image=imagxxz

                    if topp42!='' and topp42!='None' and topp42!='none':
                        imagxxas = Image.open(topp42)
                        imagxxas = imagxxas.resize((498,372))
                        imagxxas = ImageTk.PhotoImage(imagxxas)
                        text_serialno.insert(END, '-  '+dsc42+"\n")
                        text_serialno.image_create(END, image =imagxxas)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxxas.image=imagxxas

                    if topp43!='' and topp43!='None' and topp43!='none':
                        imagxxadf = Image.open(topp43)
                        imagxxadf = imagxxadf.resize((498,372))
                        imagxxadf = ImageTk.PhotoImage(imagxxadf)
                        text_serialno.insert(END, '-  '+dsc43+"\n")
                        text_serialno.image_create(END, image =imagxxadf)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxxadf.image=imagxxadf

                    if topp44!='' and topp44!='None' and topp44!='none':
                        imagxxadfg = Image.open(topp44)
                        imagxxadfg = imagxxadfg.resize((498,372))
                        imagxxadfg = ImageTk.PhotoImage(imagxxadfg)
                        text_serialno.insert(END, '-  '+dsc44+"\n")
                        text_serialno.image_create(END, image =imagxxadfg)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxxadfg.image=imagxxadfg

                    if topp45!='' and topp45!='None' and topp45!='none':
                        imagxxadfee = Image.open(topp45)
                        imagxxadfee = imagxxadfee.resize((498,372))
                        imagxxadfee = ImageTk.PhotoImage(imagxxadfee)
                        text_serialno.insert(END, '-  '+dsc45+"\n")
                        text_serialno.image_create(END, image =imagxxadfee)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxxadfee.image=imagxxadfee

                    if topp46!='' and topp46!='None' and topp46!='none':
                        imagxxadferr = Image.open(topp46)
                        imagxxadferr = imagxxadferr.resize((498,372))
                        imagxxadferr = ImageTk.PhotoImage(imagxxadferr)
                        text_serialno.insert(END, '-  '+dsc46+"\n")
                        text_serialno.image_create(END, image =imagxxadferr)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxxadferr.image=imagxxadferr

                    if topp47!='' and topp47!='None' and topp47!='none':
                        imagxxadferq = Image.open(topp47)
                        imagxxadferq = imagxxadferq.resize((498,372))
                        imagxxadferq = ImageTk.PhotoImage(imagxxadferq)
                        text_serialno.insert(END, '-  '+dsc47+"\n")
                        text_serialno.image_create(END, image =imagxxadferq)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxxadferq.image=imagxxadferq

                    if topp48!='' and topp48!='None' and topp48!='none':
                        imagxxadfebv = Image.open(topp48)
                        imagxxadfebv = imagxxadfebv.resize((498,372))
                        imagxxadfebv = ImageTk.PhotoImage(imagxxadfebv)
                        text_serialno.insert(END, '-  '+dsc48+"\n")
                        text_serialno.image_create(END, image =imagxxadfebv)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxxadfebv.image=imagxxadfebv

                    if topp49!='' and topp49!='None' and topp49!='none':
                        imagxxaddg = Image.open(topp49)
                        imagxxaddg = imagxxaddg.resize((498,372))
                        imagxxaddg = ImageTk.PhotoImage(imagxxaddg)
                        text_serialno.insert(END, '-  '+dsc49+"\n")
                        text_serialno.image_create(END, image =imagxxaddg)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxxaddg.image=imagxxaddg

                    if topp50!='' and topp50!='None' and topp50!='none':
                        imagxxag = Image.open(topp50)
                        imagxxag = imagxxag.resize((498,372))
                        imagxxag = ImageTk.PhotoImage(imagxxag)
                        text_serialno.insert(END, '-  '+dsc50+"\n")
                        text_serialno.image_create(END, image =imagxxag)
                        text_serialno.insert(END, '\n')
                        text_serialno.insert(END, '\n')
                        imagxxag.image=imagxxag


                    
                    sss.delete(0,END)





                    


                text_serialno.tag_config("PzD", font = ("arial narrow BOLD",63) ,foreground="black")

                PzD = '         SK  DENTAL  CLINIC'
                    
                offset = '+%dc' % len(PzD)
                
                pos_start = text_serialno.search(PzD, '1.0', END)
                
                while pos_start:

                    pos_end = pos_start + offset
                    text_serialno.tag_add('PzD', pos_start, pos_end)
                    pos_start = text_serialno.search(PzD, pos_end, END)
                    


                text_serialno.tag_config("PD", font = "roboto 23 bold" ,foreground="black")

                PD = 'PERSONAL  DETAILS'
                    
                offset = '+%dc' % len(PD)
                
                pos_start = text_serialno.search(PD, '1.0', END)
                
                while pos_start:

                    pos_end = pos_start + offset
                    text_serialno.tag_add('PD', pos_start, pos_end)
                    pos_start = text_serialno.search(PD, pos_end, END)                      


                text_serialno.tag_config("DD", font = "roboto 23 bold" ,foreground="black")

                DD = 'DENTAL  DETAILS'
                    
                offset = '+%dc' % len(DD)
                
                pos_start = text_serialno.search(DD, '1.0', END)
                
                while pos_start:

                    pos_end = pos_start + offset
                    text_serialno.tag_add('DD', pos_start, pos_end)
                    pos_start = text_serialno.search(DD, pos_end, END)

                    

                    
                text_serialno.tag_config("SN", font = "cambria 17" ,foreground="NAVY")      
                SN = 'SERIAL NUMBER    -'
                offset = '+%dc' % len(SN)
                pos_start = text_serialno.search(SN, '1.0', END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('SN', pos_start, pos_end)
                    pos_start = text_serialno.search(SN, pos_end, END)

                    
                  
                text_serialno.tag_config('NAME', font = "cambria 17" ,foreground="NAVY")      
                NAME = 'NAME                         -'
                offset = '+%dc' % len(NAME)
                pos_start = text_serialno.search(NAME, '1.0', END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('NAME', pos_start, pos_end)
                    pos_start = text_serialno.search(NAME, pos_end, END)                                            


              
                text_serialno.tag_config('PATIENTID', font = "cambria 17" ,foreground="NAVY")      
                PATIENTID = 'PATIENT ID             -'
                offset = '+%dc' % len(PATIENTID)
                pos_start = text_serialno.search(PATIENTID, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('PATIENTID', pos_start, pos_end)
                    pos_start = text_serialno.search(PATIENTID, pos_end, END)  



                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'DATE                         -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)



                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'AGE                            -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)
                    

                        
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'SEX                             -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)



                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'OCCUPATION         -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)


                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'ADDRESS                 -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)

                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'CONTACT                -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)



                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'CHIEF COMPLAINT   -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)

                    

                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'HISTORY OF PRESENT ILLNESS  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)  



                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'MEDICAL HISTORY  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)


                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'ALLERGIES  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)


                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'DENTAL HISTORY  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)


                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'GENERAL EXAMINATION  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)


#BLOOD PRESSURE
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'BLOOD PRESSURE  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)




#PULSE RATE
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'PULSE RATE  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)


#RESPIRATORY RATE  
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'RESPIRATORY RATE  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)


#HABITS
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'HABITS  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)


#EXTRA ORAL EXAMINATION
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'EXTRA ORAL EXAMINATION  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)


#TEMPOROMANDIBULAR JOINT
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'TEMPOROMANDIBULAR JOINT  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)

                    
#EXTRA ORAL LYMPHNODES
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'EXTRA ORAL LYMPHNODES  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)
                    
#SYMMETRY OF FACE 
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'SYMMETRY OF FACE  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)



#INTRA ORAL EXAMINATION  -  
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'INTRA ORAL EXAMINATION  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)



#GINGIVA  -
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'GINGIVA  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)



#COLOR  -
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'COLOR  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)


#CONTOUR  -
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'CONTOUR  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)




#CONSISTENCY  -  
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'CONSISTENCY  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)



#SURFACE TEXTURE  -
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'SURFACE TEXTURE  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)


#BLEEDING ON PROBING  -
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'BLEEDING ON PROBING  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)


#PERIODONTAL EXAMINATION  -
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'PERIODONTAL EXAMINATION  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)



#PERIODONTAL POCKETS  - 
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'PERIODONTAL POCKETS  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)



#STAINS  -  
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'STAINS  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)


#CALCULUS  -
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'CALCULUS  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)



#FURCATION INVOLVEMENT  - 
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'FURCATION INVOLVEMENT  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)


#RECESSION  -
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'RECESSION  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)


#OTHER FINDINGS  -
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'OTHER FINDINGS  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)


#HARD PALATE  -
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'HARD PALATE  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)


#SOFT PALATE  -
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'SOFT PALATE  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)


#BUCCAL MUCOSA  -
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'BUCCAL MUCOSA  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)
                    

#TONGUE  -
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'TONGUE  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)
                    

#FLOOR OF THE MOUTH  -
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'FLOOR OF THE MOUTH  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)
                    

#TEETH EXAMINATION  -
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'TEETH EXAMINATION  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)
                    


#NATURE OF TEETH  - 
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'NATURE OF TEETH  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)
                    
#NUMBER OF TEETH PRESENT  -   
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'NUMBER OF TEETH PRESENT  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)

#DECAYED TEETH  - 
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'DECAYED TEETH  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)
                    

#FILLED TEETH  -  
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'FILLED TEETH  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)

#TENDERNESS ON PERCUSSION  - 
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'TENDERNESS ON PERCUSSION  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)

#ROOT STUMPS  -  
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'ROOT STUMPS  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)


#MISSING  -   
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'MISSING  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)


#PROVISIONAL DIAGNOSIS  -  
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'PROVISIONAL DIAGNOSIS  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)

#FINAL DIAGNOSIS  -  
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'FINAL DIAGNOSIS  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)

#TREATMENT PLAN  -   
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'TREATMENT PLAN  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)

#ADDITIONAL INFO  - 
                text_serialno.tag_config('d', font = "cambria 17" ,foreground="NAVY")      
                d = 'ADDITIONAL INFO  -'
                offset = '+%dc' % len(d)
                pos_start = text_serialno.search(d, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('d', pos_start, pos_end)
                    pos_start = text_serialno.search(d, pos_end, END)


#IMAGES  - 
                text_serialno.tag_config('da',font = "roboto 23 bold" ,foreground="black")      
                da = '  CASE  IMAGES  '
                offset = '+%dc' % len(da)
                pos_start = text_serialno.search(da, '1.0',END)

                while pos_start: 
                    pos_end = pos_start + offset
                    text_serialno.tag_add('da', pos_start, pos_end)
                    pos_start = text_serialno.search(da, pos_end, END)                        
                    
                
                text_serialno.image_create(2.22, image = img)



                


               
                menubar = Menu(root)
                root.config(menu=menubar)

                file_menu = Menu(menubar)
                menubar.add_cascade(label="File", menu=file_menu)
                file_menu.add_command(label="printer", command=locprinter)





                
            def selectItem1():
                
                curItem = treeview.focus()
                values = treeview.item(curItem)['values']
                cc=(values[1])
                sss.delete(0,END)
                sss.insert(END,cc)

                top=Toplevel()  
             
                text_serialnos=Text(top, height=20, width=80,font=("ROBOTO BOLD",15),bg='white',fg='black')
                text_serialnos.pack()

                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor(named_tuple=True)
                cursor.execute("select* from c_h1 where serial_number ='"+ sss.get()+"' ")
                text_serialnos.delete('1.0', END)

                for row in cursor:
                    
                    
                    top='     ' 
                    sx='CASE TYPE              :    '+str(row[9])
                    sy='TREATMENT TYPE   :    '+str(row[10])
                    sz='PATIENT TYPE         :    '+str(row[11])
                    lk='STATUS                    :    '+str(row[108])
                    insertsData=( top,sx,sy,sz,lk)

                    for a in insertsData:
                                         
                        text_serialnos.insert(END, a+"\n")
                    

                    sss.delete(0,END)
                       
                con.close()


            def word():
                
                curItem = treeview.focus()
                values = treeview.item(curItem)['values']
                cc=(values[1])
                sss.delete(0,END)
                sss.insert(END,cc)         
             
                text_serialsno=Text(SEARCHPATIENT, height=80, width=100,font=("ROBOTO BOLD",15),bg='white',fg='black')
                text_serialsno.place(x=10,y=100000)
                

  
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select* from c_h1 where serial_number ='"+ sss.get()+"' ")
           
                for row in cursor:
                    top='     '
                    line=''
                    line1=''
                    pe='                                          PERSONAL DETAILS'
                    pe1=''
                    a= 'SERIAL NUMBER :       '+ str(row[0])
                    az1=''
                    b='NAME :                       '+str(row[1])
                    b1=''
                    c='PATIENT ID :              '+str(row[2])
                    c1=''
                    d='DATE :                       '+str(row[3])
                    d1=''
                    e='AGE :                         '+str(row[4])
                    e1=''
                    f='SEX :                         '+str(row[5])
                    f1=''
                    g='OCCUPATION :          '+str(row[6])
                    g1=''
                    h='ADDRESS :              '+str(row[7])
                    h1=''
                    i='CONTACT :              '+str(row[8])
                    i1=''
                    md=''
                    md1='                                                 DENTAL DETAILS    '
                    md2=''
                    j='CHIEF COMPLAINT :  '+str(row[12])
                    j1=''
                    k='HISTORY OF PRESENT ILLNESS :  '+str(row[13])
                    k1=''
                    l='MEDICAL HISTORY       :  '+str(row[14])
                    l1=''
                    m='ALLERGIES       :  '+str(row[15])
                    m1=''
                    n='DENTAL HISTORY     :  '+str(row[16])
                    n1=''
                    o='GENERAL EXAMINATION      :  '+str(row[17])
                    o1=''
                    p='BLOOD PRESSURE       :  '+str(row[18])
                    p1=''
                    q='PULSE RATE       :  '+str(row[19])
                    q1=''
                    r='RESPIRATORY RATE        :  '+str(row[20])
                    r1=''
                    s='HABITS       :  '+str(row[21])
                    s1=''
                    t='EXTRA ORAL EXAMINATION       :  '+str(row[22])
                    t1=''
                    u='TEMPOROMANDIBULAR JOINT       :  '+str(row[23])
                    u1=''
                    v='EXTRA ORAL LYMPHNODES       :  '+str(row[24])
                    v1=''
                    w='SYMMETRY OF FACE       :  '+str(row[25])
                    w1=''
                    x='INTRA ORAL EXAMINATION       :  '+str(row[26])
                    x1=''
                    y='GINGIVA       :  '+str(row[27])
                    y1=''
                    z='COLOR       :  '+str(row[28])
                    z1=''
                    aa='CONTOUR       :  '+str(row[29])
                    aa1=''
                    bb='CONSISTENCY       :  '+str(row[30])
                    bb1=''
                    cc='SURFACE TEXTURE      :  '+str(row[31])
                    cc1=''
                    dd='BLEEDING ON PROBING       :  '+str(row[32])
                    dd1=''
                    ee='PERIODONTAL EXAMINATION       :  '+str(row[33])
                    ee1=''
                    ff='PERIODONTAL POCKETS      :  '+str(row[34])
                    ff1=''
                    gg='STAINS      :  '+str(row[35])
                    gg1=''
                    hh='CALCULUS       :  '+str(row[36])
                    hh1=''
                    ii='FURCATION INVOLVEMENT       :  '+str(row[37])
                    ii1=''
                    jj='RECESSION       :  '+str(row[38])
                    jj1=''
                    kk='OTHER FINDINGS      :  '+str(row[39])
                    kk1=''
                    ll='HARD PALATE       :  '+str(row[40])
                    ll1=''
                    mm='SOFT PALATE       :  '+str(row[41])
                    mm1=''
                    nn='BUCCAL MUCOSA       :  '+str(row[42])
                    nn1=''
                    oo='TONGUE       :  '+str(row[43])
                    oo1=''
                    pp='FLOOR OF THE MOUTH       :  '+str(row[44])
                    pp1=''
                    qq='TEETH EXAMINATION       :  '+str(row[45])
                    qq1=''
                    rr='NATURE OF TEETH      :  '+str(row[46])
                    rr1=''
                    ss='NUMBER OF TEETH PRESENT       :  '+str(row[47])
                    ss1=''
                    tt='DECAYED TEETH       :  '+str(row[48])
                    tt1=''
                    uu='FILLED TEETH       :  '+str(row[49])
                    uu1=''
                    vv='TENDERNESS ON PERCUSSION       :  '+str(row[50])
                    vv1=''
                    ww='ROOT STUMPS       :  '+str(row[51])
                    ww1=''
                    xx='MISSING      :  '+str(row[52])
                    xx1=''
                    yy='PROVISIONAL DIAGNOSIS       :  '+str(row[53])
                    yy1=''
                    zz='FINAL DIAGNOSIS       :  '+str(row[54])
                    zz1=''
                    la='TREATMENT PLAN       :  '+str(row[55])
                    la1=''
                    la2='ADDITIONAL INFO       :  '+str(row[56])
                    la3=''
                    a1=str(row[58])           
                    a2=str(row[59])
                    a3=str(row[60])
                    a4=str(row[61])
                    a5=str(row[62])
                    a6=str(row[63])
                    a7=str(row[64])
                    a8=str(row[65])
                    a9=str(row[66])
                    a10=str(row[67])
                    a11=str(row[68])
                    a12=str(row[69])
                    a13=str(row[70])
                    a14=str(row[71])
                    a15=str(row[72])
                    a16=str(row[73])
                    a17=str(row[74])
                    a18=str(row[75])
                    a19=str(row[76])
                    a20=str(row[77])
                    a21=str(row[78])
                    a22=str(row[79])
                    a23=str(row[80])
                    a24=str(row[81])
                    a25=str(row[82])
                    a26=str(row[83])
                    a27=str(row[84])
                    a28=str(row[85])
                    a29=str(row[86])
                    a30=str(row[87])
                    a31=str(row[88])
                    a32=str(row[89])
                    a33=str(row[90])
                    a34=str(row[91])
                    a35=str(row[92])
                    a36=str(row[93])
                    a37=str(row[94])
                    a38=str(row[95])
                    a39=str(row[96])
                    a40=str(row[97])
                    a41=str(row[98])
                    a42=str(row[99])
                    a43=str(row[100])
                    a44=str(row[101])
                    a45=str(row[102])
                    a46=str(row[103])
                    a47=str(row[104])
                    a48=str(row[105])
                    a49=str(row[106])
                    a50=str(row[107])
                    dsc1=str(row[109])
                    dsc2=str(row[110])
                    dsc3=str(row[111])
                    dsc4=str(row[112])
                    dsc5=str(row[113])
                    dsc6=str(row[114])
                    dsc7=str(row[115])
                    dsc8=str(row[116])
                    dsc9=str(row[117])
                    dsc10=str(row[118])
                    dsc11=str(row[119])
                    dsc12=str(row[120])
                    dsc13=str(row[121])
                    dsc14=str(row[122])
                    dsc15=str(row[123])
                    dsc16=str(row[124])
                    dsc17=str(row[125])
                    dsc18=str(row[126])
                    dsc19=str(row[127])
                    dsc20=str(row[128])
                    dsc21=str(row[129])
                    dsc22=str(row[130])
                    dsc23=str(row[131])
                    dsc24=str(row[132])
                    dsc25=str(row[133])
                    dsc26=str(row[134])
                    dsc27=str(row[135])
                    dsc28=str(row[136])
                    dsc29=str(row[137])
                    dsc30=str(row[138])
                    dsc31=str(row[139])
                    dsc32=str(row[140])
                    dsc33=str(row[141])
                    dsc34=str(row[142])
                    dsc35=str(row[143])
                    dsc36=str(row[144])
                    dsc37=str(row[145])
                    dsc38=str(row[146])
                    dsc39=str(row[147])
                    dsc40=str(row[148])
                    dsc41=str(row[149])
                    dsc42=str(row[150])
                    dsc43=str(row[151])
                    dsc44=str(row[152])
                    dsc45=str(row[153])
                    dsc46=str(row[154])
                    dsc47=str(row[155])
                    dsc48=str(row[156])
                    dsc49=str(row[157])
                    dsc50=str(row[158])
                    
                    insertData=(line1,line1,line1,line1,pe,pe1,line1, a,az1,b,b1,c,c1,d,d1,e,e1,f,f1,g,g1,h,h1,i,i1,md,md1,md2,md2,j,j1,k,k1,l,l1,m,m1,n,n1,o,o1,p,p1,q,q1,r,r1,s,s1,t,t1,u,u1,v,v1,w,w1,x,x1,y,y1,z,z1,aa,aa1,bb,bb1,cc,cc1,dd,dd1,ee,ee1,ff,ff1,gg,gg1,hh,hh1,ii,ii1,jj,jj1,kk,kk1,ll,ll1,mm,mm1,nn,nn1,oo,oo1,pp,pp1,qq,qq1,rr,rr1,ss,ss1,tt,tt1,uu,uu1,vv,vv1,ww,ww1,xx,xx1,yy,yy1,zz,zz1,la,la1,la2,la3)
                        
                    for a in insertData:
                                         
                        text_serialsno.insert(END, a+"\n")
                    

                    sss.delete(0,END)                    
                    
                                         
                    eb1.insert(END, a1)
                    #a=int(len(s_NAME.get()))
                    eb2.insert(END, a2)
                    #b=int(len(s_PAT.get()))
                    eb3.insert(END, a3)
                    eb4.insert(END, a4)
                    eb5.insert(END, a5)
                    eb6.insert(END, a6)
                    eb7.insert(END, a7)
                    eb8.insert(END, a8)
                    eb9.insert(END, a9)
                    eb10.insert(END, a10)
                    eb11.insert(END, a11)
                    eb12.insert(END, a12)
                    eb13.insert(END, a13)
                    eb14.insert(END, a14)
                    eb15.insert(END, a15)
                    eb16.insert(END, a16)
                    eb17.insert(END, a17)
                    eb18.insert(END, a18)
                    eb19.insert(END, a19)
                    eb20.insert(END, a20)
                    eb21.insert(END, a21)
                    eb22.insert(END, a22)
                    eb23.insert(END, a23)
                    eb24.insert(END, a24)
                    eb25.insert(END, a25)
                    eb26.insert(END, a26)
                    eb27.insert(END, a27)
                    eb28.insert(END, a28)
                    eb29.insert(END, a29)
                    eb30.insert(END, a30)
                    eb31.insert(END, a31)
                    eb32.insert(END, a32)
                    eb33.insert(END, a33)
                    eb34.insert(END, a34)
                    eb35.insert(END, a35)
                    eb36.insert(END, a36)
                    eb37.insert(END, a37)
                    eb38.insert(END, a38)
                    eb39.insert(END, a39)
                    eb40.insert(END, a40)
                    eb41.insert(END, a41)
                    eb42.insert(END, a42)
                    eb43.insert(END, a43)
                    eb44.insert(END, a44)
                    eb45.insert(END, a45)
                    eb46.insert(END, a46)
                    eb47.insert(END, a47)
                    eb48.insert(END, a48)
                    eb49.insert(END, a49)
                    eb50.insert(END, a50)
                    

                        

                
            
                document = Document()
                run = document.add_paragraph().add_run()
                style = document.styles['Normal']
                font = style.font
                font.name = 'CALIBRI'
                font.size = Pt(15)
                document.add_picture('C:/Users/Acer/Desktop/pms files/uu1.png', width=Inches(6.5),height=Inches(1.0))
            #r.add_text(' do you like it?')    for doc word                    
                document.add_paragraph(text_serialsno.get('4.0',END)).add_run()                
           
                if eb1.get() !='None' and eb1.get() !='':
                    document.add_paragraph('                                                CASE  IMAGES').add_run()
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc1).add_run()
                    document.add_picture(eb1.get(), width=Inches(7.0),height=Inches(4.0))
                    eb1.delete(0,END)

                if eb2.get() !='None' and eb2.get() !='':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc2).add_run()
                    document.add_picture(eb2.get(), width=Inches(7.0),height=Inches(4.0))
                    eb2.delete(0,END)
                if eb3.get() !='None' and eb3.get() !='':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc3).add_run()
                    document.add_picture(eb3.get(), width=Inches(7.0),height=Inches(4.0))
                    eb3.delete(0,END)
                if eb4.get() !='None' and eb4.get() !='':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc4).add_run()
                    document.add_picture(eb4.get(), width=Inches(7.0),height=Inches(4.0))
                    eb4.delete(0,END)
                if eb5.get() !='None' and eb5.get() !='':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc5).add_run()
                    document.add_picture(eb5.get(), width=Inches(7.0),height=Inches(4.0))
                    eb5.delete(0,END)
                if eb6.get() !='None' and eb6.get() !='':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc6).add_run()
                    document.add_picture(eb6.get(), width=Inches(7.0),height=Inches(4.0))
                    eb6.delete(0,END)
                if eb7.get() !='' and eb7.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc7).add_run()
                    document.add_picture(eb7.get(), width=Inches(7.0),height=Inches(4.0))
                    eb7.delete(0,END)
                if eb8.get() !='' and eb8.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc8).add_run()
                    document.add_picture(eb8.get(), width=Inches(7.0),height=Inches(4.0))
                    eb8.delete(0,END)
                if eb9.get() !='' and eb9.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc9).add_run()
                    document.add_picture(eb9.get(), width=Inches(7.0),height=Inches(4.0))
                    eb9.delete(0,END)
                if eb10.get() !='' and eb10.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc10).add_run()
                    document.add_picture(eb10.get(), width=Inches(7.0),height=Inches(4.0))
                    eb10.delete(0,END)
                if eb11.get() !='' and eb11.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc11).add_run()
                    document.add_picture(eb11.get(), width=Inches(7.0),height=Inches(4.0))
                    eb11.delete(0,END)
                if eb12.get() !='' and eb12.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc12).add_run()
                    document.add_picture(eb12.get(), width=Inches(7.0),height=Inches(4.0))
                    eb12.delete(0,END)
                if eb13.get() !='' and eb13.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc13).add_run()
                    document.add_picture(eb13.get(), width=Inches(7.0),height=Inches(4.0))
                    eb13.delete(0,END)
                if eb14.get() !='' and eb14.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc14).add_run()
                    document.add_picture(eb14.get(), width=Inches(7.0),height=Inches(4.0))
                    eb14.delete(0,END)
                if eb15.get() !='' and eb15.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc15).add_run()
                    document.add_picture(eb15.get(), width=Inches(7.0),height=Inches(4.0))
                    eb15.delete(0,END)
                if eb16.get() !='' and eb16.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc16).add_run()
                    document.add_picture(eb16.get(), width=Inches(7.0),height=Inches(4.0))
                    eb16.delete(0,END)
                if eb17.get() !='' and eb17.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc17).add_run()
                    document.add_picture(eb17.get(), width=Inches(7.0),height=Inches(4.0))
                    eb17.delete(0,END)
                if eb18.get() !='' and eb18.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc18).add_run()
                    document.add_picture(eb18.get(), width=Inches(7.0),height=Inches(4.0))
                    eb18.delete(0,END)
                if eb19.get() !='' and eb19.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc19).add_run()
                    document.add_picture(eb19.get(), width=Inches(7.0),height=Inches(4.0))
                    eb19.delete(0,END)
                if eb20.get() !='' and eb20.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc20).add_run()
                    document.add_picture(eb20.get(), width=Inches(7.0),height=Inches(4.0))
                    eb20.delete(0,END)
                if eb21.get() !='' and eb21.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc21).add_run()
                    document.add_picture(eb21.get(), width=Inches(7.0),height=Inches(4.0))
                    eb21.delete(0,END)
                if eb22.get() !='' and eb22.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc22).add_run()
                    document.add_picture(eb22.get(), width=Inches(7.0),height=Inches(4.0))
                    eb22.delete(0,END)
                if eb23.get() !='' and eb23.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc23).add_run()
                    document.add_picture(eb23.get(), width=Inches(7.0),height=Inches(4.0))
                    eb23.delete(0,END)
                if eb24.get() !='' and eb24.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc24).add_run()
                    document.add_picture(eb24.get(), width=Inches(7.0),height=Inches(4.0))
                    eb24.delete(0,END)
                if eb25.get() !='' and eb25.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc25).add_run()
                    document.add_picture(eb25.get(), width=Inches(7.0),height=Inches(4.0))
                    eb25.delete(0,END)
                if eb26.get() !='' and eb26.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc26).add_run()
                    document.add_picture(eb26.get(), width=Inches(7.0),height=Inches(4.0))
                    eb26.delete(0,END)
                if eb27.get() !='' and eb27.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc27).add_run()
                    document.add_picture(eb27.get(), width=Inches(7.0),height=Inches(4.0))
                    eb27.delete(0,END)
                if eb28.get() !='' and eb28.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc28).add_run()
                    document.add_picture(eb28.get(), width=Inches(7.0),height=Inches(4.0))
                    eb28.delete(0,END)
                if eb29.get() !='' and eb29.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc29).add_run()
                    document.add_picture(eb29.get(), width=Inches(7.0),height=Inches(4.0))
                    eb29.delete(0,END)
                if eb30.get() !='' and eb30.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc30).add_run()
                    document.add_picture(eb30.get(), width=Inches(7.0),height=Inches(4.0))
                    eb30.delete(0,END)
                if eb31.get() !='' and eb31.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc31).add_run()
                    document.add_picture(eb31.get(), width=Inches(7.0),height=Inches(4.0))
                    eb31.delete(0,END)
                if eb32.get() !='' and eb32.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc32).add_run()
                    document.add_picture(eb32.get(), width=Inches(7.0),height=Inches(4.0))
                    eb32.delete(0,END)
                if eb33.get() !='' and eb33.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc33).add_run()
                    document.add_picture(eb33.get(), width=Inches(7.0),height=Inches(4.0))
                    eb33.delete(0,END)
                if eb34.get() !='' and eb34.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc34).add_run()
                    document.add_picture(eb34.get(), width=Inches(7.0),height=Inches(4.0))
                    eb34.delete(0,END)
                if eb35.get() !='' and eb35.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc35).add_run()
                    document.add_picture(eb35.get(), width=Inches(7.0),height=Inches(4.0))
                    eb35.delete(0,END)
                if eb36.get() !='' and eb36.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc36).add_run()
                    document.add_picture(eb36.get(), width=Inches(7.0),height=Inches(4.0))
                    eb36.delete(0,END)
                if eb37.get() !='' and eb37.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc37).add_run()
                    document.add_picture(eb37.get(), width=Inches(7.0),height=Inches(4.0))
                    eb37.delete(0,END)
                if eb38.get() !='' and eb38.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc38).add_run()
                    document.add_picture(eb38.get(), width=Inches(7.0),height=Inches(4.0))
                    eb38.delete(0,END)
                if eb39.get() !='' and eb39.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc39).add_run()
                    document.add_picture(eb39.get(), width=Inches(7.0),height=Inches(4.0))
                    eb39.delete(0,END)
                if eb40.get() !='' and eb40.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc40).add_run()
                    document.add_picture(eb40.get(), width=Inches(7.0),height=Inches(4.0))
                    eb40.delete(0,END)
                if eb41.get() !='' and eb41.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc41).add_run()
                    document.add_picture(eb41.get(), width=Inches(7.0),height=Inches(4.0))
                    eb41.delete(0,END)
                if eb42.get() !='' and eb42.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc42).add_run()
                    document.add_picture(eb42.get(), width=Inches(7.0),height=Inches(4.0))
                    eb42.delete(0,END)
                if eb43.get() !='' and eb43.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc43).add_run()
                    document.add_picture(eb43.get(), width=Inches(7.0),height=Inches(4.0))
                    eb43.delete(0,END)
                if eb44.get() !='' and eb44.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc44).add_run()
                    document.add_picture(eb44.get(), width=Inches(7.0),height=Inches(4.0))
                    eb44.delete(0,END)
                if eb45.get() !='' and eb45.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc45).add_run()
                    document.add_picture(eb45.get(), width=Inches(7.0),height=Inches(4.0))
                    eb45.delete(0,END)
                if eb46.get() !='' and eb46.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc46).add_run()
                    document.add_picture(eb46.get(), width=Inches(7.0),height=Inches(4.0))
                    eb46.delete(0,END)
                if eb47.get() !='' and eb47.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc47).add_run()
                    document.add_picture(eb47.get(), width=Inches(7.0),height=Inches(4.0))
                    eb47.delete(0,END)
                if eb48.get() !='' and eb48.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc48).add_run()
                    document.add_picture(eb48.get(), width=Inches(7.0),height=Inches(4.0))
                    eb48.delete(0,END)
                if eb49.get() !='' and eb49.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc49).add_run()
                    document.add_picture(eb49.get(), width=Inches(7.0),height=Inches(4.0))
                    eb49.delete(0,END)
                if eb50.get() !='' and eb50.get() !='None':
                    document.add_paragraph('').add_run()
                    document.add_paragraph('- '+dsc50).add_run()
                    document.add_picture(eb50.get(), width=Inches(7.0),height=Inches(4.0))
                    eb50.delete(0,END)
                    
                
                files = [('Word File', '*.docx')] 
                file = fd.asksaveasfilename(filetypes = files, defaultextension = files)


            
                document.save(file)
                con.close()


       # word photos

            inputa = StringVar()
            eb1= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputa)
            eb1.place(x=1000,y=3000)

            inputb = StringVar()
            eb2= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputb)
            eb2.place(x=1000,y=3000)

            inputc = StringVar()
            eb3= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputc)
            eb3.place(x=1000,y=3000)

            inputd = StringVar()
            eb4= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputd)
            eb4.place(x=1000,y=3000)

            inpute = StringVar()
            eb5= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inpute)
            eb5.place(x=1000,y=3000)

            inputf = StringVar()
            eb6= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputf)
            eb6.place(x=1000,y=3000)

            inputg = StringVar()
            eb7= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputg)
            eb7.place(x=1000,y=3000)

            inputh = StringVar()
            eb8= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputh)
            eb8.place(x=1000,y=3000)

            inputi = StringVar()
            eb9= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputi)
            eb9.place(x=1000,y=3000)

            inputj = StringVar()
            eb10= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputj)
            eb10.place(x=1000,y=3000)

            inputk = StringVar()
            eb11= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputk)
            eb11.place(x=1000,y=3000)

            inputl = StringVar()
            eb12= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputl)
            eb12.place(x=1000,y=3000)

            inputm = StringVar()
            eb13= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputm)
            eb13.place(x=1000,y=3000)

            inputn = StringVar()
            eb14= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputn)
            eb14.place(x=1000,y=3000)

            inputo = StringVar()
            eb15= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputo)
            eb15.place(x=1000,y=3000)

            inputp = StringVar()
            eb16= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputp)
            eb16.place(x=1000,y=3000)

            inputq = StringVar()
            eb17= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputq)
            eb17.place(x=1000,y=3000)

            inputr = StringVar()
            eb18= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputr)
            eb18.place(x=1000,y=3000) 

            inputs = StringVar()
            eb19= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputs)
            eb19.place(x=1000,y=3000)

            inputt = StringVar()
            eb20= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputt)
            eb1.place(x=1000,y=3000)

            inputu = StringVar()
            eb21= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputu)
            eb21.place(x=1000,y=3000)

            inputv = StringVar()
            eb22= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputv)
            eb22.place(x=1000,y=3000)

            inputw = StringVar()
            eb23= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputw)
            eb23.place(x=1000,y=3000)

            inputx = StringVar()
            eb24= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputx)
            eb24.place(x=1000,y=3000)
            
            inputy = StringVar()
            eb25= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputy)
            eb25.place(x=1000,y=3000)

            inputz = StringVar()
            eb26= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputz)
            eb26.place(x=1000,y=3000)

            inputaa = StringVar()
            eb27= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputaa)
            eb27.place(x=1000,y=3000)

            inputbb = StringVar()
            eb28= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputbb)
            eb28.place(x=1000,y=3000)

            inputcc = StringVar()
            eb29= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputcc)
            eb29.place(x=1000,y=3000)

            inputdd = StringVar()
            eb30= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputdd)
            eb30.place(x=1000,y=3000)

            inputee = StringVar()
            eb31= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputee)
            eb31.place(x=1000,y=3000)

            inputff = StringVar()
            eb32= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputff)
            eb32.place(x=1000,y=3000)

            inputgg = StringVar()
            eb33= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputgg)
            eb33.place(x=1000,y=3000)

            inputhh = StringVar()
            eb34= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputhh)
            eb34.place(x=1000,y=3000)

            inputii = StringVar()
            eb35= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputii)
            eb35.place(x=1000,y=3000)

            inputjj = StringVar()
            eb36= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputjj)
            eb36.place(x=1000,y=3000)

            inputkk = StringVar()
            eb37= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputkk)
            eb37.place(x=1000,y=3000)

            inputll = StringVar()
            eb38= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputll)
            eb38.place(x=1000,y=3000)

            inputmm = StringVar()
            eb39= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputmm)
            eb39.place(x=1000,y=3000)

            inputnn = StringVar()
            eb40= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputnn)
            eb40.place(x=1000,y=3000)

            inputoo = StringVar()
            eb41= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputoo)
            eb41.place(x=1000,y=3000)

            inputpp = StringVar()
            eb42= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputpp)
            eb42.place(x=1000,y=3000)

            inputqq = StringVar()
            eb43= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputqq)
            eb43.place(x=1000,y=3000)
            
            inputrr = StringVar()
            eb44= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputrr)
            eb44.place(x=1000,y=3000)

            inputss = StringVar()
            eb45= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputss)
            eb45.place(x=1000,y=3000)

            inputtt = StringVar()
            eb46= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputtt)
            eb46.place(x=1000,y=3000)
            
            inputuu = StringVar()
            eb47= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputuu)
            eb47.place(x=1000,y=3000)

            inputvv = StringVar()
            eb48= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputvv)
            eb48.place(x=1000,y=3000)

            inputww = StringVar()
            eb49= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputww)
            eb49.place(x=1000,y=3000)

            inputxx = StringVar()
            eb50= Entry(SEARCHPATIENT,font=('ARIAL BOLD',1),bg='white',fg='black',textvariable=inputxx)
            eb50.place(x=1000,y=3000)

          


            def cc_sn():
                cleartreeview()
                treeview.insert('', 'end', text='',values=('','','','','','','',''))  
                con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursor=con.cursor()
                cursor.execute("select* from c_h1 where "+saortaa.get()+" like '"+sx.get()+"%' ")

                rows=cursor.fetchall()
                
    #very long..
                for row in rows:
                    treeview.insert('', 'end', text='                                                                                                                                       '+str(row[1]),values=('|',str(row[0]),'|   '+str(row[2]),'|    '+str(row[3]), '|     '+str(row[4]),'|      '+str(row[5]),'|     '+str(row[8]),'|     '+str(row[108])))

                

            def delete():
                
                curItem = treeview.focus()
                values = treeview.item(curItem)['values']
                cc=(values[1])
                sss.insert(END,cc) 

                MsgBox = MessageBox.askquestion ('DELETE RECORD','Are you sure you want to Delete Record?',icon = 'warning')

                if MsgBox == 'yes':                 
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("delete from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")
                    cursor.execute("commit");  
                                
                    MessageBox.showinfo("Delete Status","Deleted Successfully");
                    con.close()
                    sss.delete(0,END)

            def tabb():

                curItem = treeview.focus()
                values = treeview.item(curItem)['values']
                cc=(values[1])
                sss.delete(0,END)
                sss.insert(END,cc)
                top=Toplevel()
                top.state('zoomed')
                nb1 = ttk.Notebook(top)                         #######   NOTEBOOK   #######

                PERSONAL_DETAILS = ttk.Frame(nb1)        
                MEDICAL_DETAILS = ttk.Frame(nb1)
                OTHER_DETAILS   = ttk.Frame(nb1)
                IMAGES   = ttk.Frame(nb1)
          

                nb1.add(PERSONAL_DETAILS, text='                                          PERSONAL   DETAILS                                                                     ')
                nb1.add(MEDICAL_DETAILS, text='                                    MEDICAL    DETAILS                                                          ')
                nb1.add(OTHER_DETAILS , text='                          OTHER   DETAILS                                         ')               
                nb1.add(IMAGES , text='                                                    IMAGES                                                 ')
                nb1.pack(expand=1, fill="both")

                


         #####################        PERSONAL   DETAILS       ####################
          
                frame=Frame(IMAGES,width=500,height=800)
                frame.pack(expand=True, fill=BOTH) 

                canvas=tk.Canvas(frame,width=300,height=500,scrollregion=(0,0,1500,3400))
                canvas.configure(bg='whitesmoke')
                

                img_desc1=Text(canvas,width=26,height=3)
                canvas.create_window((138,273),window=img_desc1)

                img_desc2=Text(canvas,width=26,height=3)
                canvas.create_window((438,273),window=img_desc2)


                img_desc3=Text(canvas,width=26,height=3)
                canvas.create_window((738,273),window=img_desc3)            


                img_desc4=Text(canvas,width=26,height=3)
                canvas.create_window((1038,273),window=img_desc4)


                img_desc5=Text(canvas,width=26,height=3)
                canvas.create_window((1338,273),window=img_desc5)


                img_desc6=Text(canvas,width=26,height=3)
                canvas.create_window((138,603),window=img_desc6)

                img_desc7=Text(canvas,width=26,height=3)
                canvas.create_window((438,603),window=img_desc7)

                img_desc8=Text(canvas,width=26,height=3)
                canvas.create_window((738,603),window=img_desc8)

                img_desc9=Text(canvas,width=26,height=3)
                canvas.create_window((1038,603),window=img_desc9)            

                img_desc10=Text(canvas,width=26,height=3)
                canvas.create_window((1338,603),window=img_desc10)

                img_desc11=Text(canvas,width=26,height=3)
                canvas.create_window((138,933),window=img_desc11)

                img_desc12=Text(canvas,width=26,height=3)
                canvas.create_window((438,933),window=img_desc12)


                img_desc13=Text(canvas,width=26,height=3)
                canvas.create_window((738,933),window=img_desc13)            


                img_desc14=Text(canvas,width=26,height=3)
                canvas.create_window((1038,933),window=img_desc14)


                img_desc15=Text(canvas,width=26,height=3)
                canvas.create_window((1338,933),window=img_desc15)


                img_desc16=Text(canvas,width=26,height=3)
                canvas.create_window((138,1263),window=img_desc16)

                img_desc17=Text(canvas,width=26,height=3)
                canvas.create_window((438,1263),window=img_desc17)

                img_desc18=Text(canvas,width=26,height=3)
                canvas.create_window((738,1263),window=img_desc18)

                img_desc19=Text(canvas,width=26,height=3)
                canvas.create_window((1038,1263),window=img_desc19)            

                img_desc20=Text(canvas,width=26,height=3)
                canvas.create_window((1338,1263),window=img_desc20)

                img_desc21=Text(canvas,width=26,height=3)
                canvas.create_window((138,1593),window=img_desc21)

                img_desc22=Text(canvas,width=26,height=3)
                canvas.create_window((438,1593),window=img_desc22)


                img_desc23=Text(canvas,width=26,height=3)
                canvas.create_window((738,1593),window=img_desc23)            


                img_desc24=Text(canvas,width=26,height=3)
                canvas.create_window((1038,1593),window=img_desc24)

                img_desc25=Text(canvas,width=26,height=3)
                canvas.create_window((1338,1593),window=img_desc25)


                img_desc26=Text(canvas,width=26,height=3)
                canvas.create_window((138,1923),window=img_desc26)

                img_desc27=Text(canvas,width=26,height=3)
                canvas.create_window((438,1923),window=img_desc27)

                img_desc28=Text(canvas,width=26,height=3)
                canvas.create_window((738,1923),window=img_desc28)

                img_desc29=Text(canvas,width=26,height=3)
                canvas.create_window((1038,1923),window=img_desc29)            

                img_desc30=Text(canvas,width=26,height=3)
                canvas.create_window((1338,1923),window=img_desc30)

                img_desc31=Text(canvas,width=26,height=3)
                canvas.create_window((138,2253),window=img_desc31)

                img_desc32=Text(canvas,width=26,height=3)
                canvas.create_window((438,2253),window=img_desc32)

                img_desc33=Text(canvas,width=26,height=3)
                canvas.create_window((738,2253),window=img_desc33)            

                img_desc34=Text(canvas,width=26,height=3)
                canvas.create_window((1038,2253),window=img_desc34)

                img_desc35=Text(canvas,width=26,height=3)
                canvas.create_window((1338,2253),window=img_desc35)

                img_desc36=Text(canvas,width=26,height=3)
                canvas.create_window((138,2583),window=img_desc36)

                img_desc37=Text(canvas,width=26,height=3)
                canvas.create_window((438,2583),window=img_desc37)

                img_desc38=Text(canvas,width=26,height=3)
                canvas.create_window((738,2583),window=img_desc38)

                img_desc39=Text(canvas,width=26,height=3)
                canvas.create_window((1038,2583),window=img_desc39)            

                img_desc40=Text(canvas,width=26,height=3)
                canvas.create_window((1338,2583),window=img_desc40)

                img_desc41=Text(canvas,width=26,height=3)
                canvas.create_window((138,2913),window=img_desc41)

                img_desc42=Text(canvas,width=26,height=3)
                canvas.create_window((438,2913),window=img_desc42)

                img_desc43=Text(canvas,width=26,height=3)
                canvas.create_window((738,2913),window=img_desc43)            

                img_desc44=Text(canvas,width=26,height=3)
                canvas.create_window((1038,2913),window=img_desc44)

                img_desc45=Text(canvas,width=26,height=3)
                canvas.create_window((1338,2913),window=img_desc45)

                img_desc46=Text(canvas,width=26,height=3)
                canvas.create_window((138,3243),window=img_desc46)

                img_desc47=Text(canvas,width=26,height=3)
                canvas.create_window((438,3243),window=img_desc47)

                img_desc48=Text(canvas,width=26,height=3)
                canvas.create_window((738,3243),window=img_desc48)

                img_desc49=Text(canvas,width=26,height=3)
                canvas.create_window((1038,3243),window=img_desc49)            

                img_desc50=Text(canvas,width=26,height=3)
                canvas.create_window((1338,3243),window=img_desc50)



                
                tr1=Text(canvas,width=24,height=11)
                canvas.create_window((138,130),window=tr1)
                
                tr2=Text(canvas,width=24,height=11)
                canvas.create_window((438,130),window=tr2)
                
                tr3=Text(canvas,width=24,height=11)
                canvas.create_window((738,130),window=tr3)
                
                tr4=Text(canvas,width=24,height=11)
                canvas.create_window((1038,130),window=tr4)                
                
                tr5=Text(canvas,width=24,height=11)
                canvas.create_window((1338,130),window=tr5)
                
                tr6=Text(canvas,width=24,height=11)
                canvas.create_window((138,460),window=tr6)
                
                tr7=Text(canvas,width=24,height=11)
                canvas.create_window((438,460),window=tr7)
                
                tr8=Text(canvas,width=24,height=11)
                canvas.create_window((738,460),window=tr8)
                
                tr9=Text(canvas,width=24,height=11)
                canvas.create_window((1038,460),window=tr9)                
                
                tr10=Text(canvas,width=24,height=11)
                canvas.create_window((1338,460),window=tr10)
                
                tr11=Text(canvas,width=24,height=11)
                canvas.create_window((138,790),window=tr11)
                
                tr12=Text(canvas,width=24,height=11)
                canvas.create_window((438,790),window=tr12)
                
                tr13=Text(canvas,width=24,height=11)
                canvas.create_window((738,790),window=tr13)
                
                tr14=Text(canvas,width=24,height=11)
                canvas.create_window((1038,790),window=tr14)                
                
                tr15=Text(canvas,width=24,height=11)
                canvas.create_window((1338,790),window=tr15)
                
                tr16=Text(canvas,width=24,height=11)
                canvas.create_window((138,1120),window=tr16)
                
                tr17=Text(canvas,width=24,height=11)
                canvas.create_window((438,1120),window=tr17)
                
                tr18=Text(canvas,width=24,height=11)
                canvas.create_window((738,1120),window=tr18)
                
                tr19=Text(canvas,width=24,height=11)
                canvas.create_window((1038,1120),window=tr19)                
                
                tr20=Text(canvas,width=24,height=11)
                canvas.create_window((1338,1120),window=tr20)
                
                tr21=Text(canvas,width=24,height=11)
                canvas.create_window((138,1450),window=tr21)
                
                tr22=Text(canvas,width=24,height=11)
                canvas.create_window((438,1450),window=tr22)
                
                tr23=Text(canvas,width=24,height=11)
                canvas.create_window((738,1450),window=tr23)
                
                tr24=Text(canvas,width=24,height=11)
                canvas.create_window((1038,1450),window=tr24)                
                
                tr25=Text(canvas,width=24,height=11)
                canvas.create_window((1338,1450),window=tr25)
                
                tr26=Text(canvas,width=24,height=11)
                canvas.create_window((138,1780),window=tr26)
                
                tr27=Text(canvas,width=24,height=11)
                canvas.create_window((438,1780),window=tr27)
                
                tr28=Text(canvas,width=24,height=11)
                canvas.create_window((738,1780),window=tr28)
                
                tr29=Text(canvas,width=24,height=11)
                canvas.create_window((1038,1780),window=tr29)                
                
                tr30=Text(canvas,width=24,height=11)
                canvas.create_window((1338,1780),window=tr30)
                
                tr31=Text(canvas,width=24,height=11)
                canvas.create_window((138,2110),window=tr31)
                
                tr32=Text(canvas,width=24,height=11)
                canvas.create_window((438,2110),window=tr32)
                
                tr33=Text(canvas,width=24,height=11)
                canvas.create_window((738,2110),window=tr33)
                
                tr34=Text(canvas,width=24,height=11)
                canvas.create_window((1038,2110),window=tr34)                
                
                tr35=Text(canvas,width=24,height=11)
                canvas.create_window((1338,2110),window=tr35)
                
                tr36=Text(canvas,width=24,height=11)
                canvas.create_window((138,2440),window=tr36)
                
                tr37=Text(canvas,width=24,height=11)
                canvas.create_window((438,2440),window=tr37)
                
                tr38=Text(canvas,width=24,height=11)
                canvas.create_window((738,2440),window=tr38)
            
                tr39=Text(canvas,width=24,height=11)
                canvas.create_window((1038,2440),window=tr39)                
                
                tr40=Text(canvas,width=24,height=11)
                canvas.create_window((1338,2440),window=tr40)
                
                tr41=Text(canvas,width=24,height=11)
                canvas.create_window((138,2770),window=tr41)
                
                tr42=Text(canvas,width=24,height=11)
                canvas.create_window((438,2770),window=tr42)
                
                tr43=Text(canvas,width=24,height=11)
                canvas.create_window((738,2770),window=tr43)
                
                tr44=Text(canvas,width=24,height=11)
                canvas.create_window((1038,2770),window=tr44)                
                
                tr45=Text(canvas,width=24,height=11)
                canvas.create_window((1338,2770),window=tr45)

                tr46=Text(canvas,width=24,height=11)
                canvas.create_window((138,3100),window=tr46)                
                
                tr47=Text(canvas,width=24,height=11)
                canvas.create_window((438,3100),window=tr47)

                tr48=Text(canvas,width=24,height=11)
                canvas.create_window((738,3100),window=tr48)                
                
                tr49=Text(canvas,width=24,height=11)
                canvas.create_window((1032,3100),window=tr49)

                tr50=Text(canvas,width=24,height=11)
                canvas.create_window((1338,3100),window=tr50)

                
                def phot1():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[58])
                        am1=str(row[109])
                        q1.insert(END,ma1)
                        img_desc1.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr1.image=imagxxadfeae                
                            tr1.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=fa1)
                            aA.image=p
                            canvas.create_window((138,130),window=aA)
                            p.image = p
                    con.close()
                    
                def phot1_1():                    

                        p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                        p = p.subsample(1, 1) 
                        azA=Button(canvas, text = '', image = p,bg='white',command=fa1)
                        azA.image=p
                        canvas.create_window((138,130),window=azA)
                        p.image = p                    

                
                def phot2():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[59])
                        am1=str(row[110])
                        q2.insert(END,ma1)
                        img_desc2.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr2.image=imagxxadfeae                
                            tr2.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f2)
                            aA.image=p
                            canvas.create_window((438,130),window=aA)
                            p.image = p
                    con.close()

                def phot2_2():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f2)
                    aA.image=p
                    canvas.create_window((438,130),window=aA)
                    p.image = p

                
                def phot3():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[60])
                        am1=str(row[111])
                        q3.insert(END,ma1)
                        img_desc3.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr3.image=imagxxadfeae                
                            tr3.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f3)
                            aA.image=p
                            canvas.create_window((738,130),window=aA)
                            p.image = p
                    con.close()

                    
                def phot3_3():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f3)
                    aA.image=p
                    canvas.create_window((738,130),window=aA)
                    p.image = p


                
                def phot4():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[61])
                        am1=str(row[112])
                        q4.insert(END,ma1)
                        img_desc4.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr4.image=imagxxadfeae                
                            tr4.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f4)
                            aA.image=p
                            canvas.create_window((1038,130),window=aA)
                            p.image = p
                    con.close()

                    
                def phot4_4():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f4)
                    aA.image=p
                    canvas.create_window((1038,130),window=aA)
                    p.image = p

                
                def phot5():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[62])
                        am1=str(row[113])
                        q5.insert(END,ma1)
                        img_desc5.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr5.image=imagxxadfeae                
                            tr5.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f5)
                            aA.image=p
                            canvas.create_window((1338,130),window=aA)
                            p.image = p
                    con.close()

                    
                def phot5_5():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f5)
                    aA.image=p
                    canvas.create_window((1338,130),window=aA)
                    p.image = p

                
                def phot6():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[63])
                        am1=str(row[114])
                        q6.insert(END,ma1)
                        img_desc6.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr6.image=imagxxadfeae                
                            tr6.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f6)
                            aA.image=p
                            canvas.create_window((138,460),window=aA)
                            p.image = p
                    con.close()

                    
                def phot6_6():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f6)
                    aA.image=p
                    canvas.create_window((138,460),window=aA)
                    p.image = p

                
                def phot7():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[64])
                        am1=str(row[115])
                        q7.insert(END,ma1)
                        img_desc7.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr7.image=imagxxadfeae                
                            tr7.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f7)
                            aA.image=p
                            canvas.create_window((438,460),window=aA)
                            p.image = p
                    con.close()
                    

                def phot7_7():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f7)
                    aA.image=p
                    canvas.create_window((438,460),window=aA)
                    p.image = p

                
                def phot8():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[65])
                        am1=str(row[116])
                        q8.insert(END,ma1)
                        img_desc8.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr8.image=imagxxadfeae                
                            tr8.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f8)
                            aA.image=p
                            canvas.create_window((738,460),window=aA)
                            p.image = p
                    con.close()
                    
                    
                def phot8_8():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f8)
                    aA.image=p
                    canvas.create_window((738,460),window=aA)
                    p.image = p
                
                def phot9():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[66])
                        am1=str(row[117])
                        q9.insert(END,ma1)
                        img_desc9.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr9.image=imagxxadfeae                
                            tr9.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f9)
                            aA.image=p
                            canvas.create_window((1038,460),window=aA)
                            p.image = p
                    con.close()
                    

                def phot9_9():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f9)
                    aA.image=p
                    canvas.create_window((1038,460),window=aA)
                    p.image = p
                
                def phot10():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[67])
                        am1=str(row[118])
                        q10.insert(END,ma1)
                        img_desc10.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr10.image=imagxxadfeae                
                            tr10.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f10)
                            aA.image=p
                            canvas.create_window((1338,460),window=aA)
                            p.image = p
                    con.close()
                    

                def phot10_10():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f10)
                    aA.image=p
                    canvas.create_window((1338,460),window=aA)
                    p.image = p                
                
                def phot11():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[68])
                        am1=str(row[119])
                        q11.insert(END,ma1)
                        img_desc11.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr11.image=imagxxadfeae                
                            tr11.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=fa11)
                            aA.image=p
                            canvas.create_window((138,790),window=aA)
                            p.image = p
                    con.close()
                    
                    
                def phot11_11():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=fa11)
                    aA.image=p
                    canvas.create_window((138,790),window=aA)
                    p.image = p
                
                def phot12():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[69])
                        am1=str(row[120])
                        q12.insert(END,ma1)
                        img_desc12.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr12.image=imagxxadfeae                
                            tr12.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f12)
                            aA.image=p
                            canvas.create_window((438,790),window=aA)
                            p.image = p
                    con.close()
                    

                def phot12_12():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f12)
                    aA.image=p
                    canvas.create_window((438,790),window=aA)
                    p.image = p
                
                def phot13():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[70])
                        am1=str(row[121])
                        q13.insert(END,ma1)
                        img_desc13.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr13.image=imagxxadfeae                
                            tr13.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f13)
                            aA.image=p
                            canvas.create_window((738,790),window=aA)
                            p.image = p
                    con.close()
                    

                def phot13_13():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f13)
                    aA.image=p
                    canvas.create_window((738,790),window=aA)
                    p.image = p
                
                def phot14():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[71])
                        am1=str(row[122])
                        q14.insert(END,ma1)
                        img_desc14.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr14.image=imagxxadfeae                
                            tr14.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f14)
                            aA.image=p
                            canvas.create_window((1038,790),window=aA)
                            p.image = p
                    con.close()
                    


                def phot14_14():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f14)
                    aA.image=p
                    canvas.create_window((1038,790),window=aA)
                    p.image = p
                
                def phot15():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[72])
                        am1=str(row[123])
                        q15.insert(END,ma1)
                        img_desc15.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr15.image=imagxxadfeae                
                            tr15.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f15)
                            aA.image=p
                            canvas.create_window((1338,790),window=aA)
                            p.image = p
                    con.close()
                    

                def phot15_15():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f15)
                    aA.image=p
                    canvas.create_window((1338,790),window=aA)
                    p.image = p
                
                def phot16():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[73])
                        am1=str(row[124])
                        q16.insert(END,ma1)
                        img_desc16.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr16.image=imagxxadfeae                
                            tr16.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f16)
                            aA.image=p
                            canvas.create_window((138,1120),window=aA)
                            p.image = p
                    con.close()
                    

                def phot16_16():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f16)
                    aA.image=p
                    canvas.create_window((138,1120),window=aA)
                    p.image = p
                
                def phot17():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[74])
                        am1=str(row[125])
                        q17.insert(END,ma1)
                        img_desc17.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr17.image=imagxxadfeae                
                            tr17.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f17)
                            aA.image=p
                            canvas.create_window((438,1120),window=aA)
                            p.image = p
                    con.close()
                    


                def phot17_17():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f17)
                    aA.image=p
                    canvas.create_window((438,1120),window=aA)
                    p.image = p
                
                def phot18():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[75])
                        am1=str(row[126])
                        q18.insert(END,ma1)
                        img_desc18.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr18.image=imagxxadfeae                
                            tr18.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f18)
                            aA.image=p
                            canvas.create_window((738,1120),window=aA)
                            p.image = p
                    con.close()
                    

                    
                def phot18_18():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f18)
                    aA.image=p
                    canvas.create_window((738,1120),window=aA)
                    p.image = p
                
                def phot19():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[76])
                        am1=str(row[127])
                        q19.insert(END,ma1)
                        img_desc19.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr19.image=imagxxadfeae                
                            tr19.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f19)
                            aA.image=p
                            canvas.create_window((1038,1120),window=aA)
                            p.image = p
                    con.close()
                    

                def phot19_19():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f19)
                    aA.image=p
                    canvas.create_window((1038,1120),window=aA)
                    p.image = p
                
                def phot20():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[77])
                        am1=str(row[128])
                        q20.insert(END,ma1)
                        img_desc20.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr20.image=imagxxadfeae                
                            tr20.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f20)
                            aA.image=p
                            canvas.create_window((1338,1120),window=aA)
                            p.image = p
                    con.close()
                    

                def phot20_20():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f20)
                    aA.image=p
                    canvas.create_window((1338,1120),window=aA)
                    p.image = p
                
                def phot21():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[78])
                        am1=str(row[129])
                        q21.insert(END,ma1)
                        img_desc21.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr21.image=imagxxadfeae                
                            tr21.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f21)
                            aA.image=p
                            canvas.create_window((138,1450),window=aA)
                            p.image = p
                    con.close()
                    

                def phot21_21():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f21)
                    aA.image=p
                    canvas.create_window((138,1450),window=aA)
                    p.image = p
                
                def phot22():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[79])
                        am1=str(row[130])
                        q22.insert(END,ma1)
                        img_desc22.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr22.image=imagxxadfeae                
                            tr22.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f22)
                            aA.image=p
                            canvas.create_window((438,1450),window=aA)
                            p.image = p
                    con.close()
                    

                def phot22_22():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f22)
                    aA.image=p
                    canvas.create_window((438,1450),window=aA)
                    p.image = p
                
                def phot23():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[80])
                        am1=str(row[131])
                        q23.insert(END,ma1)
                        img_desc23.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr23.image=imagxxadfeae                
                            tr23.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f23)
                            aA.image=p
                            canvas.create_window((738,1450),window=aA)
                            p.image = p
                    con.close()
                    

                def phot23_23():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f23)
                    aA.image=p
                    canvas.create_window((738,1450),window=aA)
                    p.image = p

                
                def phot24():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[81])
                        am1=str(row[132])
                        q24.insert(END,ma1)
                        img_desc24.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr24.image=imagxxadfeae                
                            tr24.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f24)
                            aA.image=p
                            canvas.create_window((1038,1450),window=aA)
                            p.image = p
                    con.close()
                    

                def phot24_24():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f24)
                    aA.image=p
                    canvas.create_window((1038,1450),window=aA)
                    p.image = p
                
                def phot25():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[82])
                        am1=str(row[133])
                        q25.insert(END,ma1)
                        img_desc25.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr25.image=imagxxadfeae                
                            tr25.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f25)
                            aA.image=p
                            canvas.create_window((1338,1450),window=aA)
                            p.image = p
                    con.close()
                    

                def phot25_25():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f25)
                    aA.image=p
                    canvas.create_window((1338,1450),window=aA)
                    p.image = p
                
                def phot26():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[83])
                        am1=str(row[134])
                        q26.insert(END,ma1)
                        img_desc26.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr26.image=imagxxadfeae                
                            tr26.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f26)
                            aA.image=p
                            canvas.create_window((138,1780),window=aA)
                            p.image = p
                    con.close()
                    

                def phot26_26():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f26)
                    aA.image=p
                    canvas.create_window((138,1780),window=aA)
                    p.image = p

                
                def phot27():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[84])
                        am1=str(row[135])
                        q27.insert(END,ma1)
                        img_desc27.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr27.image=imagxxadfeae                
                            tr27.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f27)
                            aA.image=p
                            canvas.create_window((438,1780),window=aA)
                            p.image = p
                    con.close()
                    

                def phot27_27():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f27)
                    aA.image=p
                    canvas.create_window((438,1780),window=aA)
                    p.image = p
                
                def phot28():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[85])
                        am1=str(row[136])
                        q28.insert(END,ma1)
                        img_desc28.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr28.image=imagxxadfeae                
                            tr28.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f28)
                            aA.image=p
                            canvas.create_window((738,1780),window=aA)
                            p.image = p
                    con.close()
                    

                    
                def phot28_28():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f28)
                    aA.image=p
                    canvas.create_window((738,1780),window=aA)
                    p.image = p
                
                def phot29():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[86])
                        am1=str(row[137])
                        q29.insert(END,ma1)
                        img_desc29.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr29.image=imagxxadfeae                
                            tr29.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f29)
                            aA.image=p
                            canvas.create_window((1038,1780),window=aA)
                            p.image = p
                    con.close()
                    

                def phot29_29():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f29)
                    aA.image=p
                    canvas.create_window((1038,1780),window=aA)
                    p.image = p

                
                def phot30():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[87])
                        am1=str(row[138])
                        q30.insert(END,ma1)
                        img_desc30.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr30.image=imagxxadfeae                
                            tr30.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f30)
                            aA.image=p
                            canvas.create_window((1338,1780),window=aA)
                            p.image = p
                    con.close()
                    


                def phot30_30():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f30)
                    aA.image=p
                    canvas.create_window((1338,1780),window=aA)
                    p.image = p                
                
                def phot31():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[88])
                        am1=str(row[139])
                        q31.insert(END,ma1)
                        img_desc31.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr31.image=imagxxadfeae                
                            tr31.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f31)
                            aA.image=p
                            canvas.create_window((138,2110),window=aA)
                            p.image = p
                    con.close()
                    
                    
                def phot31_31():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f31)
                    aA.image=p
                    canvas.create_window((138,2110),window=aA)
                    p.image = p
                
                def phot32():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[89])
                        am1=str(row[140])
                        q32.insert(END,ma1)
                        img_desc32.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr32.image=imagxxadfeae                
                            tr32.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f32)
                            aA.image=p
                            canvas.create_window((438,2110),window=aA)
                            p.image = p
                    con.close()
                    

                def phot32_32():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f32)
                    aA.image=p
                    canvas.create_window((438,2110),window=aA)
                    p.image = p
                
                def phot33():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[90])
                        am1=str(row[141])
                        q33.insert(END,ma1)
                        img_desc33.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr33.image=imagxxadfeae                
                            tr33.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f33)
                            aA.image=p
                            canvas.create_window((738,2110),window=aA)
                            p.image = p
                    con.close()
                    

                def phot33_33():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f33)
                    aA.image=p
                    canvas.create_window((738,2110),window=aA)
                    p.image = p

                
                def phot34():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[91])
                        am1=str(row[142])
                        q34.insert(END,ma1)
                        img_desc34.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr34.image=imagxxadfeae                
                            tr34.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f34)
                            aA.image=p
                            canvas.create_window((1038,2110),window=aA)
                            p.image = p
                    con.close()
                    

                def phot34_34():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f34)
                    aA.image=p
                    canvas.create_window((1038,2110),window=aA)
                    p.image = p
                
                def phot35():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[92])
                        am1=str(row[143])
                        q35.insert(END,ma1)
                        img_desc35.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr35.image=imagxxadfeae                
                            tr35.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f35)
                            aA.image=p
                            canvas.create_window((1338,2110),window=aA)
                            p.image = p
                    con.close()
                    

                def phot35_35():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f35)
                    aA.image=p
                    canvas.create_window((1338,2110),window=aA)
                    p.image = p
                
                def phot36():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[93])
                        am1=str(row[144])
                        q36.insert(END,ma1)
                        img_desc36.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr36.image=imagxxadfeae                
                            tr36.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f36)
                            aA.image=p
                            canvas.create_window((138,2440),window=aA)
                            p.image = p
                    con.close()
                    

                def phot36_36():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f36)
                    aA.image=p
                    canvas.create_window((138,2440),window=aA)
                    p.image = p

                
                def phot37():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[94])
                        am1=str(row[145])
                        q37.insert(END,ma1)
                        img_desc37.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr37.image=imagxxadfeae                
                            tr37.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f37)
                            aA.image=p
                            canvas.create_window((438,2440),window=aA)
                            p.image = p
                    con.close()
                    

                def phot37_37():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f37)
                    aA.image=p
                    canvas.create_window((438,2440),window=aA)
                    p.image = p
                
                def phot38():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[95])
                        am1=str(row[146])
                        q38.insert(END,ma1)
                        img_desc38.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr38.image=imagxxadfeae                
                            tr38.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f38)
                            aA.image=p
                            canvas.create_window((738,2440),window=aA)
                            p.image = p
                    con.close()
                    

                    
                def phot38_38():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f38)
                    aA.image=p
                    canvas.create_window((738,2440),window=aA)
                    p.image = p
                
                def phot39():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[96])
                        am1=str(row[147])
                        q39.insert(END,ma1)
                        img_desc39.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr39.image=imagxxadfeae                
                            tr39.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f39)
                            aA.image=p
                            canvas.create_window((1038,2440),window=aA)
                            p.image = p
                    con.close()
                    

                def phot39_39():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f39)
                    aA.image=p
                    canvas.create_window((1038,2440),window=aA)
                    p.image = p
                
                def phot40():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[97])
                        am1=str(row[148])
                        q40.insert(END,ma1)
                        img_desc40.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr40.image=imagxxadfeae                
                            tr40.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f40)
                            aA.image=p
                            canvas.create_window((1338,2440),window=aA)
                            p.image = p
                    con.close()
                    

                def phot40_40():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f40)
                    aA.image=p
                    canvas.create_window((1338,2440),window=aA)
                    p.image = p
                
                def phot41():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[98])
                        am1=str(row[149])
                        q41.insert(END,ma1)
                        img_desc41.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr41.image=imagxxadfeae                
                            tr41.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f41)
                            aA.image=p
                            canvas.create_window((138,2770),window=aA)
                            p.image = p
                    con.close()
                    

                def phot41_41():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f41)
                    aA.image=p
                    canvas.create_window((138,2770),window=aA)
                    p.image = p
                
                def phot42():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[99])
                        am1=str(row[150])
                        q42.insert(END,ma1)
                        img_desc42.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr42.image=imagxxadfeae                
                            tr42.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f42)
                            aA.image=p
                            canvas.create_window((438,2770),window=aA)
                            p.image = p
                    con.close()
                    

                def phot42_42():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f42)
                    aA.image=p
                    canvas.create_window((438,2770),window=aA)
                    p.image = p
                
                def phot43():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[100])
                        am1=str(row[151])
                        q43.insert(END,ma1)
                        img_desc43.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr43.image=imagxxadfeae                
                            tr43.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f43)
                            aA.image=p
                            canvas.create_window((738,2770),window=aA)
                            p.image = p
                    con.close()
                    

                def phot43_43():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f43)
                    aA.image=p
                    canvas.create_window((738,2770),window=aA)
                    p.image = p

                
                def phot44():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[101])
                        am1=str(row[152])
                        q44.insert(END,ma1)
                        img_desc44.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr44.image=imagxxadfeae                
                            tr44.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f44)
                            aA.image=p
                            canvas.create_window((1038,2770),window=aA)
                            p.image = p
                    con.close()
                    

                def phot44_44():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f44)
                    aA.image=p
                    canvas.create_window((1038,2770),window=aA)
                    p.image = p
                
                def phot45():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[102])
                        am1=str(row[153])
                        q45.insert(END,ma1)
                        img_desc45.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr45.image=imagxxadfeae                
                            tr45.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f45)
                            aA.image=p
                            canvas.create_window((1338,2770),window=aA)
                            p.image = p
                    con.close()
                    

                def phot45_45():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f45)
                    aA.image=p
                    canvas.create_window((1338,2770),window=aA)
                    p.image = p
                
                def phot46():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[103])
                        am1=str(row[154])
                        q46.insert(END,ma1)
                        img_desc46.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr46.image=imagxxadfeae                
                            tr46.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f46)
                            aA.image=p
                            canvas.create_window((138,3100),window=aA)
                            p.image = p
                    con.close()
                    

                def phot46_46():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f46)
                    aA.image=p
                    canvas.create_window((138,3100),window=aA)
                    p.image = p

                
                def phot47():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[104])
                        am1=str(row[155])
                        q47.insert(END,ma1)
                        img_desc47.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr47.image=imagxxadfeae                
                            tr47.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f47)
                            aA.image=p
                            canvas.create_window((438,3100),window=aA)
                            p.image = p
                    con.close()
                    

                def phot47_47():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f47)
                    aA.image=p
                    canvas.create_window((438,3100),window=aA)
                    p.image = p

                 
                def phot48():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[105])
                        am1=str(row[156])
                        q48.insert(END,ma1)
                        img_desc48.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr48.image=imagxxadfeae                
                            tr48.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f48)
                            aA.image=p
                            canvas.create_window((738,3100),window=aA)
                            p.image = p
                    con.close()
                    
                   
                def phot48_48():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f48)
                    aA.image=p
                    canvas.create_window((738,3100),window=aA)
                    p.image = p
                
                def phot49():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[106])
                        am1=str(row[157])
                        q49.insert(END,ma1)
                        img_desc49.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr49.image=imagxxadfeae                
                            tr49.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f49)
                            aA.image=p
                            canvas.create_window((1038,3100),window=aA)
                            p.image = p
                    con.close()
                    

                def phot49_49():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f49)
                    aA.image=p
                    canvas.create_window((1038,3100),window=aA)
                    p.image = p
                
                def phot50():

                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("select* from c_h1 where SERIAL_NUMBER='"+ sss.get()+"'")

                    for row in cursor:
                        ma1=str(row[107])
                        am1=str(row[158])
                        q50.insert(END,ma1)
                        img_desc50.insert(END,am1)
                    

                        if ma1!='' and ma1!='\n'  and ma1!='  '  and ma1!='null' and ma1!='none' and ma1!='None' and ma1!='NULL' and ma1!='Null' and ma1!='NONE':

                            imagxxadfeae = Image.open(ma1)
                            imagxxadfeae = imagxxadfeae.resize((196,176))
                            imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                            tr50.image=imagxxadfeae                
                            tr50.image_create(END, image =imagxxadfeae)
                            imagxxadfeae.image=imagxxadfeae

                        else:
                            
                            p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                            p = p.subsample(1, 1) 
                            aA=Button(canvas, text = '', image = p,bg='white',command=f50)
                            aA.image=p
                            canvas.create_window((1338,3100),window=aA)
                            p.image = p
                    con.close()
                    

                def phot50_50():

                    p = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/17044-200.png")
                    p = p.subsample(1, 1) 
                    aA=Button(canvas, text = '', image = p,bg='white',command=f50)
                    aA.image=p
                    canvas.create_window((1338,3100),window=aA)
                    p.image = p











                def update():

                    SER=sss.get() 
                    NAMe=e_NAME.get()
                    PAAT=e_PATIENT_ID.get()
                    DAT=e_DATE.get();
                    AGE=e_AGE.get();
                    SEX=e_SEX.get();
                    OCC=e_OCCUPATION.get('1.0', END);
                    ADD=e_ADDRESS.get('1.0', END);
                    CON=e_CONTACT.get('1.0', END);
                    CAS=e_CASETYPE.get('1.0', END);
                    TRE=e_TREATMENTTYPE.get('1.0', END);

                    listt=['A','B','C','D','E','F','G','H','I','J','K','L','M','N','O','P','Q','R','S','T','U','V','W','X','Y','Z','a','b','c','d','e','f','g','h','i','j','k','l','m','n','o','p','s','q','r','t','u','v','w','x','y','z','"','*','-','#','!','$','%','^','&','(',')','+','|','<','>',':','?','/','1','2','3','4','5','6','7','8','9','0']
                                      
                                        
                    if NAMe!='' and NAMe!=' ':
                        
                        con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                        cursor=con.cursor()
                        cursor.execute("update c_h1 set name='"+NAMe+"'where SERIAL_NUMBER='"+SER+"'")
                        cursor.execute("commit");
                        con.close()

                     
                    if PAAT!='' and PAAT!=' ':
                        con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                        cursor=con.cursor()
                        cursor.execute("update c_h1 set patient_id='"+PAAT+"'where SERIAL_NUMBER='"+SER+"'")
                        cursor.execute("commit");
                        con.close();



                    if DAT!='' and DAT!=' ':

                        if DAT[4]!="/":
                            MessageBox.showerror("Insert Status","INVALID DATE FORMAT")
                        elif DAT[7]!="/":
                            MessageBox.showerror("Insert Status","INVALID DATE FORMAT")
                        elif DAT[5]>"1":
                            MessageBox.showerror("Insert Status","INVALID DATE FORMAT")
                        elif DAT[8]>"3":
                            MessageBox.showerror("Insert Status","INVALID DATE FORMAT")
                        else:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set date='"+DAT+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();
                            

                    if AGE!='' and AGE!=' ':
                        con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                        cursor=con.cursor()
                        cursor.execute("update c_h1 set age='"+AGE+"'where SERIAL_NUMBER='"+SER+"'")
                        cursor.execute("commit");
                        con.close();
                        

                    if SEX!='' and SEX!=' ':
                        con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                        cursor=con.cursor()
                        cursor.execute("update c_h1 set sex='"+SEX+"'where SERIAL_NUMBER='"+SER+"'")
                        cursor.execute("commit");
                        con.close();

                        
                    for i in listt:
                        if i in OCC:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set occupation='"+OCC+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                        

                    for i in listt:
                        if i in ADD:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set address='"+ADD+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();


                    for i in listt:
                        if i in CON:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set contact_detail='"+CON+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();


                    for i in listt:
                        if i in CAS:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set Case_Type='"+CAS+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();


                    TREA=e_TREATMENTTYPE.get('1.0', END);

                    for i in listt:
                        if i in TREA:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set Treatment_Type='"+TREA+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();
                        

                   
                        
                    pat=e_PATIENTTYPE.get('1.0', END);
                        
                    for i in listt:
                        if i in pat:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set patient_type='"+pat+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();


                    TJU=variable.get();
                        
                    for i in listt:
                        if i in TJU:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set patient_status='"+TJU+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();
                

                        
                    CC=e_CHIEFCOMPLAINT.get('1.0', END)
      
                    for i in listt:
                        if i in CC:
                            
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set CHIEF_COMPLAINT='"+CC+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();



                 
                        
                    HOPI=e_HISTORYOFPRESENTILLNESS.get('1.0', END)
                        
                    for i in listt:
                        if i in HOPI:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set HISTORY_OF_PRESENT_ILLNESS='"+HOPI+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();



                    
                        
                    MH=e_MEDICALHISTORY.get('1.0', END);
                        
                    for i in listt:
                        if i in MH:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set MEDICAL_HISTORY='"+MH+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();



                        
                    ALL=e_ALLERGIES.get('1.0', END);
                        
                    for i in listt:
                        if i in ALL:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set ALLERGIES='"+ALL+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                             
                            
                            

                        
                    DH=e_DENTALHISTORY.get('1.0', END);
                        
                    for i in listt:
                        if i in DH:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set DENTAL_HISTORY='"+DH+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();




                        
                    GE=e_GENERALEXAMINATION.get('1.0', END);
                    
                    for i in listt:
                        if i in GE:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set GENERAL_EXAMINATION='"+GE+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();



                        
                    BP=e_BLOODPRESSURE.get('1.0', END);
                    
                    for i in listt:
                        if i in BP:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set BLOOD_PRESSURE='"+BP+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();




                    PR=e_PULSERATE.get('1.0', END);
                        
                    for i in listt:
                        if i in PR:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set PULSE_RATE='"+PR+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();



                  
                        
                    RR=e_RESPIRATORYRATE.get('1.0', END);
                        
                    for i in listt:
                        if i in RR:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set RESPIRATORY_RATE='"+RR+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();



                 
                        
                    HAB=e_HABITS.get('1.0', END);
                    
                    for i in listt:
                        if i in HAB:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set HABITS='"+HAB+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();
                        



                   
                        
                    EOE=e_EXTRAORALEXAMINATION.get('1.0', END);
                        
                    for i in listt:
                        if i in EOE:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set EXTRA_ORAL_EXAMINATION='"+EOE+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();



                  
                        
                    TJ=e_TEMPOROMANDIBULARJOINT.get('1.0', END);
                        
                    for i in listt:
                        if i in TJ:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set TEMPOROMANDIBULAR_JOINT='"+TJ+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    


                    EOL=e_EXTRAORALLYMPHNODES.get('1.0', END);
                        
                    for i in listt:
                        if i in EOL:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set EXTRA_ORAL_LYMPHNODES='"+EOL+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    


                    SF=e_SYMMETRYOFFACE.get('1.0', END);
                        
                    for i in listt:
                        if i in SF:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set SYMMETRY_OF_FACE='"+SF+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    



                    IOA=e_INTRAORALEXAMINATION.get('1.0', END);
                        
                    for i in listt:
                        if i in IOA:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set INTRA_ORAL_EXAMINATION='"+IOA+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    


                    GG=e_GINGIVA.get('1.0', END);
                        
                    for i in listt:
                        if i in GG:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set GINGIVA='"+GG+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    



                    CR=e_COLOR.get('1.0', END);
                        
                    for i in listt:
                        if i in CR:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set COLOR='"+CR+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    



                    CT=e_CONTOUR.get('1.0', END);
                        
                    for i in listt:
                        if i in CT:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set CONTOUR='"+CT+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    


                    COT=e_CONSISTENCY.get('1.0', END);
                        
                    for i in listt:
                        if i in COT:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set CONSISTENCY='"+COT+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    


                    SXT=e_SURFACETEXTURE.get('1.0', END);
                        
                    for i in listt:
                        if i in SXT:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set SURFACE_TEXTURE='"+SXT+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    



                    BOP=e_BLEEDINGONPROBING.get('1.0', END);
                        
                    for i in listt:
                        if i in BOP:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set BLEEDING_ON_PROBING='"+BOP+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    



                    PE=e_PERIODONTALEXAMINATION.get('1.0', END);
                        
                    for i in listt:
                        if i in PE:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set PERIODONTAL_EXAMINATION='"+PE+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    



                    PP=e_PERIODONTALPOCKETS.get('1.0', END);
                        
                    for i in listt:
                        if i in PP:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set PERIODONTAL_POCKETS='"+PP+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    



                    STA=e_STAINS.get('1.0', END);
                        
                    for i in listt:
                        if i in STA:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set STAINS='"+STA+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    
                    

                    CL=e_CALCULUS.get('1.0', END);
                        
                    for i in listt:
                        if i in CL:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set CALCULUS='"+CL+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    


                    FE=e_FURCATIONINVOLVEMENT.get('1.0', END);
                        
                    for i in listt:
                        if i in FE:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set FURCATION_INVOLVEMENT='"+FE+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    


                    RECC=e_RECESSION.get('1.0', END);
                        
                    for i in listt:
                        if i in RECC:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set RECESSION='"+RECC+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    



                    OF=e_OTHERFINDINGS.get('1.0', END);
                        
                    for i in listt:
                        if i in OF:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set OTHER_FINDINGS='"+OF+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    



                    HP=e_HARDPALATE.get('1.0', END);
                        
                    for i in listt:
                        if i in HP:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set HARD_PALATE='"+HP+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    



                    sp=e_SOFTPALATE.get('1.0', END);
                        
                    for i in listt:
                        if i in sp:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set SOFT_PALATE='"+sp+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    




                    BM=e_BUCCALMUCOSA.get('1.0', END);
                        
                    for i in listt:
                        if i in BM:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set BUCCAL_MUCOSA='"+BM+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    



                    TG=e_TONGUE.get('1.0', END);
                        
                    for i in listt:
                        if i in TG:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set TONGUE='"+TG+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    



                    FM=e_FLOOROFTHEMOUTH.get('1.0', END);
                        
                    for i in listt:
                        if i in FM:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set FLOOR_OF_THE_MOUTH='"+FM+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    


                    TP=e_TEETHEXAMINATION.get('1.0', END);
                        
                    for i in listt:
                        if i in TP:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set TEETH_EXAMINATION='"+TP+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    


                    NT=e_NATUREOFTEETH.get('1.0', END);
                        
                    for i in listt:
                        if i in NT:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set NATURE_OF_TEETH='"+NT+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    


                    NOTP=e_NUMBEROFTEETHPRESENT.get('1.0', END);
                        
                    for i in listt:
                        if i in NOTP:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set NUMBER_OF_TEETH_PRESENT='"+NOTP+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    



                    DAAT=e_DECAYEDTEETH.get('1.0', END);
                        
                    for i in listt:
                        if i in DAAT:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set DECAYED_TEETH='"+DAAT+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    


                    TAJ=e_FILLEDTEETH.get('1.0', END);
                        
                    for i in listt:
                        if i in TAJ:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set FILLED_TEETH='"+TAJ+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    


                    ATJ=e_TENDERNESSONPERCUSSION.get('1.0', END);
                        
                    for i in listt:
                        if i in ATJ:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set TENDERNESS_ON_PERCUSSION='"+ATJ+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    


                    TJZ=e_ROOTSTUMPS.get('1.0', END);
                        
                    for i in listt:
                        if i in TJZ:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set ROOT_STUMPS='"+TJZ+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    


                    QTJ=e_MISSING.get('1.0', END);
                        
                    for i in listt:
                        if i in QTJ:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set MISSING='"+QTJ+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    


                    TJC=e_PROVISIONALDIAGNOSIS.get('1.0', END);
                        
                    for i in listt:
                        if i in TJC:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set PROVISIONAL_DIAGNOSIS='"+TJC+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    


                    TJN=e_DIAGNOSTICAIDS.get('1.0', END);
                        
                    for i in listt:
                        if i in TJN:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set DIAGNOSTIC_AIDS='"+TJN+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    


                    TJP=e_FINALDIAGNOSIS.get('1.0', END);
                        
                    for i in listt:
                        if i in TJP:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set FINAL_DIAGNOSIS='"+TJP+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    



                    TJO=e_TREATMENTPLAN.get('1.0', END);
                        
                    for i in listt:
                        if i in TJO:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set TREATMENT_PLAN='"+TJO+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                    



                    TJU=e_ADDITIONALINFO.get('1.0', END);
                        
                    for i in listt:
                        if i in TJU:
                            con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                            cursor=con.cursor()
                            cursor.execute("update c_h1 set ADDITIONAL_INFO='"+TJU+"'where SERIAL_NUMBER='"+SER+"'")
                            cursor.execute("commit");
                            con.close();

                  
                        
                            
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image='"+q1.get(1.0,'1.end')+"',desc1='"+img_desc1.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image1='"+q2.get(1.0,'1.end')+"',desc2='"+img_desc2.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image2='"+q3.get(1.0,'1.end')+"',desc3='"+img_desc3.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image3='"+q4.get(1.0,'1.end')+"',desc4='"+img_desc4.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image4='"+q5.get(1.0,'1.end')+"',desc5='"+img_desc5.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                        
                            
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image5='"+q6.get(1.0,'1.end')+"',desc6='"+img_desc6.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image6='"+q7.get(1.0,'1.end')+"',desc7='"+img_desc7.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image7='"+q8.get(1.0,'1.end')+"',desc8='"+img_desc8.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image8='"+q9.get(1.0,'1.end')+"',desc9='"+img_desc9.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image9='"+q10.get(1.0,'1.end')+"',desc10='"+img_desc10.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                        
                            
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image10='"+q11.get(1.0,'1.end')+"',desc11='"+img_desc11.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image11='"+q12.get(1.0,'1.end')+"',desc12='"+img_desc12.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image12='"+q13.get(1.0,'1.end')+"',desc13='"+img_desc13.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image13='"+q14.get(1.0,'1.end')+"',desc14='"+img_desc14.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image14='"+q15.get(1.0,'1.end')+"',desc15='"+img_desc15.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                        
                            
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image15='"+q16.get(1.0,'1.end')+"',desc16='"+img_desc16.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image16='"+q17.get(1.0,'1.end')+"',desc17='"+img_desc17.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image17='"+q18.get(1.0,'1.end')+"',desc18='"+img_desc18.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image18='"+q19.get(1.0,'1.end')+"',desc19='"+img_desc19.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image19='"+q20.get(1.0,'1.end')+"',desc20='"+img_desc20.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                        
                            
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image20='"+q21.get(1.0,'1.end')+"',desc21='"+img_desc21.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image21='"+q22.get(1.0,'1.end')+"',desc22='"+img_desc22.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image22='"+q23.get(1.0,'1.end')+"',desc23='"+img_desc23.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image23='"+q24.get(1.0,'1.end')+"',desc24='"+img_desc24.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image24='"+q25.get(1.0,'1.end')+"',desc25='"+img_desc25.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                        
                            
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image25='"+q26.get(1.0,'1.end')+"',desc26='"+img_desc26.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image26='"+q27.get(1.0,'1.end')+"',desc27='"+img_desc27.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image27='"+q28.get(1.0,'1.end')+"',desc28='"+img_desc28.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image28='"+q29.get(1.0,'1.end')+"',desc29='"+img_desc29.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image29='"+q30.get(1.0,'1.end')+"',desc30='"+img_desc30.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                        
                            
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image30='"+q31.get(1.0,'1.end')+"',desc31='"+img_desc31.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image31='"+q32.get(1.0,'1.end')+"',desc32='"+img_desc32.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image32='"+q33.get(1.0,'1.end')+"',desc33='"+img_desc33.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image33='"+q34.get(1.0,'1.end')+"',desc34='"+img_desc34.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image34='"+q35.get(1.0,'1.end')+"',desc35='"+img_desc35.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                        
                            
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image35='"+q36.get(1.0,'1.end')+"',desc36='"+img_desc36.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image36='"+q37.get(1.0,'1.end')+"',desc37='"+img_desc37.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image37='"+q38.get(1.0,'1.end')+"',desc38='"+img_desc38.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image38='"+q39.get(1.0,'1.end')+"',desc39='"+img_desc39.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image39='"+q40.get(1.0,'1.end')+"',desc40='"+img_desc40.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                        
                            
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image40='"+q41.get(1.0,'1.end')+"',desc41='"+img_desc41.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image41='"+q42.get(1.0,'1.end')+"',desc42='"+img_desc42.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image42='"+q43.get(1.0,'1.end')+"',desc43='"+img_desc43.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image43='"+q44.get(1.0,'1.end')+"',desc44='"+img_desc44.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image44='"+q45.get(1.0,'1.end')+"',desc45='"+img_desc45.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                        
                            
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image45='"+q46.get(1.0,'1.end')+"',desc46='"+img_desc46.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image46='"+q47.get(1.0,'1.end')+"',desc47='"+img_desc47.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image47='"+q48.get(1.0,'1.end')+"',desc48='"+img_desc48.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image48='"+q49.get(1.0,'1.end')+"',desc49='"+img_desc49.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();

                  
                    con = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                    cursor=con.cursor()
                    cursor.execute("update c_h1 set image49='"+q50.get(1.0,'1.end')+"',desc50='"+img_desc50.get(1.0,END)+"'where SERIAL_NUMBER='"+SER+"'")
                    cursor.execute("commit");
                    con.close();


                    MessageBox.showinfo('UPDATE STATUS','Updated Successfully!')
                    top.destroy()
                    sss.delete(0,END)



             
                 
                    

                
                canvas_PERSONAL_DETAILS = Canvas(PERSONAL_DETAILS, width = 800, height = 400)
                #canvas_PERSONAL_DETAILS.configure(bg='lightsteelblue')
                canvas_PERSONAL_DETAILS.pack(side = LEFT, fill = BOTH, expand = True)



                
                entryContainerZr = Frame(canvas_PERSONAL_DETAILS)                # FRAME FOR CANVAS
                entryContainerZr.pack(fill = BOTH)

                gif1sX = Image.open("C:/Users/Acer/Desktop/pms files/ImageBlur (5).png")    # PERSONAL DETAILS BG IMAGE
                gif1sX = gif1sX.resize((1550, 765))
                imageD = ImageTk.PhotoImage(gif1sX)
                labels2 = tk.Label(canvas_PERSONAL_DETAILS,image=imageD)
                labels2.image=imageD
                canvas_PERSONAL_DETAILS.create_window((0, 0), anchor=NW, window=labels2)
                imageD.image = imageD        

                                                                            # TOP LABEL WITH updatePATIENT TEXT       
                photoiQmage = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/update.png") 
                aj=Button(canvas_PERSONAL_DETAILS, text = '', image = photoiQmage)
                aj.image=photoiQmage
                canvas_PERSONAL_DETAILS.create_window((8, 15), anchor=NW, window=aj)
                photoiQmage.image = photoiQmage

                





        # DATA  INSERT  BUTTONS  AND  ENTRY  BOXES
                
 
                

                widget = Button(canvas_PERSONAL_DETAILS, text='      NAME       ',font=('arial ',16), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_PERSONAL_DETAILS.create_window(270, 245, window=widget)    
                e_NAME=Entry(canvas_PERSONAL_DETAILS,font=('verdana ',13),bg='snow')
                canvas_PERSONAL_DETAILS.create_window((520,245),window=e_NAME)
                

                
                widget = Button(canvas_PERSONAL_DETAILS, text='  PATIENT ID   ',font=('arial ',16), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_PERSONAL_DETAILS.create_window(270, 335, window=widget)
                e_PATIENT_ID=Entry(canvas_PERSONAL_DETAILS,font=('verdana',11),bg='snow')
                canvas_PERSONAL_DETAILS.create_window((520,335),window=e_PATIENT_ID)
                    

                
                widget = Button(canvas_PERSONAL_DETAILS, text='       DATE        ',font=('arial',16), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_PERSONAL_DETAILS.create_window(270, 420, window=widget)
                e_DATE=Entry(canvas_PERSONAL_DETAILS,font=('verdana',11),bg='snow')
                canvas_PERSONAL_DETAILS.create_window((520,420),window=e_DATE)

                

                
                widget = Button(canvas_PERSONAL_DETAILS, text='        AGE         ',font=('arial',16), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_PERSONAL_DETAILS.create_window(270, 500, window=widget)
                e_AGE=Entry(canvas_PERSONAL_DETAILS,font=('verdana',11),bg='snow')
                canvas_PERSONAL_DETAILS.create_window((520,500),window=e_AGE)

                coni = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursoir=coni.cursor()
                cursoir.execute("select* from c_h1 where SERIAL_NUMBER='"+sss.get()+"'")
                for row in cursoir:
                    koap=str(row[5])

                    widget = Button(canvas_PERSONAL_DETAILS, text='        SEX         ',font=('arial',16), fg='white',bg='darkturquoise')
                    widget.pack()
                    canvas_PERSONAL_DETAILS.create_window(270, 580, window=widget)
                    e_SEX = StringVar(canvas_PERSONAL_DETAILS)
                    e_SEX.set(koap)
                    waas = OptionMenu(canvas_PERSONAL_DETAILS, e_SEX, "Male", "Female")
                    waas.config(font='arial')
                    canvas_PERSONAL_DETAILS.create_window((460,580),window=waas)
                

               
                widget = Button(canvas_PERSONAL_DETAILS, text='  OCCUPATION  ',font=('arial',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_PERSONAL_DETAILS.create_window(960, 250, window=widget)
                e_OCCUPATION=Text(canvas_PERSONAL_DETAILS, height=3, width=40,font=('VERDANA',11),bg='snow')
                canvas_PERSONAL_DETAILS.create_window((1280,250),window=e_OCCUPATION)
                
                
                widget = Button(canvas_PERSONAL_DETAILS, text='     ADDRESS     ',font=('arial',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_PERSONAL_DETAILS.create_window(960,350, window=widget)
                e_ADDRESS=Text(canvas_PERSONAL_DETAILS, height=3, width=40,font=('verdana',11),bg='snow')
                canvas_PERSONAL_DETAILS.create_window((1280,350),window=e_ADDRESS)
                

                
                widget = Button(canvas_PERSONAL_DETAILS, text='     CONTACT     ',font=('arial',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_PERSONAL_DETAILS.create_window(960, 450, window=widget)
                e_CONTACT=Text(canvas_PERSONAL_DETAILS, height=3, width=40,font=('VERDANA',11),bg='snow')
                canvas_PERSONAL_DETAILS.create_window((1280,450),window=e_CONTACT)


                
         # ADD  RECORD  BUTTON         
                
                button = Button(canvas_PERSONAL_DETAILS, text=" UPDATE   ", font=("ROBOTO BOLD",18),bg='dodgerblue',fg='snow',command=update)
                canvas_PERSONAL_DETAILS.create_window((1360, 130), anchor=NW, window=button)



         #####################        MEDICAL   DETAILS       ####################





                canvas_MEDICAL_DETAILS = Canvas(MEDICAL_DETAILS, width = 800, height = 400,scrollregion=(0,0,4650,6500))
                canvas_MEDICAL_DETAILS.pack(side = LEFT, fill = BOTH, expand = True)
                #canvas_MEDICAL_DETAILS.config(bg='gainsboro')

                vsb = Scrollbar(canvas_MEDICAL_DETAILS, command=canvas_MEDICAL_DETAILS.yview)                        # SCROLLBAR FOR CANVAS
                canvas_MEDICAL_DETAILS.config(yscrollcommand = vsb.set)
                canvas_MEDICAL_DETAILS.pack(side = LEFT, fill = BOTH, expand = True)
                vsb.pack(side = RIGHT, fill = Y)


                
                entryContainerAZr = Frame(canvas_MEDICAL_DETAILS)                # FRAME FOR CANVAS
                entryContainerAZr.pack(fill = BOTH)





        # DATA  INSERT  BUTTONS  AND  ENTRY  BOXES

                widget = Button(canvas_MEDICAL_DETAILS, text='CHIEF COMPLAINT',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(174, 100, window=widget)
                e_CHIEFCOMPLAINT=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,100),window=e_CHIEFCOMPLAINT)

            

                
                widget = Button(canvas_MEDICAL_DETAILS, text='HISTORY OF PRESENT ILLNESS',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(174, 240, window=widget)
                e_HISTORYOFPRESENTILLNESS=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,240),window=e_HISTORYOFPRESENTILLNESS) 
               

               
                widget = Button(canvas_MEDICAL_DETAILS, text='MEDICAL HISTORY',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170,380, window=widget)
                e_MEDICALHISTORY=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,380),window=e_MEDICALHISTORY)
               
               
                widget = Button(canvas_MEDICAL_DETAILS, text='ALLERGIES',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 520, window=widget)
                e_ALLERGIES=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,520),window=e_ALLERGIES)
                

                
                widget = Button(canvas_MEDICAL_DETAILS, text='DENTAL HISTORY',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 660, window=widget)
                e_DENTALHISTORY=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,660),window=e_DENTALHISTORY)
                

                
                widget = Button(canvas_MEDICAL_DETAILS, text='GENERAL EXAMINATION',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 800, window=widget)
                e_GENERALEXAMINATION=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,800),window=e_GENERALEXAMINATION)
                

                
                widget = Button(canvas_MEDICAL_DETAILS, text='BLOOD PRESSURE',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 940, window=widget)
                e_BLOODPRESSURE=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,940),window=e_BLOODPRESSURE)
                

               
                widget = Button(canvas_MEDICAL_DETAILS, text='PULSE RATE',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 1080, window=widget)
                e_PULSERATE=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,1080),window=e_PULSERATE)
                

               
                widget = Button(canvas_MEDICAL_DETAILS, text='RESPIRATORY RATE',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 1220, window=widget)
                e_RESPIRATORYRATE=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,1220),window=e_RESPIRATORYRATE)
                

               
                widget = Button(canvas_MEDICAL_DETAILS, text='HABITS',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 1360, window=widget)
                e_HABITS=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,1360),window=e_HABITS)
               

                
                widget = Button(canvas_MEDICAL_DETAILS, text='EXTRA ORAL EXAMINATION',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170,1500, window=widget)
                e_EXTRAORALEXAMINATION=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,1500),window=e_EXTRAORALEXAMINATION)
                

               
                widget = Button(canvas_MEDICAL_DETAILS, text='TEMPOROMANDIBULAR JOINT',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 1640, window=widget)
                e_TEMPOROMANDIBULARJOINT=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,1640),window=e_TEMPOROMANDIBULARJOINT)
                

                
                widget = Button(canvas_MEDICAL_DETAILS, text='EXTRA ORAL LYMPHNODES',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 1780, window=widget)
                e_EXTRAORALLYMPHNODES=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,1780),window=e_EXTRAORALLYMPHNODES)
                

                
                widget = Button(canvas_MEDICAL_DETAILS, text='SYMMETRY OF FACE',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 1920, window=widget)
                e_SYMMETRYOFFACE=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,1920),window=e_SYMMETRYOFFACE)
               

                
                widget = Button(canvas_MEDICAL_DETAILS, text='INTRA ORAL EXAMINATION',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 2060, window=widget)
                e_INTRAORALEXAMINATION=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,2060),window=e_INTRAORALEXAMINATION)
                

               
                widget = Button(canvas_MEDICAL_DETAILS, text='GINGIVA',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 2200, window=widget)
                e_GINGIVA=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,2200),window=e_GINGIVA)
                

                widget = Button(canvas_MEDICAL_DETAILS, text='COLOR',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 2340, window=widget)
                e_COLOR=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,2340),window=e_COLOR)
                
              

                widget = Button(canvas_MEDICAL_DETAILS, text='CONTOUR',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 2480, window=widget)
                e_CONTOUR=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,2480),window=e_CONTOUR)
              


                widget = Button(canvas_MEDICAL_DETAILS, text='CONSISTENCY',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 2620, window=widget)
                e_CONSISTENCY=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,2620),window=e_CONSISTENCY)
                
                

                widget = Button(canvas_MEDICAL_DETAILS, text='SURFACE TEXTURE',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 2760, window=widget)
                e_SURFACETEXTURE=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,2760),window=e_SURFACETEXTURE)
               


                widget = Button(canvas_MEDICAL_DETAILS, text='BLEEDING ON PROBING',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 2900, window=widget)
                e_BLEEDINGONPROBING=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,2900),window=e_BLEEDINGONPROBING)
                


                widget = Button(canvas_MEDICAL_DETAILS, text='PERIODONTAL EXAMINATION',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 3040, window=widget)
                e_PERIODONTALEXAMINATION=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,3040),window=e_PERIODONTALEXAMINATION)
              


                widget = Button(canvas_MEDICAL_DETAILS, text='PERIODONTAL POCKETS',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 3180, window=widget)
                e_PERIODONTALPOCKETS=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,3180),window=e_PERIODONTALPOCKETS)



                widget = Button(canvas_MEDICAL_DETAILS, text='STAINS',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 3320, window=widget)
                e_STAINS=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,3320),window=e_STAINS)



                widget = Button(canvas_MEDICAL_DETAILS, text='CALCULUS',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 3460, window=widget)
                e_CALCULUS=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,3460),window=e_CALCULUS)



                widget = Button(canvas_MEDICAL_DETAILS, text='FURCATION INVOLVEMENT',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 3600, window=widget)
                e_FURCATIONINVOLVEMENT=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,3600),window=e_FURCATIONINVOLVEMENT)
              


                widget = Button(canvas_MEDICAL_DETAILS, text='RECESSION',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 3740, window=widget)
                e_RECESSION=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,3740),window=e_RECESSION)
               

                widget = Button(canvas_MEDICAL_DETAILS, text='OTHER FINDINGS',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 3880, window=widget)
                e_OTHERFINDINGS=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,3880),window=e_OTHERFINDINGS)
               


                widget = Button(canvas_MEDICAL_DETAILS, text='HARD PALATE',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 4020, window=widget)
                e_HARDPALATE=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,4020),window=e_HARDPALATE)
              


                widget = Button(canvas_MEDICAL_DETAILS, text='SOFT PALATE',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 4160, window=widget)
                e_SOFTPALATE=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,4160),window=e_SOFTPALATE)



                widget = Button(canvas_MEDICAL_DETAILS, text='BUCCAL MUCOSA',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 4300, window=widget)
                e_BUCCALMUCOSA=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,4300),window=e_BUCCALMUCOSA)



                widget = Button(canvas_MEDICAL_DETAILS, text='TONGUE',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 4440, window=widget)
                e_TONGUE=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,4440),window=e_TONGUE)


                widget = Button(canvas_MEDICAL_DETAILS, text='FLOOR OF THE MOUTH',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 4580, window=widget)
                e_FLOOROFTHEMOUTH=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,4580),window=e_FLOOROFTHEMOUTH)



                widget = Button(canvas_MEDICAL_DETAILS, text='TEETH EXAMINATION',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 4720, window=widget)
                e_TEETHEXAMINATION=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,4720),window=e_TEETHEXAMINATION)



                widget = Button(canvas_MEDICAL_DETAILS, text='NATURE OF TEETH',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 4860, window=widget)
                e_NATUREOFTEETH=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,4860),window=e_NATUREOFTEETH)
               


                widget = Button(canvas_MEDICAL_DETAILS, text='NUMBER OF TEETH PRESENT',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 5000, window=widget)
                e_NUMBEROFTEETHPRESENT=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,5000),window=e_NUMBEROFTEETHPRESENT)
                


                widget = Button(canvas_MEDICAL_DETAILS, text='DECAYED TEETH',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 5140, window=widget)
                e_DECAYEDTEETH=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,5140),window=e_DECAYEDTEETH)
              

                widget = Button(canvas_MEDICAL_DETAILS, text='FILLED TEETH',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 5280, window=widget)
                e_FILLEDTEETH=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,5280),window=e_FILLEDTEETH)
               


                widget = Button(canvas_MEDICAL_DETAILS, text='TENDERNESS ON PERCUSSION',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 5420, window=widget)
                e_TENDERNESSONPERCUSSION=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,5420),window=e_TENDERNESSONPERCUSSION)



                widget = Button(canvas_MEDICAL_DETAILS, text='ROOT STUMPS',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 5560, window=widget)
                e_ROOTSTUMPS=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,5560),window=e_ROOTSTUMPS)



                widget = Button(canvas_MEDICAL_DETAILS, text='MISSING',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 5700, window=widget)
                e_MISSING=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,5700),window=e_MISSING)



                widget = Button(canvas_MEDICAL_DETAILS, text='PROVISIONAL DIAGNOSIS',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 5840, window=widget)
                e_PROVISIONALDIAGNOSIS=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,5840),window=e_PROVISIONALDIAGNOSIS)



                widget = Button(canvas_MEDICAL_DETAILS, text='DIAGNOSTIC AIDS',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 5980, window=widget)
                e_DIAGNOSTICAIDS=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,5980),window=e_DIAGNOSTICAIDS)



                widget = Button(canvas_MEDICAL_DETAILS, text='FINAL DIAGNOSIS',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 6120, window=widget)
                e_FINALDIAGNOSIS=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,6120),window=e_FINALDIAGNOSIS)



                widget = Button(canvas_MEDICAL_DETAILS, text='TREATMENT PLAN',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 6260, window=widget)
                e_TREATMENTPLAN=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,6260),window=e_TREATMENTPLAN)



                widget = Button(canvas_MEDICAL_DETAILS, text='ADDITIONAL INFO',font=('arial ',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_MEDICAL_DETAILS.create_window(170, 6400, window=widget)
                e_ADDITIONALINFO=Text(canvas_MEDICAL_DETAILS, height=5, width=113,font=('VERDANA',11))
                canvas_MEDICAL_DETAILS.create_window((920,6400),window=e_ADDITIONALINFO)




         #####################        OTHER   DETAILS       ####################




                canvas_OTHER_DETAILS = Canvas(OTHER_DETAILS, width = 800, height = 400)
                canvas_OTHER_DETAILS.pack(side = LEFT, fill = BOTH, expand = True)



                
                AentryContainerAZr = Frame(canvas_OTHER_DETAILS)                # FRAME FOR CANVAS
                AentryContainerAZr.pack(fill = BOTH)

                                               


        # DATA  INSERT  BUTTONS  AND  ENTRY  BOXES



                widget = Button(canvas_OTHER_DETAILS, text='    CASE TYPE    ',font=('VERDANA',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_OTHER_DETAILS.create_window(320, 80, window=widget)
                e_CASETYPE=Text(canvas_OTHER_DETAILS, height=3, width=45,font=('VERDANA',11),bg='snow')
                canvas_OTHER_DETAILS.create_window((690,80),window=e_CASETYPE)
              



                widget = Button(canvas_OTHER_DETAILS, text='TREATMENT TYPE',font=('VERDANA',14),fg='white',bg='darkturquoise')
                widget.pack()
                canvas_OTHER_DETAILS.create_window(320, 180, window=widget)
                e_TREATMENTTYPE=Text(canvas_OTHER_DETAILS, height=3, width=45,font=('VERDANA',11),bg='snow')
                canvas_OTHER_DETAILS.create_window((690,180),window=e_TREATMENTTYPE)



                widget = Button(canvas_OTHER_DETAILS, text='  PATIENT TYPE  ',font=('VERDANA',14), fg='white',bg='darkturquoise')
                widget.pack()
                canvas_OTHER_DETAILS.create_window(320, 280, window=widget)
                e_PATIENTTYPE=Text(canvas_OTHER_DETAILS, height=3, width=45,font=('VERDANA',11),bg='snow')
                canvas_OTHER_DETAILS.create_window((690,280),window=e_PATIENTTYPE)

                coni = mysql.connect(host="localhost",user="root",password="admin",database="patientsoftware1")
                cursoir=coni.cursor()
                cursoir.execute("select* from c_h1 where SERIAL_NUMBER='"+sss.get()+"'")
                for row in cursoir:
                    kop=str(row[108])
    

                    widgetb = Button(canvas_OTHER_DETAILS, text='PATIENT STATUS',font=('VERDANA',14), fg='white',bg='darkturquoise')
                    widgetb.pack()
                    canvas_OTHER_DETAILS.create_window(320, 380, window=widgetb)
                    variable = StringVar(canvas_OTHER_DETAILS)
                    variable.set(kop)
                    was = OptionMenu(canvas_OTHER_DETAILS, variable, "Active", "Recovered",'Referred')
                    was.config(font='arial')
                    canvas_OTHER_DETAILS.create_window((550,380),window=was)
                con.close()

        #############################  IMAGE  ADD   ############################
                q1=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q1)

                q2=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q2)

                q3=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q3)

                q4=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q4)

                q5=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q5)

                q6=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q6)

                q7=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q7)

                q8=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q8)

                q9=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q9)

                q10=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q10)

                q11=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q11)

                q12=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q12)

                q13=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q13)

                q14=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q14)

                q15=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q15)

                q16=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q16)

                q17=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q17)

                q18=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q18)

                q19=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q19)

                q20=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q20)

                q21=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q21)

                q22=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q22)

                q23=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q23)

                q24=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q24)

                q25=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q25)

                q26=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q26)

                q27=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q27)

                q28=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q28)

                q29=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q29)

                q30=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q30)

                q31=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q31)

                q32=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q32)

                q33=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q33)

                q34=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q34)

                q35=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q35)

                q36=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q36)

                q37=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q37)

                q38=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q38)

                q39=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q39)

                q40=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q40)

                q41=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q41)

                q42=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q42)

                q43=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q43)

                q44=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q44)

                q45=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q45)

                q46=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q46)

                q47=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q47)

                q48=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q48)

                q49=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q49)

                q50=Text(canvas_OTHER_DETAILS,height=3,width=100)
                canvas_OTHER_DETAILS.create_window((10000,20000),window=q50)

                
              

                def fa1():
                    filename = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q1.insert(END,filename)
                    
                    imagxxadfeae = Image.open(filename)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t1=Text(canvas,width=24,height=11)
                    canvas.create_window((138,130),window=t1)
                    t1.image=imagxxadfeae
                    
                    t1.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae



                def f2():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q2.insert(END,filenamee)
                    
                    imagxxadfeae1 = Image.open(filenamee)
                    imagxxadfeae1 = imagxxadfeae1.resize((196,176))
                    imagxxadfeae1 = ImageTk.PhotoImage(imagxxadfeae1)
                    
                    t2=Text(canvas,width=24,height=11)
                    canvas.create_window((438,130),window=t2)                
                    t2.image=imagxxadfeae1
                    
                    t2.image_create(END, image =imagxxadfeae1)
                    imagxxadfeae1.image=imagxxadfeae1


                def f3():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q3.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t3=Text(canvas,width=24,height=11)
                    canvas.create_window((738,130),window=t3)                
                    t3.image=imagxxadfeae
                    
                    t3.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae


                def f4():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q4.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t4=Text(canvas,width=24,height=11)
                    canvas.create_window((1038,130),window=t4)                
                    t4.image=imagxxadfeae
                    
                    t4.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae


                def f5():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q5.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t5=Text(canvas,width=24,height=11)
                    canvas.create_window((1338,130),window=t5)                
                    t5.image=imagxxadfeae
                    
                    t5.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae

                def f6():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q6.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t6=Text(canvas,width=24,height=11)
                    canvas.create_window((138,460),window=t6)                
                    t6.image=imagxxadfeae
                    
                    t6.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae                

                def f7():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q7.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t7=Text(canvas,width=24,height=11)
                    canvas.create_window((438,460),window=t7)                
                    t7.image=imagxxadfeae
                    
                    t7.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae


                def f8():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q8.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t8=Text(canvas,width=24,height=11)
                    canvas.create_window((738,460),window=t8)                
                    t8.image=imagxxadfeae
                    
                    t8.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae


                def f9():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q9.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t9=Text(canvas,width=24,height=11)
                    canvas.create_window((1038,460),window=t9)                
                    t9.image=imagxxadfeae
                    
                    t9.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae
                    
                def f10():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q10.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t10=Text(canvas,width=24,height=11)
                    canvas.create_window((1338,460),window=t10)                
                    t10.image=imagxxadfeae
                    
                    t10.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae

                    
                def fa11():
                    filename = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q11.insert(END,filename)
                    
                    imagxxadfeae = Image.open(filename)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t11=Text(canvas,width=24,height=11)
                    canvas.create_window((138,790),window=t11)
                    t11.image=imagxxadfeae
                    
                    t11.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae




                def f12():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q12.insert(END,filenamee)
                    
                    imagxxadfeae1 = Image.open(filenamee)
                    imagxxadfeae1 = imagxxadfeae1.resize((196,176))
                    imagxxadfeae1 = ImageTk.PhotoImage(imagxxadfeae1)
                    
                    t12=Text(canvas,width=24,height=11)
                    canvas.create_window((438,790),window=t12)                
                    t12.image=imagxxadfeae1
                    
                    t12.image_create(END, image =imagxxadfeae1)
                    imagxxadfeae1.image=imagxxadfeae1


                def f13():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q13.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t13=Text(canvas,width=24,height=11)
                    canvas.create_window((738,790),window=t13)                
                    t13.image=imagxxadfeae
                    
                    t13.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae


                def f14():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q14.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t14=Text(canvas,width=24,height=11)
                    canvas.create_window((1038,790),window=t14)                
                    t14.image=imagxxadfeae
                    
                    t14.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae


                def f15():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q15.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t15=Text(canvas,width=24,height=11)
                    canvas.create_window((1338,790),window=t15)                
                    t15.image=imagxxadfeae
                    
                    t15.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae

                def f16():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q16.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t16=Text(canvas,width=24,height=11)
                    canvas.create_window((138,1120),window=t16)                
                    t16.image=imagxxadfeae
                    
                    t16.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae                

                def f17():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q17.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t17=Text(canvas,width=24,height=11)
                    canvas.create_window((438,1120),window=t17)                
                    t17.image=imagxxadfeae
                    
                    t17.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae


                def f18():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q18.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t18=Text(canvas,width=24,height=11)
                    canvas.create_window((738,1120),window=t18)                
                    t18.image=imagxxadfeae
                    
                    t18.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae


                def f19():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q19.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t19=Text(canvas,width=24,height=11)
                    canvas.create_window((1038,1120),window=t19)                
                    t19.image=imagxxadfeae
                    
                    t19.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae
                    
                def f20():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q20.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t20=Text(canvas,width=24,height=11)
                    canvas.create_window((1338,1120),window=t20)                
                    t20.image=imagxxadfeae
                    
                    t20.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae

                    
                def f21():
                    filename = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q21.insert(END,filename)
                    
                    imagxxadfeae = Image.open(filename)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t21=Text(canvas,width=24,height=11)
                    canvas.create_window((138,1450),window=t21)
                    t21.image=imagxxadfeae
                    
                    t21.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae




                def f22():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q22.insert(END,filenamee)
                    
                    imagxxadfeae1 = Image.open(filenamee)
                    imagxxadfeae1 = imagxxadfeae1.resize((196,176))
                    imagxxadfeae1 = ImageTk.PhotoImage(imagxxadfeae1)
                    
                    t12=Text(canvas,width=24,height=11)
                    canvas.create_window((438,1450),window=t12)                
                    t12.image=imagxxadfeae1
                    
                    t12.image_create(END, image =imagxxadfeae1)
                    imagxxadfeae1.image=imagxxadfeae1


                def f23():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q23.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t13=Text(canvas,width=24,height=11)
                    canvas.create_window((738,1450),window=t13)                
                    t13.image=imagxxadfeae
                    
                    t13.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae


                def f24():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q24.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t14=Text(canvas,width=24,height=11)
                    canvas.create_window((1038,1450),window=t14)                
                    t14.image=imagxxadfeae
                    
                    t14.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae


                def f25():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q25.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t15=Text(canvas,width=24,height=11)
                    canvas.create_window((1338,1450),window=t15)                
                    t15.image=imagxxadfeae
                    
                    t15.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae

                def f26():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q26.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t16=Text(canvas,width=24,height=11)
                    canvas.create_window((138,1780),window=t16)                
                    t16.image=imagxxadfeae
                    
                    t16.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae                

                def f27():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q27.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t17=Text(canvas,width=24,height=11)
                    canvas.create_window((438,1780),window=t17)                
                    t17.image=imagxxadfeae
                    
                    t17.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae


                def f28():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q28.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t18=Text(canvas,width=24,height=11)
                    canvas.create_window((738,1780),window=t18)                
                    t18.image=imagxxadfeae
                    
                    t18.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae


                def f29():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q29.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t19=Text(canvas,width=24,height=11)
                    canvas.create_window((1038,1780),window=t19)                
                    t19.image=imagxxadfeae
                    
                    t19.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae
                    
                def f30():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q30.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t20=Text(canvas,width=24,height=11)
                    canvas.create_window((1338,1780),window=t20)                
                    t20.image=imagxxadfeae
                    
                    t20.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae

                    
                def f31():
                    filename = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q31.insert(END,filename)
                    
                    imagxxadfeae = Image.open(filename)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t1=Text(canvas,width=24,height=11)
                    canvas.create_window((138,2110),window=t1)
                    t1.image=imagxxadfeae
                    
                    t1.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae




                def f32():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q32.insert(END,filenamee)
                    
                    imagxxadfeae1 = Image.open(filenamee)
                    imagxxadfeae1 = imagxxadfeae1.resize((196,176))
                    imagxxadfeae1 = ImageTk.PhotoImage(imagxxadfeae1)
                    
                    t2=Text(canvas,width=24,height=11)
                    canvas.create_window((438,2110),window=t2)                
                    t2.image=imagxxadfeae1
                    
                    t2.image_create(END, image =imagxxadfeae1)
                    imagxxadfeae1.image=imagxxadfeae1


                def f33():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q33.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t3=Text(canvas,width=24,height=11)
                    canvas.create_window((738,2110),window=t3)                
                    t3.image=imagxxadfeae
                    
                    t3.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae


                def f34():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q34.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t4=Text(canvas,width=24,height=11)
                    canvas.create_window((1038,2110),window=t4)                
                    t4.image=imagxxadfeae
                    
                    t4.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae


                def f35():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q35.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t5=Text(canvas,width=24,height=11)
                    canvas.create_window((1338,2110),window=t5)                
                    t5.image=imagxxadfeae
                    
                    t5.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae

                def f36():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q36.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t6=Text(canvas,width=24,height=11)
                    canvas.create_window((138,2440),window=t6)                
                    t6.image=imagxxadfeae
                    
                    t6.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae                

                def f37():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q37.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t7=Text(canvas,width=24,height=11)
                    canvas.create_window((438,2440),window=t7)                
                    t7.image=imagxxadfeae
                    
                    t7.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae


                def f38():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q38.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t8=Text(canvas,width=24,height=11)
                    canvas.create_window((738,2440),window=t8)                
                    t8.image=imagxxadfeae
                    
                    t8.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae


                def f39():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q39.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t9=Text(canvas,width=24,height=11)
                    canvas.create_window((1038,2440),window=t9)                
                    t9.image=imagxxadfeae
                    
                    t9.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae
                    
                def f40():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q40.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t10=Text(canvas,width=24,height=11)
                    canvas.create_window((1338,2440),window=t10)                
                    t10.image=imagxxadfeae
                    
                    t10.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae

                    
                def f41():
                    filename = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q41.insert(END,filename)
                    
                    imagxxadfeae = Image.open(filename)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t11=Text(canvas,width=24,height=11)
                    canvas.create_window((138,2770),window=t11)
                    t11.image=imagxxadfeae
                    
                    t11.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae




                def f42():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q42.insert(END,filenamee)
                    
                    imagxxadfeae1 = Image.open(filenamee)
                    imagxxadfeae1 = imagxxadfeae1.resize((196,176))
                    imagxxadfeae1 = ImageTk.PhotoImage(imagxxadfeae1)
                    
                    t12=Text(canvas,width=24,height=11)
                    canvas.create_window((438,2770),window=t12)                
                    t12.image=imagxxadfeae1
                    
                    t12.image_create(END, image =imagxxadfeae1)
                    imagxxadfeae1.image=imagxxadfeae1


                def f43():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q43.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t13=Text(canvas,width=24,height=11)
                    canvas.create_window((738,2770),window=t13)                
                    t13.image=imagxxadfeae
                    
                    t13.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae


                def f44():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q44.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t14=Text(canvas,width=24,height=11)
                    canvas.create_window((1038,2770),window=t14)                
                    t14.image=imagxxadfeae
                    
                    t14.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae


                def f45():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q45.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t15=Text(canvas,width=24,height=11)
                    canvas.create_window((1338,2770),window=t15)                
                    t15.image=imagxxadfeae
                    
                    t15.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae

                def f46():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q46.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t16=Text(canvas,width=24,height=11)
                    canvas.create_window((138,3100),window=t16)                
                    t16.image=imagxxadfeae
                    
                    t16.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae                

                def f47():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q47.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t17=Text(canvas,width=24,height=11)
                    canvas.create_window((438,3100),window=t17)                
                    t17.image=imagxxadfeae
                    
                    t17.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae


                def f48():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q48.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t18=Text(canvas,width=24,height=11)
                    canvas.create_window((738,3100),window=t18)                
                    t18.image=imagxxadfeae
                    
                    t18.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae


                def f49():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q49.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t19=Text(canvas,width=24,height=11)
                    canvas.create_window((1038,3100),window=t19)                
                    t19.image=imagxxadfeae
                    
                    t19.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae
                    
                def f50():
                    filenamee = fd.askopenfilename(parent=canvas,title='Choose a file')
                    q50.insert(END,filenamee)
                    
                    imagxxadfeae = Image.open(filenamee)
                    imagxxadfeae = imagxxadfeae.resize((196,176))
                    imagxxadfeae = ImageTk.PhotoImage(imagxxadfeae)
                    
                    t40=Text(canvas,width=24,height=11)
                    canvas.create_window((1338,3100),window=t40)                
                    t40.image=imagxxadfeae
                    
                    t40.image_create(END, image =imagxxadfeae)
                    imagxxadfeae.image=imagxxadfeae

                     
                def d1():
                    phot1_1()
                    tr1.delete(1.0,END)
                    q1.delete(1.0,END)


                def d2():
                    phot2_2()
                    tr2.delete(1.0,END)
                    q2.delete(1.0,END)


                def d3():
                    phot3_3()
                    q3.delete(1.0,END)
                    tr3.delete(1.0,END)

                def d4():
                    phot4_4()
                    q4.delete(1.0,END)
                    tr4.delete(1.0,END)

                def d5():
                    phot5_5()
                    q5.delete(1.0,END)
                    tr5.delete(1.0,END)

                def d6():
                    phot6_6()
                    q6.delete(1.0,END)
                    tr6.delete(1.0,END)
                        
                def d7():
                    phot7_7()
                    q7.delete(1.0,END)
                    tr7.delete(1.0,END)

                def d8():
                    phot8_8()
                    q8.delete(1.0,END)
                    tr8.delete(1.0,END)

                def d9():
                    phot9_9()
                    q9.delete(1.0,END)
                    tr9.delete(1.0,END)

                def d10():
                    phot10_10()
                    q10.delete(1.0,END)
                    tr10.delete(1.0,END)

                def d11():
                    phot11_11()
                    q11.delete(1.0,END)
                    tr11.delete(1.0,END)


                def d12():
                    phot12_12()
                    q12.delete(1.0,END)
                    tr12.delete(1.0,END)


                def d13():
                    phot13_13()
                    q13.delete(1.0,END)
                    tr13.delete(1.0,END)

                def d14():
                    phot14_14()
                    q14.delete(1.0,END)
                    tr14.delete(1.0,END)

                def d15():
                    phot15_15()
                    q15.delete(1.0,END)
                    tr15.delete(1.0,END)

                def d16():
                    phot16_6()
                    q16.delete(1.0,END)
                    tr16.delete(1.0,END)
                        
                def d17():
                    phot17_17()
                    q17.delete(1.0,END)
                    tr17.delete(1.0,END)

                def d18():
                    phot18_18()
                    q18.delete(1.0,END)
                    tr18.delete(1.0,END)

                def d19():
                    phot19_19()
                    q19.delete(1.0,END)
                    tr19.delete(1.0,END)

                def d20():
                    phot20_20()
                    q20.delete(1.0,END)
                    tr20.delete(1.0,END)

                     
                def d21():
                    phot21_21()
                    q21.delete(1.0,END)
                    tr21.delete(1.0,END)


                def d22():
                    phot22_22()
                    q22.delete(1.0,END)
                    tr22.delete(1.0,END)


                def d23():
                    phot23_23()
                    q23.delete(1.0,END)
                    tr23.delete(1.0,END)

                def d24():
                    phot24_24()
                    q24.delete(1.0,END)
                    tr24.delete(1.0,END)
                    
                def d25():
                    phot25_25()
                    q25.delete(1.0,END)
                    tr25.delete(1.0,END)

                def d26():
                    phot26_26()
                    q26.delete(1.0,END)
                    tr26.delete(1.0,END)
                        
                def d27():
                    phot27_27()
                    q27.delete(1.0,END)
                    tr27.delete(1.0,END)

                def d28():
                    phot28_28()
                    q28.delete(1.0,END)
                    tr28.delete(1.0,END)

                def d29():
                    phot29_29()
                    q29.delete(1.0,END)
                    tr29.delete(1.0,END)

                def d30():
                    phot30_30()
                    q30.delete(1.0,END)
                    tr30.delete(1.0,END)

                def d31():
                    phot31_31()
                    q31.delete(1.0,END)
                    tr31.delete(1.0,END)


                def d32():
                    phot32_32()
                    q32.delete(1.0,END)
                    tr32.delete(1.0,END)


                def d33():
                    phot33_33()
                    q33.delete(1.0,END)
                    tr33.delete(1.0,END)

                def d34():
                    phot34_34()
                    q34.delete(1.0,END)
                    tr34.delete(1.0,END)

                def d35():
                    phot35_35()
                    q35.delete(1.0,END)
                    tr35.delete(1.0,END)

                def d36():
                    phot36_36()
                    q36.delete(1.0,END)
                    tr36.delete(1.0,END)
                        
                def d37():
                    phot37_37()
                    q37.delete(1.0,END)
                    tr37.delete(1.0,END)

                def d38():
                    phot38_38()
                    q38.delete(1.0,END)
                    tr38.delete(1.0,END)

                def d39():
                    phot39_39()
                    q39.delete(1.0,END)
                    tr39.delete(1.0,END)

                def d40():
                    phot40_40()
                    q40.delete(1.0,END)
                    tr40.delete(1.0,END)

                def d41():
                    phot41_41()
                    q41.delete(1.0,END)
                    tr41.delete(1.0,END)


                def d42():
                    phot42_42()
                    q42.delete(1.0,END)
                    tr42.delete(1.0,END)


                def d43():
                    phot43_43()
                    q43.delete(1.0,END)
                    tr43.delete(1.0,END)

                def d44():
                    phot44_44()
                    q44.delete(1.0,END)
                    tr44.delete(1.0,END)

                def d45():
                    phot45_45()
                    q45.delete(1.0,END)
                    tr45.delete(1.0,END)

                def d46():
                    phot46_46()
                    q46.delete(1.0,END)
                    tr46.delete(1.0,END)
                        
                def d47():
                    phot47_47()
                    q47.delete(1.0,END)
                    tr47.delete(1.0,END)

                def d48():
                    phot48_48()
                    q48.delete(1.0,END)
                    tr48.delete(1.0,END)

                def d49():
                    phot49_49()
                    q49.delete(1.0,END)
                    tr49.delete(1.0,END)

                def d50():
                    phot50_50()
                    q50.delete(1.0,END)
                    tr50.delete(1.0,END)



                    




            

        # add image icon
                    
                phot1()
                phot2()
                phot3()
                phot4()
                phot5()
                phot6()
                phot7()
                phot8()
                phot9()
                phot10()
                phot11()
                phot12()
                phot13()
                phot14()
                phot15()
                phot16()
                phot17()
                phot18()
                phot19()
                phot20()            
                phot21()
                phot22()
                phot23()
                phot24()
                phot25()
                phot26()
                phot27()
                phot28()
                phot29()
                phot30()
                phot31()
                phot32()
                phot33()
                phot34()
                phot35()
                phot36()
                phot37()
                phot38()
                phot39()
                phot40()
                phot41()
                phot42()
                phot43()
                phot44()
                phot45()
                phot46()
                phot47()
                phot48()
                phot49()
                phot50()

                 
        # DELETE BUTTONS

                de=Button(canvas,text='Delete',command=d1)
                canvas.create_window((138,21),window=de)

                de=Button(canvas,text='Delete',command=d2)
                canvas.create_window((438,21),window=de)

                de=Button(canvas,text='Delete',command=d3)
                canvas.create_window((738,21),window=de)

                de=Button(canvas,text='Delete',command=d4)
                canvas.create_window((1038,21),window=de)

                de=Button(canvas,text='Delete',command=d5)
                canvas.create_window((1338,21),window=de)
                
                de=Button(canvas,text='Delete',command=d6)
                canvas.create_window((138,350),window=de)

                de=Button(canvas,text='Delete',command=d7)
                canvas.create_window((438,350),window=de)

                de=Button(canvas,text='Delete',command=d8)
                canvas.create_window((738,350),window=de)
                
                de=Button(canvas,text='Delete',command=d9)
                canvas.create_window((1038,350),window=de)

                de=Button(canvas,text='Delete',command=d10)
                canvas.create_window((1338,350),window=de)
                
                de=Button(canvas,text='Delete',command=d11)
                canvas.create_window((138,679),window=de)

                de=Button(canvas,text='Delete',command=d12)
                canvas.create_window((438,679),window=de)

                de=Button(canvas,text='Delete',command=d13)
                canvas.create_window((738,679),window=de)

                de=Button(canvas,text='Delete',command=d14)
                canvas.create_window((1038,679),window=de)

                de=Button(canvas,text='Delete',command=d15)
                canvas.create_window((1338,679),window=de)
                
                de=Button(canvas,text='Delete',command=d16)
                canvas.create_window((138,1008),window=de)

                de=Button(canvas,text='Delete',command=d17)
                canvas.create_window((438,1008),window=de)

                de=Button(canvas,text='Delete',command=d18)
                canvas.create_window((738,1008),window=de)
                
                de=Button(canvas,text='Delete',command=d19)
                canvas.create_window((1038,1008),window=de)

                de=Button(canvas,text='Delete',command=d20)
                canvas.create_window((1338,1008),window=de)

                de=Button(canvas,text='Delete',command=d21)
                canvas.create_window((138,1337),window=de)

                de=Button(canvas,text='Delete',command=d22)
                canvas.create_window((438,1337),window=de)

                de=Button(canvas,text='Delete',command=d23)
                canvas.create_window((738,1337),window=de)

                de=Button(canvas,text='Delete',command=d24)
                canvas.create_window((1038,1337),window=de)

                de=Button(canvas,text='Delete',command=d25)
                canvas.create_window((1338,1337),window=de)
                
                de=Button(canvas,text='Delete',command=d26)
                canvas.create_window((138,1666),window=de)

                de=Button(canvas,text='Delete',command=d27)
                canvas.create_window((438,1666),window=de)

                de=Button(canvas,text='Delete',command=d28)
                canvas.create_window((738,1666),window=de)
                
                de=Button(canvas,text='Delete',command=d29)
                canvas.create_window((1038,1666),window=de)

                de=Button(canvas,text='Delete',command=d30)
                canvas.create_window((1338,1666),window=de)
                
                de=Button(canvas,text='Delete',command=d31)
                canvas.create_window((138,1995),window=de)

                de=Button(canvas,text='Delete',command=d32)
                canvas.create_window((438,1995),window=de)

                de=Button(canvas,text='Delete',command=d33)
                canvas.create_window((738,1995),window=de)

                de=Button(canvas,text='Delete',command=d34)
                canvas.create_window((1038,1995),window=de)

                de=Button(canvas,text='Delete',command=d35)
                canvas.create_window((1338,1995),window=de)
                
                de=Button(canvas,text='Delete',command=d36)
                canvas.create_window((138,2324),window=de)

                de=Button(canvas,text='Delete',command=d37)
                canvas.create_window((438,2324),window=de)

                de=Button(canvas,text='Delete',command=d38)
                canvas.create_window((738,2324),window=de)
                
                de=Button(canvas,text='Delete',command=d39)
                canvas.create_window((1038,2324),window=de)

                de=Button(canvas,text='Delete',command=d40)
                canvas.create_window((1338,2324),window=de)

                de=Button(canvas,text='Delete',command=d41)
                canvas.create_window((138,2653),window=de)

                de=Button(canvas,text='Delete',command=d42)
                canvas.create_window((438,2653),window=de)

                de=Button(canvas,text='Delete',command=d43)
                canvas.create_window((738,2653),window=de)

                de=Button(canvas,text='Delete',command=d44)
                canvas.create_window((1038,2653),window=de)

                de=Button(canvas,text='Delete',command=d45)
                canvas.create_window((1338,2653),window=de)
                
                de=Button(canvas,text='Delete',command=d46)
                canvas.create_window((138,2982),window=de)

                de=Button(canvas,text='Delete',command=d47)
                canvas.create_window((438,2982),window=de)

                de=Button(canvas,text='Delete',command=d48)
                canvas.create_window((738,2982),window=de)
                
                de=Button(canvas,text='Delete',command=d49)
                canvas.create_window((1038,2982),window=de)

                de=Button(canvas,text='Delete',command=d50)
                canvas.create_window((1338,2982),window=de)
                



                v1=Button(top,text=' 1 ',bg='dodgerblue',fg='white')
                canvas.create_window((15,120),window=v1)

                v1=Button(top,text=' 2 ',bg='dodgerblue',fg='white')
                canvas.create_window((310,120),window=v1)

                v1=Button(top,text=' 3 ',bg='dodgerblue',fg='white')
                canvas.create_window((613,120),window=v1)

                v1=Button(top,text=' 4 ',bg='dodgerblue',fg='white')
                canvas.create_window((915,120),window=v1)

                v1=Button(top,text=' 5 ',bg='dodgerblue',fg='white')
                canvas.create_window((1215,120),window=v1)
                
                v1=Button(top,text=' 6 ',bg='dodgerblue',fg='white')
                canvas.create_window((15,460),window=v1)

                v1=Button(top,text=' 7 ',bg='dodgerblue',fg='white')
                canvas.create_window((310,460),window=v1)

                v1=Button(top,text=' 8 ',bg='dodgerblue',fg='white')
                canvas.create_window((613,460),window=v1)

                v1=Button(top,text=' 9 ',bg='dodgerblue',fg='white')
                canvas.create_window((915,460),window=v1)

                v1=Button(top,text=' 10 ',bg='dodgerblue',fg='white')
                canvas.create_window((1215,460),window=v1)
                
                v1=Button(top,text=' 11 ',bg='dodgerblue',fg='white')
                canvas.create_window((15,790),window=v1)

                v1=Button(top,text=' 12 ',bg='dodgerblue',fg='white')
                canvas.create_window((310,790),window=v1)

                v1=Button(top,text=' 13 ',bg='dodgerblue',fg='white')
                canvas.create_window((613,790),window=v1)

                v1=Button(top,text=' 14 ',bg='dodgerblue',fg='white')
                canvas.create_window((915,790),window=v1)

                v1=Button(top,text=' 15 ',bg='dodgerblue',fg='white')
                canvas.create_window((1215,790),window=v1)
                
                v1=Button(top,text=' 16 ',bg='dodgerblue',fg='white')
                canvas.create_window((15,1110),window=v1)

                v1=Button(top,text=' 17 ',bg='dodgerblue',fg='white')
                canvas.create_window((310,1110),window=v1)

                v1=Button(top,text=' 18 ',bg='dodgerblue',fg='white')
                canvas.create_window((613,1110),window=v1)

                v1=Button(top,text=' 19 ',bg='dodgerblue',fg='white')
                canvas.create_window((915,1110),window=v1)

                v1=Button(top,text=' 20 ',bg='dodgerblue',fg='white')
                canvas.create_window((1215,1110),window=v1)
                
                v1=Button(top,text=' 21 ',bg='dodgerblue',fg='white')
                canvas.create_window((15,1440),window=v1)

                v1=Button(top,text=' 22 ',bg='dodgerblue',fg='white')
                canvas.create_window((310,1440),window=v1)

                v1=Button(top,text=' 23 ',bg='dodgerblue',fg='white')
                canvas.create_window((613,1440),window=v1)

                v1=Button(top,text=' 24 ',bg='dodgerblue',fg='white')
                canvas.create_window((915,1440),window=v1)

                v1=Button(top,text=' 25 ',bg='dodgerblue',fg='white')
                canvas.create_window((1215,1440),window=v1)
                
                v1=Button(top,text=' 26 ',bg='dodgerblue',fg='white')
                canvas.create_window((15,1770),window=v1)

                v1=Button(top,text=' 27 ',bg='dodgerblue',fg='white')
                canvas.create_window((310,1770),window=v1)

                v1=Button(top,text=' 28 ',bg='dodgerblue',fg='white')
                canvas.create_window((613,1770),window=v1)

                v1=Button(top,text=' 29 ',bg='dodgerblue',fg='white')
                canvas.create_window((915,1770),window=v1)

                v1=Button(top,text=' 30 ',bg='dodgerblue',fg='white')
                canvas.create_window((1215,1770),window=v1)
                
                v1=Button(top,text=' 31 ',bg='dodgerblue',fg='white')
                canvas.create_window((15,2100),window=v1)

                v1=Button(top,text=' 32 ',bg='dodgerblue',fg='white')
                canvas.create_window((310,2100),window=v1)

                v1=Button(top,text=' 33 ',bg='dodgerblue',fg='white')
                canvas.create_window((613,2100),window=v1)

                v1=Button(top,text=' 34 ',bg='dodgerblue',fg='white')
                canvas.create_window((915,2100),window=v1)

                v1=Button(top,text=' 35 ',bg='dodgerblue',fg='white')
                canvas.create_window((1215,2100),window=v1)
                
                v1=Button(top,text=' 36 ',bg='dodgerblue',fg='white')
                canvas.create_window((15,2430),window=v1)

                v1=Button(top,text=' 37 ',bg='dodgerblue',fg='white')
                canvas.create_window((310,2430),window=v1)

                v1=Button(top,text=' 38 ',bg='dodgerblue',fg='white')
                canvas.create_window((613,2430),window=v1)

                v1=Button(top,text=' 39 ',bg='dodgerblue',fg='white')
                canvas.create_window((915,2430),window=v1)

                v1=Button(top,text=' 40 ',bg='dodgerblue',fg='white')
                canvas.create_window((1215,2430),window=v1)
                
                v1=Button(top,text=' 41 ',bg='dodgerblue',fg='white')
                canvas.create_window((15,2760),window=v1)

                v1=Button(top,text=' 42 ',bg='dodgerblue',fg='white')
                canvas.create_window((310,2760),window=v1)

                v1=Button(top,text=' 43 ',bg='dodgerblue',fg='white')
                canvas.create_window((613,2760),window=v1)

                v1=Button(top,text=' 44 ',bg='dodgerblue',fg='white')
                canvas.create_window((915,2760),window=v1)

                v1=Button(top,text=' 45 ',bg='dodgerblue',fg='white')
                canvas.create_window((1215,2760),window=v1)

                v1=Button(top,text=' 46 ',bg='dodgerblue',fg='white')
                canvas.create_window((15,3080),window=v1)

                v1=Button(top,text=' 47 ',bg='dodgerblue',fg='white')
                canvas.create_window((310,3080),window=v1)

                v1=Button(top,text=' 48 ',bg='dodgerblue',fg='white')
                canvas.create_window((613,3080),window=v1)

                v1=Button(top,text=' 49 ',bg='dodgerblue',fg='white')
                canvas.create_window((915,3080),window=v1)

                v1=Button(top,text=' 50 ',bg='dodgerblue',fg='white')
                canvas.create_window((1215,3080),window=v1)


                    
                vbar=Scrollbar(frame,orient=VERTICAL,width=25)
                vbar.pack(side=RIGHT,fill=Y)
                vbar.config(command=canvas.yview)
                canvas.bind('<4>', lambda event : canvas.yview('scroll', -1, 'units'))
                canvas.bind('<5>', lambda event : canvas.yview('scroll', 1, 'units'))
         
           


                canvas.config(width=1550,height=1250)
                canvas.config( yscrollcommand=vbar.set)
                canvas.pack(side=LEFT,expand=True,fill=BOTH)

                
                    

          
            b=Button(SEARCHPATIENT,text='MEDICAL  DETAILS',font=('Bahnschrift Light SemiCondensed',14),bg='DARKTURQUOISE',foreground='ghostwhite',command=selectItem)
            b.place(x=800,y=690)

            bs=Button(SEARCHPATIENT,text='OTHER  DETAILS',font=('Bahnschrift Light SemiCondensed',14),bg='DARKTURQUOISE',foreground='ghostwhite',command=selectItem1)
            bs.place(x=1050,y=690)
 
            bqs=Button(SEARCHPATIENT,text='  SORT  ',font=('Bahnschrift Light SemiCondensed',12),bg='DODGERBLUE',foreground='white',command=myy)
            bqs.place(x=430,y=182)

            baqs=Button(SEARCHPATIENT,text='REFRESH',font=('Bahnschrift Light SemiCondensed',12),command=refresh)
            baqs.place(x=1450,y=708)
            
            bl=Button(SEARCHPATIENT,text='SEARCH',font=('Bahnschrift Light SemiCondensed',12),bg='DODGERBLUE',foreground='white',command=cc_sn)
            bl.place(x=430,y=65)

            ASSX=Button(SEARCHPATIENT,text=' SAVE  AS  DOCX ',font=('Bahnschrift Light SemiCondensed',14),bg='DODGERBLUE',foreground='white',command=word)
            ASSX.place(x=30,y=350)

            button1 = Button(SEARCHPATIENT, text=" DELETE " , font=('Bahnschrift Light SemiCondensed',14),bg='DODGERBLUE',fg='WHITE',command=delete)
            button1.place(x=240,y=350)

            button1 = Button(SEARCHPATIENT, text=" UPDATE " , font=('Bahnschrift Light SemiCondensed',14),bg='DODGERBLUE',fg='WHITE',command=tabb)
            button1.place(x=390,y=350)

            sorta = StringVar(SEARCHPATIENT)
            sorta.set("Date")
            waaasa = OptionMenu(SEARCHPATIENT,sorta, "Serial_Number","Patient_Id","Name", "Date","Age","Sex")
            waaasa.config(fg='BLACK',font='arial')
            waaasa.place(x=20,y=180)

            sortaa = StringVar(SEARCHPATIENT)
            sortaa.set("Asc")
            waaaasa = OptionMenu(SEARCHPATIENT,sortaa, "Asc", "Desc")
            waaaasa.config(fg='BLACK',font='arial')
            waaaasa.place(x=240,y=180) 

            saortaa = StringVar(SEARCHPATIENT)
            saortaa.set("Serial_Number")
            awaaaasa = OptionMenu(SEARCHPATIENT,saortaa,  "Name","Serial_Number",'Patient_Id','Date','Age','Sex','Contact_Detail','Patient_Status')
            awaaaasa.config(font='arial 16')
            awaaaasa.place(x=20,y=60)

            
            sx=Entry(SEARCHPATIENT,font='ARIAL 10')
            sx.place(x=240,y=70)
            sx.bind("<Return>", (lambda event:cc_sn()))
             

            pizlimage = Image.open("C:/Users/Acer/Desktop/pms files/gg.png")
            pizlage = pizlimage.resize((120, 110))
            img = ImageTk.PhotoImage(pizlage)

        
     

       
  
       
            
        showrecords_serialsno()



#########################################################################################################
####################################          UPDATE      TAB       #####################################
#########################################################################################################


        


            



#########################################################################################################
################################        APPOINTMENTS      TAB   #########################################
#########################################################################################################


         
        gsif1saad = Image.open("C:/Users/Acer/Desktop/pms files/ImageBlur (5).png")    #  BG IMAGE
        gsif1saad = gsif1saad.resize((1520, 765))
        iqmageed = ImageTk.PhotoImage(gsif1saad)
        label2d = tk.Label(APPOINTMENTS,image=iqmageed)
        label2d.image=iqmageed 
        label2d.pack()
        iqmageed.image = iqmageed

    
        
























        

#########################################################################################################
#########################################################################################################
#########################################################################################################
#########################################################################################################
#########################################################################################################
#########################################################################################################
#########################################################################################################


        
    else:                                          # IF  LOGIN  PASSWORD  INCORRECT  THEN  MESSAGEBOX
        MessageBox.showerror("LOGIN ERROR","Incorrect Password",parent=root)
        
    con.close()


##############################################  ROOT  WINDOW  ##########################################

    
root=tk.Tk()
root.geometry("1550x1250")                                                 # ROOT  WINDOW  FORMED
root.title('')
root.configure(background='WHITE' )
root.state('zoomed')






pilimage = Image.open("C:/Users/Acer/Desktop/pms files/ImageBlur.png")     # BG IMAGE FOR ROOT WINDOW
pilage = pilimage.resize((1600, 780))
image = ImageTk.PhotoImage(pilage)
label2 = tk.Label(root,image=image)
label2.pack(padx=5, pady=5)


                                      
buttonpms = Button(root, text="                                     PATIENT  MANAGEMENT  SYSTEM                                     ",font=("ROBOTO BOLD",30),bg='darkturquoise',state='disabled',disabledforeground='WHITE')
buttonpms.place(x=10,y=10)



pilage = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/user icon.png")      #USERNAME  ICON
photoimage = pilage.subsample(35, 35) 
Button(root, text = '', image = photoimage,state='disabled').place(x=500,y=300)



DROPDOWN = tk.StringVar(root)                                   # DROP DOWN MENU FOR USERNAME
DROPDOWN.set("Dr. Ravjot Singh")
DDOptions = OptionMenu(root, DROPDOWN, "Dr. Ravjot Singh", "        User       ")
DDOptions.config(bg = "WHITE",fg='BLACK',font='arial')
DDOptions.place(x=610,y=315)





pilaage = PhotoImage(file = "C:/Users/Acer/Desktop/pms files/download.png")       # PASSWORD  ICON
photoiamage = pilaage.subsample(3, 3) 
Button(root, text = '', image = photoiamage,state='disabled').place(x=490,y=400)


p=StringVar
passw=Entry(root,bg='white', show="*",font=("ROBOTO",15),textvariable=p)          # PASSWORD ENTRY BOX
passw.place(x=610,y=415)
passw.bind("<Return>", (lambda event: login()))



                                              # LOGIN  BUTTON THAT OPENS TOPLEVEL WINDOW
buttonlg = Button(root, text=" Login ",font=("ROBOTO BOLD",15),bg='darkturquoise',fg='white',command=login)
buttonlg.place(x=680,y=490)


                                              # PASSWORD  RESET   BUTTON  [HEADING] 
buttonPR = Button(root, text=" Reset Password ",font=("ROBOTO BOLD",15),bg='blue',state='disabled',disabledforeground='white')
buttonPR.place(x=1220,y=120)



                                              # GET  OTP  ON  MAIL  [BUTTON]                                          
buttonOTP = Button(root, text=" Get OTP ",font=("ROBOTO BOLD",15),bg='dodgerblue',fg='white',command=mail)
buttonOTP.place(x=1120,y=190)



OTPTYPE=Entry(root,font=("ROBOTO BOLD",15))             #  OTP  TYPE  ENTRY  BOX
OTPTYPE.place(x=1260,y=195)
OTPTYPE.bind("<Return>", (lambda event: top_for_FP()))


                                    # SUBMIT  BUTTON  TO  VERIFY  OTP  AND  OPEN  NEW  PASWORD  TOPLEVEL
buttonsubmit = Button(root, text=" Submit ",font=("ROBOTO BOLD",15),bg='dodgerblue',fg='white',command=top_for_FP)
buttonsubmit.place(x=1300,y=240)
          
root.mainloop()



# SINGLE COTES AND MORE THAN 1 SPACE IN FILENAME NOT TO BE USED. USE _ FOR FILE OR 1 SPACE
# PASSWORD IS CASE SENSITIVE
# PATIENT DATA NOT CASE SENSITIVE
# WHILE ADDING IMAGE PATHS IN TEXTBOX IN ADD RECORD, 1 LINE 1 PATH














